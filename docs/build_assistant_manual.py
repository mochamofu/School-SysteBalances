# -*- coding: utf-8 -*-
"""積立金入力アシスタント 手順書（PDF元のdocx）を生成するスクリプト

docs/build_manual.py と同じヘルパーを流用する。
生成: docs/assistant_manual.docx → LibreOfficeでPDF化して
      docs/入力アシスタント_手順書.pdf として保存する。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"


def set_japanese_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text, size=20, color=(0x20, 0x38, 0x64)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_japanese_font(r, size=size, bold=True, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_japanese_font(r, size=15, bold=True, color=(0x20, 0x38, 0x64))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("■ " + text)
    set_japanese_font(r, size=12.5, bold=True, color=(0x44, 0x72, 0xC4))
    return p


def add_body(doc, text, size=10.5, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_japanese_font(r, size=size, bold=bold, color=color)
    return p


def add_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"Step {num}　")
    set_japanese_font(r, size=10.5, bold=True, color=(0x20, 0x38, 0x64))
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p


def add_note(doc, text, label="⚠ 注意"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FFF2CC')
    pPr.append(shd)
    r = p.add_run(f"{label}: ")
    set_japanese_font(r, size=10, bold=True, color=(0xC0, 0x50, 0x00))
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r2 = p.add_run(line)
        set_japanese_font(r2, size=10)
    return p


def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("☐ ")
    set_japanese_font(r, size=10.5, bold=True)
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_japanese_font(r, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '203864')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_japanese_font(r, size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ============ 表紙 ============
for _ in range(6):
    doc.add_paragraph()
add_title(doc, "積立金会計 入力アシスタント", size=24)
add_title(doc, "導入・操作手順書", size=18, color=(0x44, 0x72, 0xC4))
doc.add_paragraph()
add_title(doc, "～ 「令和○年度生積立金」への入力作業を、ボタン操作だけにする ～",
          size=11, color=(0x60, 0x60, 0x60))
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("対象ファイル\n積立金入力アシスタント.xlsx（→保存後 .xlsm）\nvba/A00〜A07 の8モジュール\n練習用ダミーファイル3点")
set_japanese_font(r, size=11, color=(0x20, 0x38, 0x64))
doc.add_page_break()

# ============ 0. これは何か ============
add_h1(doc, "0．このパッケージでできること")
add_body(doc, "実物の「令和○年度生積立金」（学校徴収金個人別管理票）はそのまま使い続けます。"
              "このアシスタントは、そのファイルの隣に置く「入力の代行係」です。"
              "マスターの行・列・数式には一切触らず、決まった場所に安全に書き込みます。")
add_table(doc,
          ["これまでの作業", "アシスタント導入後"],
          [["クラス替えのたび、名簿を見ながら320人分の組・番号を手で修正", "名簿を貼り付けて2クリック（氏名で自動照合）"],
           ["支出のたび、同じ金額を315人分入力し例外を1人ずつ修正", "件名と単価を入力、例外の生徒だけ表に書いて1クリック"],
           ["振替結果を見ながら入金額を1人ずつ入力", "未納者の精算番号だけ書けば残り全員に自動入力"],
           ["支出承認書・収入承認書に同じ情報を二重記入", "一括入力と同時に承認書が自動で埋まる"],
           ["決算書は電卓で項目合計を出して直打ち", "1クリックで全項目の人数・単価・総額を一覧化"],
           ["精算書は紙への一括印刷のみ", "生徒別PDF（組-番号_氏名.pdf）の一括保存も可能"]],
          col_widths=[8.2, 8.2])
add_note(doc, "どの書き込み機能も、実行した瞬間にマスターのバックアップ（日時付きコピー）を自動で作ります。"
              "失敗してもバックアップから戻せます。", label="✔ 安心")

# ============ 1. 導入 ============
add_h1(doc, "1．導入（最初に1回だけ・約10分）")
add_h2(doc, "1-1．ファイルを学校PCに置く")
add_step(doc, 1, "「積立金入力アシスタント.xlsx」と vba フォルダ（A00〜A07.bas の8個）を学校PCの同じフォルダにコピー")
add_step(doc, 2, "マスター（令和○年度生積立金.xlsx）の場所（フルパス）を控えておく")
add_h2(doc, "1-2．マクロを組み込む")
add_step(doc, 3, "アシスタントブックを開き、Alt + F11 でVBAエディタを開く")
add_step(doc, 4, "「ファイル」→「ファイルのインポート」で A00〜A07 の8個を順番に全部インポート")
add_step(doc, 5, "「ファイル」→「閉じてMicrosoft Excelへ戻る」")
add_step(doc, 6, "「名前を付けて保存」→ 種類を「Excelマクロ有効ブック(*.xlsm)」にして保存")
add_step(doc, 7, "Alt + F8 → 「初期設定」を実行 → メニューシートにボタン①〜⑨が並ぶ")
add_h2(doc, "1-3．設定シートを埋める")
add_step(doc, 8, "「設定」シートのC3にマスターファイルのフルパス、C5に年度（例：7）を入力")
add_body(doc, "C4（バックアップ先）とC6（PDF保存先）は空欄でOKです。マスターと同じ場所に自動でフォルダが作られます。")
add_note(doc, "本物のマスターを触る前に、必ず「練習用_令和X年度生積立金.xlsx」（架空の氏名のダミー）をC3に設定して、"
              "2章以降の操作を一度リハーサルしてください。")

# ============ 2. クラス替え ============
add_h1(doc, "2．毎年4月：クラス替え名簿の取込")
add_step(doc, 1, "掲示用名簿（新クラス発表のExcel）を開き、シート全体をコピー（Ctrl+A → Ctrl+C）")
add_step(doc, 2, "アシスタントの「名簿貼付」シートのA1に貼り付け")
add_step(doc, 3, "メニューの「① 名簿を解析して照合する」を押す")
add_body(doc, "→「名簿一覧」シートに全生徒の 組・番号・氏名 と照合結果が並びます。")
add_table(doc,
          ["D列の表示", "意味", "やること"],
          [["一致", "マスターに同じ氏名が1人だけいた", "何もしない（自動で精算番号が入る）"],
           ["見つからず", "マスターに同じ氏名がない（転入生・改姓・表記ゆれ）", "E列に正しい精算番号を手入力"],
           ["複数候補（同姓同名）", "同じ読みの生徒が2人以上いる", "E列にどちらの精算番号かを手入力"]],
          col_widths=[3.8, 6.6, 6.0])
add_step(doc, 4, "E列がすべて埋まったら「② クラス替えをマスターに反映する」を押し、新しい学年（例：2）を入力")
add_body(doc, "→ マスターの 学年・組・番号 が一括更新されます。精算番号と氏名の並び順は変わりません（精算書の数式を守るため）。")
add_h2(doc, "新入生の場合（マスターがまだ空の年）")
add_body(doc, "Step 1〜3 は同じ。最後に「③ 新入生としてマスターに登録する」を押すと、名簿の順に精算番号1から登録されます。")

# ============ 3. 支出 ============
add_h1(doc, "3．支出のたび：一括入力＋承認書")
add_step(doc, 1, "「支出入力」シートの黄色いセルを埋める（支出No・件名・支払先・日付・一人あたり金額・対象）")
add_step(doc, 2, "例外の生徒だけ下の例外表に書く")
add_table(doc,
          ["例外の書き方", "意味"],
          [["精算番号45 ／ 金額 0", "45番は対象外（転退学・給付型奨学金など）"],
           ["精算番号12 ／ 金額 -1040", "12番に1,040円を返金（マイナスで記録）"],
           ["精算番号300 ／ 金額 260", "300番だけ260円（途中参加の半額など）"]],
          col_widths=[6.0, 10.4])
add_step(doc, 3, "メニューの「④ 支出をマスターへ一括入力」を押す")
add_body(doc, "→ 在籍生徒全員に金額が入り、例外だけ上書きされます。件名・日付・No.も6〜8行目に自動で入ります。"
              "対象人数と合計金額が表示されるので、承認書の金額と一致するか確認してください。")
add_step(doc, 4, "「支出承認書」シートが自動で埋まっているので、確認して印刷（黄色いセルだけ手で補記）")
add_note(doc, "「全員」は氏名が入っている生徒だけが対象です。転退学で氏名を消した行には書き込まれません。\n"
              "すでに金額が入っている支出Noを指定すると、人数を示して上書き確認が出ます。")

# ============ 4. 収入 ============
add_h1(doc, "4．口座振替のたび：収入の一括入力")
add_step(doc, 1, "「⑥ 収入枠の一覧を表示」で空いている枠Noを確認（使用中の枠と人数が一覧で出ます）")
add_step(doc, 2, "「収入入力」シートの黄色いセルを埋める（収入枠No・件名・日付・一人あたり金額）")
add_step(doc, 3, "銀行の振替結果で「振替できなかった生徒」の精算番号だけを未納者表に書く")
add_step(doc, 4, "メニューの「⑤ 収入をマスターへ一括入力」を押す")
add_body(doc, "→ 未納者以外の在籍生徒全員に金額が入ります。未納者は空欄のままなので、"
              "マスターH列の「未納」の印が自動で立ち、督促状づくりにそのまま使えます。"
              "「収入承認書」シートも自動で埋まります。")

# ============ 5. 年度末 ============
add_h1(doc, "5.年度末：決算集計・点検・精算書PDF")
add_h2(doc, "5-1．決算集計（⑦）")
add_body(doc, "「⑦ 決算用の集計を実行」を押すと、「決算集計」シートに使用中の全項目について "
              "件名・日付・対象人数・一人あたり（最も多くの生徒に入っている金額）・執行総額 が並びます。"
              "この数字をマスターの決算書（1年決算書シート）へ転記すれば、電卓での手集計は不要です。")
add_h2(doc, "5-2．整合性チェック（⑧）")
add_body(doc, "「⑧ マスターの整合性をチェック」は次の4点を点検し、「チェック結果」シートに書き出します。")
add_table(doc,
          ["点検", "見つけられる問題"],
          [["収入−支出＝残金 の確認", "生徒ごと・全体の金額のズレ、行の挿入削除による集計崩れ"],
           ["氏名空欄なのに金額あり", "転退学処理のし忘れ"],
           ["金額あるのに件名なし", "6行目の件名の書き忘れ（精算書に空欄項目が出る原因）"],
           ["未納者一覧", "督促状を出すべき生徒のリストアップ"]],
          col_widths=[5.5, 10.9])
add_h2(doc, "5-3．精算書の一括PDF（⑨）")
add_body(doc, "「⑨ 精算書を一括PDF保存」で範囲（例：1-320）を指定すると、生徒1人につき1つのPDF"
              "（ファイル名：組-番号_氏名.pdf）が保存されます。転退学者1人分だけの出力（例：45）もできます。"
              "紙への一括印刷は、従来どおりマスターの一括印刷ボタンも使えます。")

# ============ 6. 安全のしくみ ============
add_h1(doc, "6．安全のしくみ（読んでおくと安心）")
add_table(doc,
          ["しくみ", "内容"],
          [["自動バックアップ", "書き込み系の機能は実行のたびに「バックアップ」フォルダへ日時付きコピーを作ってから書き込む"],
           ["構造チェック", "マスターの「データ」シートの形が想定と違うファイルには書き込まない（誤爆防止）"],
           ["上書き確認", "すでに金額が入っている列には、人数を示して確認してから書き込む"],
           ["触らない場所", "精算番号・氏名・数式列（H,I,BA〜BC,FC,FD）・行や列の挿入削除は一切行わない"]],
          col_widths=[4.0, 12.4])
add_body(doc, "元に戻したいとき：マスターを閉じて、「バックアップ」フォルダの直近のコピーを元の名前に戻すだけです。")

# ============ 7. 困ったとき ============
add_h1(doc, "7．うまく動かないときは")
add_table(doc,
          ["症状", "原因と対処"],
          [["ボタンがない", "Alt+F8 →「初期設定」を実行するとメニューにボタンが並びます"],
           ["「設定が足りません」と出る", "「設定」シートC3にマスターのフルパスを入力してください"],
           ["「構造エラー」と出る", "指定したファイルが積立金マスターではない、または行や列を挿入してしまっています。バックアップから戻してください"],
           ["「〜組が見つからない」と出る", "掲示用名簿を見出し（1年1組 など）ごと貼り付けてください"],
           ["マクロが実行できない", "ファイルが .xlsm で保存されているか、「コンテンツの有効化」を押したかを確認"],
           ["見つからずが多い", "名簿とマスターで氏名の表記（かな/漢字）が違う年は、E列の精算番号手入力で確定してください"]],
          col_widths=[5.5, 10.9])

# ============ 8. 本番前チェックリスト ============
add_h1(doc, "8．本番前チェックリスト")
for item in [
    "練習用ダミーファイルで ①〜⑨ を一通り実行し、エラーなく動いた",
    "練習用マスターの「バックアップ」フォルダに日時付きコピーができていた",
    "本物のマスターの控えを手動でも1部コピーして別の場所に置いた",
    "「設定」シートC3を本物のマスターのパスに書き換えた",
    "最初の一括入力のあと、対象人数と合計金額が承認書・銀行資料と一致することを確認した",
    "チェック機能（⑧）で⚠が出ないことを確認した",
]:
    add_check(doc, item)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("以上")
set_japanese_font(r, size=11, bold=True)

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assistant_manual.docx')
doc.save(out_path)
print(f"saved {out_path}")
