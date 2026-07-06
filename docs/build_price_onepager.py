# -*- coding: utf-8 -*-
"""お客様配布用の価格・サービス比較資料（A4・稟議添付用）を生成する。

生成: scratchpad に docx を作り、LibreOffice で PDF 化して
      docs/価格とサービスのご説明.pdf として保存する。
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)


def set_jp(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10, bold=False, color=None, space_after=4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=12, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths, header_fill="203864", size=8.5,
          highlight_row=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_jp(r, size=size, bold=True, color=(0xFF, 0xFF, 0xFF))
        shade(cell, header_fill)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            hl = (highlight_row is not None and i == highlight_row)
            r = cell.paragraphs[0].add_run(str(v))
            set_jp(r, size=size, bold=hl)
            if hl:
                shade(cell, "FFF2CC")
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("積立金会計 入力アシスタント　価格とサービスのご説明")
    set_jp(r, size=15, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("― ご予算の検討・稟議資料としてお使いください ―")
    set_jp(r, size=9)

    para(doc,
         "積立金は保護者の皆さまからの預かり金です。本サービスの費用は「便利な道具の代金」ではなく、"
         "預かり金会計の事故防止と、担当の先生おひとりに業務が集中しない体制づくりに対するものです。"
         "以下、何と比べていただくのが適切かを整理しました。", size=9.5, space_after=8)

    h1(doc, "1. いまのやり方を続けた場合にかかっているもの")
    for t in [
        "・入力・照合・転記の手作業：月およそ5〜10時間（人件費に換算すると月1〜2万円相当）",
        "・誤請求・未納見落としが1件起きたときの対応：再計算・確認・保護者への説明に複数名で数時間＋信用面の影響",
        "・入力手順が担当者おひとりの経験に依存（異動・退職・急なお休みで業務が止まるリスク）",
        "・マスターに組み込まれた既存マクロは1999年製で、保守できる方が既にいらっしゃいません",
    ]:
        para(doc, t, size=9.5, space_after=2)
    para(doc, "→ 「追加費用ゼロ」に見えて、毎月1〜2万円相当の人件費と、金額に換算しにくいリスクを負担し続けている状態です。",
         size=9.5, bold=True, space_after=8)

    h1(doc, "2. 選択肢の比較")
    table(
        doc,
        ["選択肢", "費用のめやす", "いまのマスターを\nそのまま使えるか", "保守・年度更新"],
        [
            ["このまま手入力を続ける", "追加費用0円\n（人件費 月1〜2万円相当）", "―", "なし（担当者依存のまま）"],
            ["市販Excelテンプレート（買い切り）", "数千円", "×　様式が合わず手直しが必要", "なし"],
            ["校務支援・徴収金管理システム", "月2〜3万円＋初期1〜30万円\n生徒単価型は約300円/人・月\n（320名で月約96,000円）", "×　システムへの移行・\nデータ載せ替えが前提", "ベンダー保守あり"],
            ["本サービス（スタンダード）", "月19,800円（税別）・初期0円\n年払い198,000円（2ヶ月分お得）", "○　行も列も1つも変えず\nそのまま使い続けられます", "保守・年度更新・再レクチャー\nすべて月額に含む"],
        ],
        widths=[4.6, 5.2, 4.6, 4.2],
        highlight_row=3,
    )
    para(doc, "生徒1人あたりに換算すると月約62円（320名の場合）。校務支援システムの生徒単価相場（約300円）の約5分の1です。",
         size=9, space_after=8)

    h1(doc, "3. 月額に含まれるもの（ファイルの代金ではなく、伴走の費用です）")
    para(doc, "お渡しする実体は、いまお使いの積立金マスターの「隣に置く」マクロ付きExcelブック1つです。"
              "サーバーもインストールも不要で、月額はその後の伴走に対する費用です。", size=9.5, space_after=4)
    for t in [
        "① 入力自動化一式（ボタン15個）… 名簿照合・クラス替え反映／支出・収入の一括入力／振替結果の自動照合／承認書の自動作成／決算集計／精算書PDF／業者別集計と翌年度への引き継ぎ",
        "② 事故防止装置 … 書き込み前の自動バックアップ・誤ったファイルへの書き込み防止・収支の整合性チェック",
        "③ エラー対応窓口 … 画面の写真をお送りいただくだけ。修正版の差し替えまで対応",
        "④ 年度更新の立ち会い … 新年度の項目引き継ぎ（残す項目に○を付けるだけ）を毎年サポート",
        "⑤ ご担当者の交代時の再レクチャー、様式・項目が変わった際のマクロ改修",
    ]:
        para(doc, t, size=9.5, space_after=2)
    para(doc, "", size=4, space_after=2)

    h1(doc, "4. 品質の裏付け（検証記録をご提示できます）")
    for t in [
        "・実物と同形式のファイルで、支出全項目×全生徒分の金額転記を検証し、1円の差もないことを確認済み",
        "・全校320名規模の試験データで、名簿照合320名一致／口座振替321件の照合（未納5件・不明口座1件を正しく検出）",
        "・わざと異常なデータを流し込むテストを15項目実施し、発見された不具合はすべて修正済み（検証ログあり）",
    ]:
        para(doc, t, size=9.5, space_after=2)

    h1(doc, "5. 安心してお試しいただくために")
    for t in [
        "・30日間の無料トライアル：実ファイルの「コピー」でお使いいただき、手入力の結果と突き合わせてください。1円でも合わなければご契約は不要です。",
        "・初期費用0円・月単位のご契約：いつでも解約でき、解約後もいまのマスターは何も変わらず残ります（囲い込みがありません）。",
    ]:
        para(doc, t, size=9.5, space_after=2)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("お問い合わせ：株式会社ココロラボ　info@cocorolab.co.jp")
    set_jp(r, size=9.5, bold=True)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
        check=True, env={**os.environ, "HOME": tempfile.gettempdir()},
    )
    produced = os.path.splitext(docx_path)[0] + ".pdf"
    os.replace(produced, pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "price_onepager.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "価格とサービスのご説明.pdf"))
