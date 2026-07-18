# -*- coding: utf-8 -*-
"""かんたん導入ガイド（完全図解版）を生成する。

「1クリック＝1手順＝1つの絵」で、パソコンが苦手な人でも
上から順になぞるだけで導入が終わることを目指した詳細版。
USBを差すところから、右クリックメニュー、Excelのどのボタンを
押すかまで、全手順に画面の絵と赤い丸印を付ける。

生成: docs/04_オンライン導入/かんたん導入ガイド_完全図解版.pdf
"""
import os
import math
import subprocess
import tempfile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PDF = os.path.join(HERE, "..", "04_オンライン導入", "かんたん導入ガイド_完全図解版.pdf")

FONT_PATH = "/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf"
FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GREEN = (0x1E, 0x7A, 0x3C)
NAVY_S = "#203864"
BLUE_S = "#2F5FA8"
AMBER_S = "#C58F00"
YELLOW_S = "#FFF3B0"
RED_S = "#D02020"
GRAY_S = "#666666"
LGRAY_S = "#E8EDF5"
WIN_S = "#F3F3F3"     # Windows背景
SEL_S = "#CCE4FF"     # 選択色

_workdir = None


def F(size):
    return ImageFont.truetype(FONT_PATH, size)


def _save(img, name):
    path = os.path.join(_workdir, name)
    img.save(path)
    return path


def txt(d, x, y, s, size, fill=NAVY_S, stroke=0):
    d.text((x, y), s, font=F(size), fill=fill, stroke_width=stroke, stroke_fill=fill)


def ctext(d, cx, y, s, size, fill=NAVY_S, stroke=0):
    w = d.textlength(s, font=F(size))
    d.text((cx - w / 2, y), s, font=F(size), fill=fill, stroke_width=stroke, stroke_fill=fill)


def circle(d, box, width=10):
    """クリック位置を示す赤い丸"""
    d.ellipse(box, outline=RED_S, width=width)


def arrow(d, x1, y1, x2, y2, color=RED_S, width=9):
    d.line((x1, y1, x2, y2), fill=color, width=width)
    ang = math.atan2(y2 - y1, x2 - x1)
    L = 30
    for da in (2.6, -2.6):
        d.line((x2, y2, x2 - L * math.cos(ang + da), y2 - L * math.sin(ang + da)),
               fill=color, width=width)


def window(d, box, title, title_fill=NAVY_S):
    """タイトルバー付きウィンドウ"""
    x1, y1, x2, y2 = box
    d.rectangle(box, fill="white", outline="#999999", width=3)
    d.rectangle((x1, y1, x2, y1 + 52), fill=title_fill)
    txt(d, x1 + 20, y1 + 12, title, 26, "white")
    # 閉じるボタン等
    for i, ch in enumerate("―□×"):
        d.text((x2 - 130 + i * 42, y1 + 10), ch, font=F(26), fill="white")


def context_menu(d, x, y, items, highlight_idx, w=360):
    """右クリックメニュー。highlight_idx の項目を黄色にする"""
    row_h = 56
    d.rectangle((x, y, x + w, y + row_h * len(items)), fill="white", outline="#888888", width=3)
    for i, it in enumerate(items):
        yy = y + i * row_h
        if i == highlight_idx:
            d.rectangle((x + 3, yy + 3, x + w - 3, yy + row_h - 3), fill=YELLOW_S)
        txt(d, x + 28, yy + 12, it, 28, NAVY_S if i == highlight_idx else "#444444")
    return (x, y + highlight_idx * row_h, x + w, y + (highlight_idx + 1) * row_h)


def folder_icon(d, x, y, w=170, label=None, label_size=22):
    h = int(w * 0.72)
    d.polygon([(x, y + 18), (x + w * 0.38, y + 18), (x + w * 0.46, y), (x + w * 0.46, y),
               (x + w * 0.46, y + 18)], fill="#FFD35C")
    d.polygon([(x, y + 18), (x + w * 0.38, y + 18), (x + w * 0.46, y + 34), (x + w, y + 34),
               (x + w, y + h), (x, y + h)], fill="#FFD35C", outline=AMBER_S, width=4)
    if label:
        ctext(d, x + w / 2, y + h + 8, label, label_size)


def excel_file_icon(d, x, y, w=120, label=None):
    h = int(w * 1.25)
    d.rectangle((x, y, x + w, y + h), fill="white", outline="#1D6F42", width=5)
    d.rectangle((x, y, x + int(w * 0.42), y + h), fill="#1D6F42")
    d.text((x + 8, y + h // 2 - 26), "X", font=F(52), fill="white")
    if label:
        ctext(d, x + w / 2, y + h + 8, label, 22)


# ==================================================================
# 手順ごとの図
# ==================================================================
def s01_insert():
    img = Image.new("RGB", (1500, 480), "white")
    d = ImageDraw.Draw(img)
    # ノートPC側面
    d.polygon([(300, 260), (1050, 260), (1130, 380), (220, 380)], fill=LGRAY_S, outline=NAVY_S, width=5)
    d.rectangle((300, 120, 1050, 260), fill="white", outline=NAVY_S, width=5)
    ctext(d, 675, 170, "パソコン", 34)
    # USBポート
    d.rectangle((1090, 300, 1150, 336), fill="black")
    # USB本体
    d.rounded_rectangle((1230, 288, 1420, 350), radius=12, fill="white", outline=NAVY_S, width=5)
    d.rectangle((1180, 300, 1230, 338), fill=LGRAY_S, outline=NAVY_S, width=4)
    d.ellipse((1260, 270, 1330, 340), fill=BLUE_S, outline=NAVY_S, width=4)
    ctext(d, 1295, 284, "①", 42, "white", 1)
    arrow(d, 1180, 318, 1160, 318)
    circle(d, (1060, 270, 1180, 366))
    ctext(d, 750, 420, "パソコンの横（または後ろ）の差し込み口に、USB①をそのまま差します", 30, RED_S, 1)
    return _save(img, "s01.png")


def s02_explorer():
    img = Image.new("RGB", (1500, 560), "white")
    d = ImageDraw.Draw(img)
    # タスクバー
    d.rectangle((60, 420, 1440, 500), fill="#E8E8E8", outline="#AAAAAA", width=3)
    # スタートボタン風
    d.rectangle((90, 438, 150, 482), fill=BLUE_S)
    # フォルダアイコン（エクスプローラー）
    folder_icon(d, 210, 430, 78)
    circle(d, (185, 415, 315, 505))
    txt(d, 340, 445, "← タスクバーの「黄色いフォルダ」をクリック", 28, RED_S)
    # 開いたウィンドウ
    window(d, (500, 40, 1440, 390), "エクスプローラー")
    txt(d, 540, 120, "PC", 28, GRAY_S)
    d.rectangle((530, 170, 1410, 240), fill=SEL_S, outline=BLUE_S, width=3)
    d.rounded_rectangle((560, 185, 640, 226), radius=8, fill="white", outline=NAVY_S, width=4)
    txt(d, 660, 186, "USB ドライブ (D:) など", 30)
    circle(d, (520, 160, 1100, 252))
    ctext(d, 970, 300, "「USBドライブ」をダブルクリックして開く", 28, RED_S, 1)
    return _save(img, "s02.png")


def s03_copy():
    img = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(img)
    window(d, (60, 40, 820, 580), "USB ドライブ (D:)")
    folder_icon(d, 130, 150, 190, "1_積立金入力アシスタント", 24)
    txt(d, 100, 380, "① このフォルダを", 28, RED_S)
    txt(d, 100, 420, "　 右クリック", 28, RED_S)
    menu_box = context_menu(d, 440, 260, ["開く", "コピー", "切り取り", "名前の変更"], 1)
    circle(d, (menu_box[0] - 15, menu_box[1] - 10, menu_box[2] + 15, menu_box[3] + 10))
    txt(d, 420, 510, "② 出てきたメニューの「コピー」をクリック", 30, RED_S)
    # 右: 説明
    txt(d, 880, 120, "マウスの右側のボタンを", 32)
    txt(d, 880, 170, "カチッと押すと", 32)
    txt(d, 880, 220, "メニューが出ます。", 32)
    txt(d, 880, 320, "※「コピー」が絵（アイコン）だけの", 26, GRAY_S)
    txt(d, 880, 360, "　場合もあります（四角が2枚の絵）", 26, GRAY_S)
    return _save(img, "s03.png")


def s04_paste():
    img = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((60, 40, 820, 580), fill="#EAF2E6", outline="#2E7D32", width=4)
    ctext(d, 440, 60, "デスクトップ（何もないところ）", 30, "#2E7D32")
    txt(d, 110, 140, "① 何もない場所で 右クリック", 30, RED_S)
    menu_box = context_menu(d, 300, 220, ["表示", "並べ替え", "貼り付け", "新規作成"], 2)
    circle(d, (menu_box[0] - 15, menu_box[1] - 10, menu_box[2] + 15, menu_box[3] + 10))
    txt(d, 300, 480, "② 「貼り付け」をクリック", 30, RED_S)
    # 右: 完成形
    txt(d, 880, 100, "しばらく待つと、デスクトップに", 32)
    folder_icon(d, 950, 180, 200, "1_積立金入力アシスタント", 26)
    txt(d, 880, 420, "このフォルダが現れれば成功です。", 32)
    txt(d, 880, 480, "（数分かかることがあります）", 26, GRAY_S)
    return _save(img, "s04.png")


def s05_open_folder():
    img = Image.new("RGB", (1500, 560), "white")
    d = ImageDraw.Draw(img)
    txt(d, 100, 60, "デスクトップにできたフォルダを ダブルクリック（すばやく2回）で開くと…", 30)
    window(d, (100, 130, 1400, 520), "1_積立金入力アシスタント")
    excel_file_icon(d, 200, 230, 110, "積立金入力アシスタント.xlsm")
    excel_file_icon(d, 560, 230, 110, "口座マスターひな形.xlsx")
    folder_icon(d, 900, 250, 130, "練習用データ")
    circle(d, (160, 200, 420, 460))
    txt(d, 620, 430, "← この「積立金入力アシスタント.xlsm」をこれから使います", 28, RED_S)
    return _save(img, "s05.png")


def s06_properties_menu():
    img = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(img)
    excel_file_icon(d, 150, 120, 130, "積立金入力アシスタント.xlsm")
    txt(d, 100, 320, "① このファイルを", 28, RED_S)
    txt(d, 100, 360, "　 右クリック", 28, RED_S)
    items = ["開く", "コピー", "切り取り", "名前の変更", "プロパティ"]
    menu_box = context_menu(d, 450, 120, items, 4)
    circle(d, (menu_box[0] - 15, menu_box[1] - 10, menu_box[2] + 15, menu_box[3] + 10))
    txt(d, 450, 460, "② いちばん下の「プロパティ」をクリック", 30, RED_S)
    txt(d, 920, 150, "※メニューの中身は", 26, GRAY_S)
    txt(d, 920, 190, "　パソコンによって多少違いますが、", 26, GRAY_S)
    txt(d, 920, 230, "　「プロパティ」は必ずいちばん下に", 26, GRAY_S)
    txt(d, 920, 270, "　あります。", 26, GRAY_S)
    txt(d, 920, 350, "※Windows 11 で見つからないときは", 26, GRAY_S)
    txt(d, 920, 390, "　「その他のオプションを表示」を先に", 26, GRAY_S)
    txt(d, 920, 430, "　クリックしてください。", 26, GRAY_S)
    return _save(img, "s06.png")


def s07_properties_check():
    img = Image.new("RGB", (1500, 660), "white")
    d = ImageDraw.Draw(img)
    window(d, (430, 30, 1230, 620), "積立金入力アシスタント.xlsm のプロパティ")
    # タブ
    for i, t in enumerate(["全般", "セキュリティ", "詳細"]):
        fill = "white" if i == 0 else "#DDDDDD"
        d.rectangle((460 + i * 160, 100, 380 + (i + 1) * 160 - 8, 150), fill=fill, outline="#999999", width=2)
        ctext(d, 380 + i * 160 + 76, 112, t, 26)
    d.line((450, 150, 1210, 150), fill="#999999", width=2)
    txt(d, 480, 180, "種類:  Microsoft Excel マクロ有効ワークシート", 24, GRAY_S)
    txt(d, 480, 230, "場所:  C:\\Users\\...\\Desktop", 24, GRAY_S)
    d.line((470, 300, 1190, 300), fill="#CCCCCC", width=2)
    txt(d, 480, 330, "セキュリティ: このファイルは他のコンピューターから", 26)
    txt(d, 480, 370, "取得したものです。…", 26)
    # チェックボックス
    d.rectangle((510, 430, 558, 478), outline=RED_S, width=7)
    d.line((520, 452, 534, 468), fill=RED_S, width=9)
    d.line((534, 468, 552, 438), fill=RED_S, width=9)
    txt(d, 580, 434, "許可する", 32, RED_S)
    circle(d, (490, 415, 780, 495))
    txt(d, 800, 440, "← ①ここにチェック", 30, RED_S)
    # OK
    d.rounded_rectangle((1010, 540, 1190, 600), radius=8, fill=YELLOW_S, outline=RED_S, width=5)
    ctext(d, 1100, 554, "OK", 30, RED_S, 1)
    circle(d, (990, 525, 1210, 615))
    txt(d, 460, 630, "②「OK」を押して閉じる", 28, RED_S)
    txt(d, 30, 60, "※「セキュリティ:」の行が", 26, GRAY_S)
    txt(d, 30, 100, "　見当たらない場合は、", 26, GRAY_S)
    txt(d, 30, 140, "　何もせず右上の×で閉じて", 26, GRAY_S)
    txt(d, 30, 180, "　次の手順へ進んで大丈夫です。", 26, GRAY_S)
    return _save(img, "s07.png")


def s08_open_excel():
    img = Image.new("RGB", (1500, 560), "white")
    d = ImageDraw.Draw(img)
    # タスクバー＋スタート
    d.rectangle((60, 420, 1440, 500), fill="#E8E8E8", outline="#AAAAAA", width=3)
    d.rectangle((100, 436, 165, 484), fill=BLUE_S)
    ctext(d, 132, 444, "田", 34, "white")
    circle(d, (80, 420, 185, 500))
    txt(d, 90, 510, "① 画面左下の「スタート」をクリック", 28, RED_S)
    # スタートメニュー
    window(d, (250, 60, 800, 400), "スタートメニュー", "#444444")
    for i, name in enumerate(["Excel", "Word", "メール"]):
        yy = 140 + i * 80
        if i == 0:
            d.rectangle((280, yy - 8, 770, yy + 60), fill=YELLOW_S)
        d.rectangle((300, yy, 356, yy + 52), fill="#1D6F42" if i == 0 else "#888888")
        ctext(d, 328, yy + 8, "X" if i == 0 else " ", 32, "white")
        txt(d, 390, yy + 6, name, 30)
    circle(d, (270, 120, 790, 210))
    txt(d, 850, 160, "②「Excel」をクリックして開く", 30, RED_S)
    txt(d, 850, 240, "※一覧に無ければ、スタートを押して", 26, GRAY_S)
    txt(d, 850, 280, "　そのまま「excel」と入力すると", 26, GRAY_S)
    txt(d, 850, 320, "　出てきます。", 26, GRAY_S)
    return _save(img, "s08.png")


def s09_options():
    img = Image.new("RGB", (1500, 640), "white")
    d = ImageDraw.Draw(img)
    window(d, (150, 40, 1350, 600), "Excel（最初の画面）", "#1D6F42")
    # 左メニュー
    d.rectangle((153, 92, 420, 597), fill="#F0F6F0")
    for i, name in enumerate(["ホーム", "新規", "開く", "", "アカウント", "その他...", "オプション"]):
        if not name:
            continue
        yy = 130 + i * 62
        if name == "オプション":
            d.rectangle((165, yy - 6, 410, yy + 44), fill=YELLOW_S)
        txt(d, 200, yy, name, 28)
    circle(d, (155, 495, 425, 570))
    txt(d, 470, 500, "← 左メニューの いちばん下 にある", 30, RED_S)
    txt(d, 470, 545, "「オプション」をクリック", 30, RED_S)
    txt(d, 470, 150, "※「オプション」が見えないときは", 26, GRAY_S)
    txt(d, 470, 190, "　「その他...」の中にあります。", 26, GRAY_S)
    return _save(img, "s09.png")


def s10_trust_center():
    img = Image.new("RGB", (1500, 640), "white")
    d = ImageDraw.Draw(img)
    window(d, (100, 40, 1400, 600), "Excel のオプション")
    # 左リスト
    d.rectangle((103, 92, 420, 597), fill="#F7F7F7")
    for i, name in enumerate(["全般", "数式", "保存", "言語", "詳細設定", "トラスト センター"]):
        yy = 120 + i * 70
        if i == 5:
            d.rectangle((115, yy - 6, 410, yy + 48), fill=YELLOW_S)
        txt(d, 140, yy, name, 28)
    circle(d, (105, 460, 425, 545))
    txt(d, 130, 555, "① 左の いちばん下", 28, RED_S)
    # 右ペインのボタン
    txt(d, 500, 160, "Microsoft Excel トラスト センター", 28)
    d.rounded_rectangle((870, 300, 1330, 370), radius=8, fill=YELLOW_S, outline=RED_S, width=5)
    ctext(d, 1100, 318, "トラスト センターの設定...", 28, RED_S)
    circle(d, (850, 285, 1350, 385))
    txt(d, 620, 420, "② 右側のこのボタンをクリック", 30, RED_S)
    return _save(img, "s10.png")


def s11_trusted_location():
    img = Image.new("RGB", (1500, 640), "white")
    d = ImageDraw.Draw(img)
    window(d, (100, 40, 1400, 600), "トラスト センター")
    d.rectangle((103, 92, 460, 597), fill="#F7F7F7")
    for i, name in enumerate(["信頼できる発行元", "信頼できる場所", "信頼済みドキュメント", "マクロの設定"]):
        yy = 130 + i * 76
        if i == 1:
            d.rectangle((115, yy - 6, 450, yy + 50), fill=YELLOW_S)
        txt(d, 140, yy, name, 28)
    circle(d, (105, 196, 465, 266))
    txt(d, 500, 216, "← ① 左の「信頼できる場所」をクリック", 28, RED_S)
    # 下部のボタン
    d.rounded_rectangle((520, 480, 940, 550), radius=8, fill=YELLOW_S, outline=RED_S, width=5)
    ctext(d, 730, 498, "新しい場所の追加(A)...", 28, RED_S)
    circle(d, (500, 465, 960, 565))
    txt(d, 990, 500, "② このボタンをクリック", 30, RED_S)
    txt(d, 520, 160, "信頼できる場所の一覧（表）が出ます", 26, GRAY_S)
    return _save(img, "s11.png")


def s12_browse():
    img = Image.new("RGB", (1500, 640), "white")
    d = ImageDraw.Draw(img)
    window(d, (250, 40, 1250, 430), "Microsoft Office の信頼できる場所")
    txt(d, 290, 130, "パス(P):", 26, GRAY_S)
    d.rectangle((290, 170, 1000, 225), fill="white", outline="#999999", width=3)
    txt(d, 305, 182, "C:\\Users\\...\\Desktop\\1_積立金入力アシスタント", 24)
    d.rounded_rectangle((1030, 165, 1210, 230), radius=8, fill=YELLOW_S, outline=RED_S, width=5)
    ctext(d, 1120, 182, "参照(B)...", 26, RED_S)
    circle(d, (1010, 150, 1230, 245))
    txt(d, 290, 270, "①「参照」を押して、デスクトップ →「1_積立金入力アシスタント」を選ぶ", 28, RED_S)
    d.rounded_rectangle((950, 330, 1090, 395), radius=8, fill=YELLOW_S, outline=RED_S, width=5)
    ctext(d, 1020, 346, "OK", 28, RED_S)
    circle(d, (930, 315, 1110, 410))
    txt(d, 350, 345, "② 選べたら「OK」", 28, RED_S)
    # OK3連打
    txt(d, 300, 480, "③ そのあと、開いている画面すべてで「OK」を押して閉じます（合計3回くらい）", 30)
    for i in range(3):
        d.rounded_rectangle((350 + i * 260, 540, 540 + i * 260, 600), radius=8, fill=LGRAY_S, outline=NAVY_S, width=4)
        ctext(d, 445 + i * 260, 554, f"OK （{i + 1}回目）", 26)
    return _save(img, "s12.png")


def s13_open_check():
    img = Image.new("RGB", (1500, 620), "white")
    d = ImageDraw.Draw(img)
    txt(d, 100, 50, "① フォルダに戻り、「積立金入力アシスタント.xlsm」を ダブルクリック で開く", 30)
    window(d, (100, 120, 1400, 560), "積立金入力アシスタント.xlsm - Excel", "#1D6F42")
    # 警告バー無しゾーン
    d.rectangle((110, 180, 1390, 240), fill="#EAF6EA", outline="#2E7D32", width=3)
    ctext(d, 750, 192, "← ここに黄色や赤の帯（警告バー）が出ていなければ、設定成功です！", 27, "#2E7D32")
    # メニューのボタン
    for i in range(3):
        for j in range(2):
            d.rounded_rectangle((260 + i * 340, 300 + j * 90, 540 + i * 340, 360 + j * 90),
                                radius=8, fill=LGRAY_S, outline=BLUE_S, width=3)
    ctext(d, 750, 500, "メニューにボタンが並んでいます（押すのはZoomレクチャーの日でOK）", 26, GRAY_S)
    return _save(img, "s13.png")


def s14_enable_content():
    img = Image.new("RGB", (1500, 480), "white")
    d = ImageDraw.Draw(img)
    txt(d, 100, 40, "もし黄色い帯が出たら → ボタンを1回押すだけで消えます", 32, RED_S)
    window(d, (100, 120, 1400, 420), "積立金入力アシスタント.xlsm - Excel", "#1D6F42")
    d.rectangle((110, 180, 1390, 250), fill="#FFF4CE", outline="#C8A200", width=3)
    txt(d, 140, 198, "セキュリティの警告  マクロが無効にされました。", 28)
    d.rounded_rectangle((1020, 190, 1330, 242), radius=6, fill="white", outline=RED_S, width=5)
    ctext(d, 1175, 200, "コンテンツの有効化", 27, RED_S)
    circle(d, (1000, 175, 1350, 258))
    txt(d, 640, 310, "←「コンテンツの有効化」をクリック", 30, RED_S)
    txt(d, 140, 360, "※赤い帯（ブロックされました）が出た場合は、手順6〜7（プロパティ→許可する）に戻ってください。", 24, GRAY_S)
    return _save(img, "s14.png")


def s15_usb2():
    img = Image.new("RGB", (1500, 480), "white")
    d = ImageDraw.Draw(img)
    ctext(d, 750, 30, "もう1台（受取用パソコン）は、これだけで終わりです", 34, NAVY_S, 1)
    # USB2 → コピー → デスクトップ
    d.rounded_rectangle((130, 170, 320, 232), radius=12, fill="white", outline=NAVY_S, width=5)
    d.ellipse((160, 150, 230, 220), fill=AMBER_S, outline=NAVY_S, width=4)
    ctext(d, 195, 164, "②", 42, "white", 1)
    ctext(d, 225, 250, "USB②を差す", 28)
    arrow(d, 380, 200, 520, 200, AMBER_S)
    folder_icon(d, 560, 140, 190, "1_積立金データ受け渡し", 24)
    arrow(d, 800, 200, 940, 200, AMBER_S)
    ctext(d, 862, 130, "コピー", 26, RED_S)
    d.rectangle((980, 100, 1400, 320), fill="#EAF2E6", outline="#2E7D32", width=4)
    ctext(d, 1190, 120, "デスクトップ", 26, "#2E7D32")
    folder_icon(d, 1090, 170, 180)
    ctext(d, 750, 380, "手順3〜4と同じ「右クリック→コピー」「右クリック→貼り付け」です（マクロ設定は不要）", 27, GRAY_S)
    return _save(img, "s15.png")


# ==================================================================
# docx 組み立て
# ==================================================================
def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=4, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def step(doc, num, title, img_path, ok_text=None, width_cm=16.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"手順{num}　{title}")
    set_jp(r, size=13.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '203864')
    pPr.append(shd)

    ip = doc.add_paragraph()
    ip.paragraph_format.space_after = Pt(2)
    ip.paragraph_format.keep_with_next = bool(ok_text)
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(img_path, width=Cm(width_cm))

    if ok_text:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run("✔ こうなればOK： " + ok_text)
        set_jp(r2, size=10.5, bold=True, color=GREEN)


def build():
    global _workdir
    _workdir = tempfile.mkdtemp(prefix="fullsetup_")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.3))

    para(doc, "かんたん導入ガイド（完全図解版）", size=19, bold=True, color=NAVY, center=True, space_after=2)
    para(doc, "1つの手順に、絵が1枚ずつ付いています。上から順に、絵のとおりに操作してください。",
         size=11, center=True, space_after=2)
    para(doc, "全部で15手順・約15分です。途中でうまくいかなくなったら、そのままにして大丈夫。"
         "Zoomレクチャーの日に一緒にやります。", size=10, center=True, color=(0xC0, 0, 0), space_after=8)

    step(doc, 1, "USB①「記録側PC用」を、保管用パソコン（積立金マスターがある方）に差す",
         s01_insert(), "「ピロン」と音が鳴る、または画面の右下に小さな表示が出ます")
    step(doc, 2, "「エクスプローラー」を開いて、USBドライブをダブルクリック",
         s02_explorer(), "USBの中身（フォルダなど）が表示されます")
    step(doc, 3, "「1_積立金入力アシスタント」フォルダを右クリック →「コピー」",
         s03_copy())
    step(doc, 4, "デスクトップの何もない場所を右クリック →「貼り付け」",
         s04_paste(), "デスクトップに同じ名前のフォルダが現れます")
    step(doc, 5, "デスクトップのフォルダをダブルクリックで開く",
         s05_open_folder(), "「積立金入力アシスタント.xlsm」という緑のファイルが見えます")
    step(doc, 6, "「積立金入力アシスタント.xlsm」を右クリック → いちばん下の「プロパティ」",
         s06_properties_menu())
    step(doc, 7, "いちばん下の「許可する」にチェックを入れて「OK」",
         s07_properties_check(), "何も表示が変わらなくてOK（設定は済んでいます）")
    step(doc, 8, "Excelを開く（スタート → Excel）",
         s08_open_excel(), "Excelの緑の最初の画面が開きます")
    step(doc, 9, "左メニューのいちばん下「オプション」をクリック",
         s09_options(), "「Excel のオプション」という画面が開きます")
    step(doc, 10, "左の「トラスト センター」→ 右の「トラスト センターの設定...」",
         s10_trust_center())
    step(doc, 11, "左の「信頼できる場所」→ 下の「新しい場所の追加(A)...」",
         s11_trusted_location())
    step(doc, 12, "「参照」でデスクトップの「1_積立金入力アシスタント」を選び、OKで全部閉じる",
         s12_browse(), "画面がすべて閉じて、Excelの最初の画面に戻ります")
    step(doc, 13, "フォルダに戻って「積立金入力アシスタント.xlsm」をダブルクリックで開く",
         s13_open_check(), "警告の帯が出ずに開けば、このパソコンの準備は完了です！")
    step(doc, 14, "（もし黄色い帯が出たら）「コンテンツの有効化」を1回押す",
         s14_enable_content(), "帯が消えて、次からは何も出なくなります")
    step(doc, 15, "もう1台（受取用パソコン）にUSB②を差して、フォルダをコピーするだけ",
         s15_usb2(), "デスクトップに「1_積立金データ受け渡し」フォルダがあれば、すべて完了です")

    # 最終ページ: 完了チェック
    doc.add_page_break()
    para(doc, "さいごに ― 完了チェック", size=14, bold=True, color=NAVY, space_after=6)
    for t in [
        "□　保管用パソコンのデスクトップに「1_積立金入力アシスタント」フォルダがある",
        "□　「積立金入力アシスタント.xlsm」が、警告の帯なしで開ける",
        "□　受取用パソコンのデスクトップに「1_積立金データ受け渡し」フォルダがある",
        "□　USB2本は、なくさない場所に保管した（今後の受け渡しにも使います）",
    ]:
        para(doc, t, size=11.5, space_after=6)
    para(doc, "この4つに印が付けば、あとはZoomレクチャーの日を待つだけです。おつかれさまでした！",
         size=11.5, bold=True, color=GREEN, space_after=10)
    para(doc, "うまくいかなかった手順があっても、そのままで大丈夫です。当日、画面を見ながら一緒に行います。",
         size=10.5, color=(0x88, 0x88, 0x88))

    docx_path = os.path.join(_workdir, "full_setup.docx")
    doc.save(docx_path)
    env = {**os.environ, "HOME": _workdir}
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", _workdir, docx_path],
                   check=True, env=env)
    os.replace(os.path.join(_workdir, "full_setup.pdf"), os.path.abspath(OUT_PDF))
    print("pdf saved:", os.path.abspath(OUT_PDF))


if __name__ == "__main__":
    build()
