# -*- coding: utf-8 -*-
"""モニター校覚書（ひな形）を生成する。

無償モニター提供の条件を書面化するためのたたき台。
学校に提示する前に、必要に応じて専門家（行政書士等）の確認を推奨。

生成: docs/02_営業・商談資料/モニター校覚書_ひな形.pdf
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GRAY = (0x88, 0x88, 0x88)


def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=5, center=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def article(doc, title, items):
    para(doc, title, size=11.5, bold=True, color=NAVY, space_after=3)
    for t in items:
        para(doc, t, size=10.5, space_after=3, indent=0.4)


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(2.0))

    para(doc, "「積立金会計 入力アシスタント」モニター利用に関する覚書（案）",
         size=15, bold=True, color=NAVY, center=True, space_after=10)

    para(doc, "＿＿＿＿＿＿＿＿＿＿（以下「学校」という）と ＿＿＿＿＿＿＿＿＿＿（以下「提供者」という）は、"
              "積立金会計 入力アシスタント（以下「本製品」という）のモニター利用について、以下のとおり合意する。",
         size=10.5, space_after=10)

    article(doc, "第1条（目的）", [
        "提供者は、本製品の実務における有用性の検証を目的として、学校に対し本製品を無償で提供し、"
        "学校はこれを利用して検証に協力する。",
    ])
    article(doc, "第2条（提供物と期間）", [
        "1. 提供物：本製品一式（Excelブック・練習用データ・各種手順書）、導入時の説明（訪問またはオンライン）、"
        "および期間中のサポート（メール・オンライン）。",
        "2. モニター期間：＿＿＿＿年＿＿月＿＿日から1年間とする。",
        "3. 期間中の利用料は無償とする。学校に金銭的負担は発生しない。",
    ])
    article(doc, "第3条（学校の協力事項）", [
        "1. 導入前後の関連業務の作業時間について、提供者が用意する簡易な様式により月1回程度記録し、提供者に共有する。",
        "2. 本製品の改善要望・不具合について、月1回程度メール等でフィードバックする。",
        "3. 提供者が本製品の導入事例として紹介することを許諾する。校名の掲載可否は次のいずれかを選択する：",
        "　　□ 校名を掲載してよい　　□ 匿名（例：都内公立高校A校）での掲載に限る",
    ])
    article(doc, "第4条（データの取り扱い）", [
        "1. 生徒に関する情報は学校の管理するパソコン内でのみ取り扱われ、提供者がこれを受領・持ち出すことはない。",
        "2. 提供者が事例紹介・検証結果の公表を行う場合も、生徒個人を特定できる情報は一切含めない。",
        "3. 第3条の作業時間記録に、生徒の個人情報は含まれない。",
    ])
    article(doc, "第5条（期間終了後の取り扱い）", [
        "1. モニター期間の終了1ヶ月前までに、提供者は有償での継続条件を提示し、学校は継続または終了を選択する。"
        "自動的に有償契約へ移行することはない。",
        "2. 終了を選択した場合、学校は本製品の利用を停止する。学校が本製品を用いて作成・更新したデータおよび"
        "既存の積立金マスターファイルは、終了後もそのまま学校が利用できる。",
    ])
    article(doc, "第6条（免責等）", [
        "1. 本製品は書き込み前に自動バックアップを作成する設計であるが、学校は重要データについて自らも適宜バックアップを行う。",
        "2. 提供者は、故意または重過失による場合を除き、本製品の利用により生じた損害について責任を負わない。",
        "3. 本覚書に定めのない事項は、両者が誠意をもって協議のうえ定める。",
    ])

    para(doc, "", space_after=8)
    para(doc, "本覚書の成立を証するため、本書2通を作成し、双方記名押印のうえ各1通を保有する。", size=10.5, space_after=14)
    para(doc, "＿＿＿＿年＿＿月＿＿日", size=10.5, space_after=14)
    para(doc, "学校　　：住所・名称・代表者　＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿　印", size=10.5, space_after=10)
    para(doc, "提供者　：住所・名称・代表者　＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿　印", size=10.5, space_after=14)

    para(doc, "※本書はひな形（案）です。締結前に文言・法的事項をご確認ください。", size=9, color=GRAY)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, env={**os.environ, "HOME": tempfile.gettempdir()})
    os.replace(os.path.splitext(docx_path)[0] + ".pdf", pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "monitor_agreement.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "..", "02_営業・商談資料", "モニター校覚書_ひな形.pdf"))
