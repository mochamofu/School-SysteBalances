# -*- coding: utf-8 -*-
"""動作確認チェックシート（実機テスト手順・期待値つき）を生成する。

出荷前の実機デバッグと、学校側の受け入れ確認の両方に使う。
ボタン①〜⑮を順に実行し、期待値どおりかを○×で記録する形式。

生成: docs/01_学校向けマニュアル/動作確認チェックシート.pdf
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
RED = (0xC0, 0x00, 0x00)


def set_jp(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10, bold=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=12, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def check_table(doc, rows, widths=(1.2, 6.6, 7.4, 1.4, 1.6)):
    headers = ["No", "操作", "期待どおりの結果", "結果", "メモ"]
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_jp(r, size=9, bold=True, color=(0xFF, 0xFF, 0xFF))
        shade(cell, "203864")
    for i, (op, expect) in enumerate(rows):
        vals = [str(i + 1), op, expect, "", ""]
        for j, v in enumerate(vals):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(v)
            set_jp(r, size=9)
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.3))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("動作確認チェックシート")
    set_jp(r, size=16, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("実施日（　　　／　　　／　　　）　実施者（　　　　　　　　）　Excelバージョン（　　　　　　　　）")
    set_jp(r, size=9.5)

    para(doc, "上から順に実行し、結果欄に ○（期待どおり）／×（違った）を記入してください。"
              "×の場合はメモ欄に何が起きたかをひとこと書き、エラーメッセージは画面を撮影して保存してください。"
              "使用するのはすべて架空データです。本物のマスターには一切触れません。", size=9.5, space_after=6)

    # ============ A. 導入確認 ============
    h1(doc, "A. 導入確認（最初に1回）")
    check_table(doc, [
        ("積立金入力アシスタント.xlsm を開く",
         "警告バー（赤・黄）が出ずに開く。※出た場合は導入手順書の「マクロを許可する」を実施"),
        ("メニューシートを見る",
         "ボタン①〜⑮が並んでいる（並んでいなければ Alt+F8 →「初期設定」を1回実行）"),
        ("設定シートに 練習用_令和X年度生積立金.xlsx のフルパス（C3）、年度（C5）、"
         "練習用_口座マスター.xlsx のパス（C7）を入力", "入力できる（この3つで準備完了）"),
    ])

    # ============ B. 練習用データ（80名）での一巡 ============
    h1(doc, "B. 基本動作（練習用データ・架空の80名）")
    check_table(doc, [
        ("名簿貼付シートに 練習用_掲示用名簿 のシート全体を貼り付け → ①名簿を解析して照合する",
         "「4クラス80名」を検出し、名簿一覧で全員が「一致」（赤い行なし）"),
        ("②クラス替えをマスターに反映する",
         "80名分の組・番号が更新される。バックアップ作成のメッセージが出る"),
        ("支出入力シートに 支出No=空き番号／件名=校外学習バス代／日付=当日／金額=3500／対象=全員 → ④",
         "「対象80名・合計280,000円」。支出承認書シートが自動で埋まる"),
        ("振替結果取込シートB12に 練習用_振替結果 の4列×80行を貼り付け → ⑪振替結果を照合",
         "読取80件／振替済78件／未納2名／不明口座0件。未納の精算番号は 7 と 44"),
        ("収入入力シートを見る（⑪の直後）",
         "未納者表に精算番号7・44の2名が自動で入っている"),
        ("収入入力シートに 枠No=空き枠／件名=口座振替／金額=76000 → ⑤収入をマスターへ一括入力",
         "「入金あり78名・未納2名」。練習用マスターのH列で2名に未納の印"),
        ("⑥収入枠の一覧を表示", "いま入力した枠が「78名」で表示される"),
        ("⑦決算用の集計を実行", "決算集計シートに支出1件・収入1件が人数・総額つきで並ぶ"),
        ("⑧マスターの整合性をチェック", "チェック結果シートに結果が出る（重大な不整合なし）"),
        ("⑨精算書を一括PDF保存", "「精算書PDF」フォルダに 組-番号_氏名.pdf が80枚できる"),
        ("年間予定表の1行目に予定を書き → ⑫予定を入力フォームへ転送（行No=1）",
         "支出入力（または収入入力）シートに件名・金額が自動で入る"),
        ("⑬支出項目を読み込む", "支出項目一覧に④で入れた項目が 人数80・総額280,000 で出る"),
        ("⑭業者別に集計する", "右側に業者別の年間合計が出る"),
        ("⑮来年度の予定へ引き継ぐ（H列に○を付けてから）", "○の項目だけが年間予定表に追加される"),
    ])

    # ============ C. フルスケール（320名） ============
    h1(doc, "C. 実物と同じ規模（検証用データ・架空の320名）")
    para(doc, "設定シートC3を 検証用_令和X年度生積立金.xlsx、C7を 検証用_口座マスター.xlsx に切り替えてから実施。",
         size=9.5, space_after=3)
    check_table(doc, [
        ("名簿貼付シートに 検証用_新年度名簿 を貼り付け → ①",
         "8クラス320名を検出・全員一致"),
        ("振替結果取込シートに 検証用_振替結果（321行）を貼り付け → ⑪",
         "読取321／振替済315／未納5（精算番号7・44・159・241・312）／不明口座1件"),
    ])

    # ============ D. わざと間違える（安全装置の確認） ============
    h1(doc, "D. 安全装置（わざと間違えて、止まることを確認）")
    check_table(doc, [
        ("支出入力の例外表に 精算番号400 を入れて ④",
         "「範囲外です」のエラーで止まる（マスターには何も書かれない）"),
        ("設定シートC3を空欄にして ④",
         "「設定を入力してください」の案内で止まる"),
        ("すでに金額が入っている支出Noを指定して ④",
         "「すでに○名分の金額が入っています。上書きしてよいですか？」の確認が出る（いいえで中止できる）"),
        ("普通のExcelファイル（マスター以外）をC3に指定して ④",
         "「データシートの形が想定と違います」で拒否される"),
    ])

    # ============ 不具合の報告 ============
    h1(doc, "×が付いたときの報告方法")
    for t in [
        "次の4点をメールでお送りください。最短で修正版をお返しします。",
        "　1. どのボタンか（例：⑪振替結果を照合）",
        "　2. エラーメッセージの画面写真（スマートフォン撮影で可）",
        "　3. 直前に行った操作（例：練習用_振替結果を貼り付けた直後）",
        "　4. 使っていたデータ（練習用80名／検証用320名）",
    ]:
        para(doc, t, size=9.5, space_after=2)

    para(doc, "A〜Dがすべて○になれば、実運用を開始できる状態です。", size=10.5, bold=True, color=RED, space_after=2)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, env={**os.environ, "HOME": tempfile.gettempdir()})
    os.replace(os.path.splitext(docx_path)[0] + ".pdf", pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "acceptance_checklist.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "..", "01_学校向けマニュアル", "動作確認チェックシート.pdf"))
