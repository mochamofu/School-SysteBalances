# -*- coding: utf-8 -*-
"""スタートガイド（図解版）とUSBラベル印刷シートを生成する。

郵送を受け取った学校の担当者が、封を開けて最初に見る紙。
「同梱物の確認 → どのUSBをどのPCへ → 3ステップ」を図で示す。
文字を読み込まなくても図だけで流れが分かることを最優先にする。

生成:
  docs/04_オンライン導入/スタートガイド_図解版.pdf   （印刷して同封する紙）
  docs/04_オンライン導入/USBラベル印刷シート.pdf     （手元用・ラベル切り貼り）
"""
import os
import subprocess
import tempfile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(HERE, "..", "04_オンライン導入")

FONT_PATH = "/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf"
FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
NAVY_S = "#203864"
BLUE_S = "#2F5FA8"
AMBER_S = "#C58F00"
YELLOW_S = "#FFF3B0"
RED_S = "#C00000"
GRAY_S = "#666666"
LGRAY_S = "#E8EDF5"

_workdir = None  # build() が設定


def F(size):
    return ImageFont.truetype(FONT_PATH, size)


def _save(img, name):
    path = os.path.join(_workdir, name)
    img.save(path)
    return path


def rounded(d, box, fill, outline, width=3, radius=18):
    d.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def center_text(d, cx, y, text, size, fill=NAVY_S, bold_stroke=0):
    f = F(size)
    w = d.textlength(text, font=f)
    d.text((cx - w / 2, y), text, font=f, fill=fill,
           stroke_width=bold_stroke, stroke_fill=fill)


def usb_icon(d, x, y, label, tag_color):
    """USBメモリの絵（本体＋端子＋大きな番号タグ）"""
    rounded(d, (x, y + 26, x + 150, y + 96), "white", NAVY_S, 4, 14)      # 本体
    d.rectangle((x + 150, y + 44, x + 186, y + 78), outline=NAVY_S, width=4, fill=LGRAY_S)  # 端子
    d.ellipse((x + 8, y, x + 76, y + 68), fill=tag_color, outline=NAVY_S, width=4)  # タグ
    center_text(d, x + 42, y + 12, label, 40, "white", 1)


def pc_icon(d, cx, y, w=260, h=170, screen_text=""):
    """ノートPCの絵"""
    d.rectangle((cx - w // 2, y, cx + w // 2, y + h), outline=NAVY_S, width=5, fill="white")
    if screen_text:
        f = F(24)
        lines = screen_text.split("\n")
        ty = y + h // 2 - len(lines) * 15
        for ln in lines:
            tw = d.textlength(ln, font=f)
            d.text((cx - tw / 2, ty), ln, font=f, fill=NAVY_S)
            ty += 32
    d.polygon([(cx - w // 2 - 30, y + h + 40), (cx + w // 2 + 30, y + h + 40),
               (cx + w // 2, y + h), (cx - w // 2, y + h)], outline=NAVY_S, fill=LGRAY_S, width=4)


def arrow(d, x1, y1, x2, y2, color=AMBER_S, width=8):
    d.line((x1, y1, x2, y2), fill=color, width=width)
    import math
    ang = math.atan2(y2 - y1, x2 - x1)
    L = 24
    for da in (2.6, -2.6):
        d.line((x2, y2, x2 - L * math.cos(ang + da), y2 - L * math.sin(ang + da)),
               fill=color, width=width)


# ------------------------------------------------------------------
# 図1: 同梱物の確認
# ------------------------------------------------------------------
def fig_contents():
    img = Image.new("RGB", (1500, 460), "white")
    d = ImageDraw.Draw(img)
    center_text(d, 750, 16, "【同梱物】 そろっているか確認してください", 40, NAVY_S, 1)

    usb_icon(d, 180, 130, "①", BLUE_S)
    center_text(d, 280, 250, "USBメモリ①", 30)
    center_text(d, 280, 292, "「記録側PC用」", 28, BLUE_S)

    usb_icon(d, 620, 130, "②", AMBER_S)
    center_text(d, 720, 250, "USBメモリ②", 30)
    center_text(d, 720, 292, "「送信側PC用」", 28, AMBER_S)

    # 紙2枚
    for i, (t1, t2) in enumerate([("スタートガイド", "（この紙）"), ("かんたん導入ガイド", "（完全図解版）")]):
        x = 1020 + i * 230
        d.rectangle((x, 110, x + 180, 340), outline=NAVY_S, width=4, fill="white")
        d.rectangle((x + 14, 110 + 14, x + 194, 354), outline=GRAY_S, width=2)
        center_text(d, x + 90, 180, t1, 26)
        center_text(d, x + 90, 220, t2, 26)
    center_text(d, 1250, 380, "印刷物 2部", 30)
    return _save(img, "sg_contents.png")


# ------------------------------------------------------------------
# 図2: どのUSBをどのパソコンへ
# ------------------------------------------------------------------
def fig_which_pc():
    img = Image.new("RGB", (1500, 560), "white")
    d = ImageDraw.Draw(img)
    center_text(d, 750, 10, "【いちばん大事】 USBを差すパソコンが決まっています", 40, RED_S, 1)

    pc_icon(d, 400, 120, screen_text="積立金マスターが\nあるパソコン\n（保管用）")
    usb_icon(d, 180, 380, "①", BLUE_S)
    arrow(d, 400, 470, 400, 350)
    center_text(d, 400, 480, "USB①「記録側PC用」を差す", 32, BLUE_S, 1)

    pc_icon(d, 1100, 120, screen_text="銀行データや名簿を\n受け取るパソコン\n（受取用）")
    usb_icon(d, 880, 380, "②", AMBER_S)
    arrow(d, 1100, 470, 1100, 350)
    center_text(d, 1100, 480, "USB②「送信側PC用」を差す", 32, AMBER_S, 1)
    return _save(img, "sg_which.png")


# ------------------------------------------------------------------
# 図3: ステップ1 コピー
# ------------------------------------------------------------------
def fig_copy():
    img = Image.new("RGB", (1500, 420), "white")
    d = ImageDraw.Draw(img)

    # USBの中身ウィンドウ
    rounded(d, (60, 40, 640, 380), "white", NAVY_S, 4, 10)
    d.rectangle((60, 40, 640, 90), fill=NAVY_S)
    center_text(d, 350, 52, "USBドライブの中", 26, "white")
    # フォルダアイコン
    d.polygon([(120, 150), (200, 150), (215, 170), (360, 170), (360, 300), (120, 300)],
              fill="#FFD35C", outline=AMBER_S, width=4)
    center_text(d, 240, 320, "1_積立金入力アシスタント", 24)

    arrow(d, 680, 210, 850, 210)
    center_text(d, 765, 150, "丸ごとコピー", 30, RED_S, 1)

    # デスクトップ
    rounded(d, (880, 40, 1450, 380), "#EAF2E6", "#2E7D32", 4, 10)
    d.rectangle((880, 40, 1450, 90), fill="#2E7D32")
    center_text(d, 1165, 52, "パソコンのデスクトップ", 26, "white")
    d.polygon([(990, 150), (1070, 150), (1085, 170), (1230, 170), (1230, 300), (990, 300)],
              fill="#FFD35C", outline=AMBER_S, width=4)
    center_text(d, 1110, 320, "ここに置く（USBから直接開かない）", 24, "#2E7D32")
    return _save(img, "sg_copy.png")


# ------------------------------------------------------------------
# 図4: ステップ2a プロパティで許可
# ------------------------------------------------------------------
def fig_properties():
    img = Image.new("RGB", (1500, 560), "white")
    d = ImageDraw.Draw(img)
    # ダイアログ
    rounded(d, (330, 20, 1170, 520), "white", NAVY_S, 4, 10)
    d.rectangle((330, 20, 1170, 74), fill=NAVY_S)
    center_text(d, 750, 32, "積立金入力アシスタント.xlsm のプロパティ", 26, "white")
    f = F(26)
    d.text((380, 110), "全般タブの いちばん下 を見る", font=f, fill=GRAY_S)
    d.line((370, 160, 1130, 160), fill="#CCCCCC", width=2)
    d.text((380, 190), "セキュリティ: このファイルは他のコンピューター", font=F(28), fill=NAVY_S)
    d.text((380, 230), "から取得したものです。…", font=F(28), fill=NAVY_S)
    # チェックボックス
    d.rectangle((420, 300, 462, 342), outline=RED_S, width=6)
    d.line((428, 320, 440, 334), fill=RED_S, width=8)
    d.line((440, 334, 458, 306), fill=RED_S, width=8)
    d.text((480, 302), "許可する ← ここにチェック", font=F(32), fill=RED_S)
    # OKボタン
    rounded(d, (900, 420, 1090, 486), YELLOW_S, RED_S, 5, 8)
    center_text(d, 995, 436, "OK を押す", 30, RED_S, 1)
    d.text((340, 530), "※この表示が無ければ、何もしなくてOKです（そのまま次へ）", font=F(24), fill=GRAY_S)
    return _save(img, "sg_props.png")


# ------------------------------------------------------------------
# 図5: ステップ2b 信頼できる場所
# ------------------------------------------------------------------
def fig_trust():
    img = Image.new("RGB", (1700, 460), "white")
    d = ImageDraw.Draw(img)
    steps = ["ファイル", "オプション", "トラスト\nセンター", "トラスト センター\nの設定", "信頼できる場所", "新しい場所の追加"]
    x = 40
    for i, t in enumerate(steps):
        w = 190 if i < 3 else 250
        color = LGRAY_S if i < len(steps) - 1 else YELLOW_S
        outline = NAVY_S if i < len(steps) - 1 else RED_S
        rounded(d, (x, 60, x + w, 190), color, outline, 4, 12)
        lines = t.split("\n")
        ty = 100 if len(lines) == 1 else 80
        for ln in lines:
            center_text(d, x + w // 2, ty, ln, 26, NAVY_S if i < len(steps) - 1 else RED_S)
            ty += 36
        if i < len(steps) - 1:
            arrow(d, x + w + 4, 125, x + w + 32, 125, width=6)
        x += w + 36
    d.text((60, 250), "「参照」で、デスクトップにコピーした", font=F(30), fill=NAVY_S)
    d.text((60, 295), "「1_積立金入力アシスタント」フォルダを選んで OK", font=F(30), fill=NAVY_S)
    d.text((60, 370), "※Excelに元からある設定画面です。パソコン自体の設定は変えません。", font=F(24), fill=GRAY_S)
    return _save(img, "sg_trust.png")


# ------------------------------------------------------------------
# 図6: ステップ3 開いて確認
# ------------------------------------------------------------------
def fig_done():
    img = Image.new("RGB", (1500, 420), "white")
    d = ImageDraw.Draw(img)
    rounded(d, (200, 30, 1300, 360), "white", NAVY_S, 4, 10)
    d.rectangle((200, 30, 1300, 84), fill=NAVY_S)
    center_text(d, 750, 42, "積立金入力アシスタント.xlsm", 26, "white")
    # 警告バーが無いことを示す
    d.rectangle((220, 100, 1280, 150), fill="#EAF6EA", outline="#2E7D32", width=3)
    center_text(d, 750, 110, "← ここに黄色や赤の警告バーが出なければ準備完了！", 28, "#2E7D32")
    # ボタン群
    for i in range(3):
        for j in range(2):
            rounded(d, (300 + i * 340, 190 + j * 70, 580 + i * 340, 240 + j * 70),
                    LGRAY_S, BLUE_S, 3, 8)
    center_text(d, 750, 330, "メニューにボタンが並んでいれば、あとはZoomの日を待つだけです", 28, NAVY_S)
    return _save(img, "sg_done.png")


# ==================================================================
# docx ヘルパー
# ==================================================================
def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=4, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def step_head(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=13, bold=True, color=(0xFF, 0xFF, 0xFF))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '203864')
    pPr.append(shd)
    return p


def pic(doc, path, width_cm=17.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, env={**os.environ, "HOME": tempfile.gettempdir()})
    os.replace(os.path.splitext(docx_path)[0] + ".pdf", os.path.abspath(pdf_out))
    print("pdf saved:", os.path.abspath(pdf_out))


# ==================================================================
# スタートガイド本体
# ==================================================================
def build_guide():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.3))

    para(doc, "スタートガイド ― 封を開けたら、まずこの紙", size=19, bold=True, color=NAVY, center=True, space_after=2)
    para(doc, "むずかしい作業はありません。この紙のとおりに進めれば約15分で準備が終わります。",
         size=10.5, center=True, space_after=8)

    pic(doc, fig_contents(), 17.0)
    pic(doc, fig_which_pc(), 17.0)
    para(doc, "※どちらのUSBにも生徒の実データは入っていません（練習用の架空データのみ）。"
              "いまお使いの積立金ファイルには何も手を加えません。", size=9.5, color=(0xC0, 0, 0), space_after=2)

    # ---- 2ページ目: 3ステップ ----
    doc.add_page_break()
    step_head(doc, "ステップ1　フォルダをデスクトップにコピーする（両方のパソコンで）")
    para(doc, "USB①の「1_積立金入力アシスタント」、USB②の「1_積立金データ受け渡し」を、"
              "それぞれのパソコンのデスクトップへ丸ごとコピーします。", size=10)
    pic(doc, fig_copy(), 16.0)

    step_head(doc, "ステップ2　マクロを許可する（保管用パソコンだけ・2分）")
    para(doc, "(1) デスクトップにコピーした「積立金入力アシスタント.xlsm」を右クリック →「プロパティ」。", size=10)
    pic(doc, fig_properties(), 14.5)
    para(doc, "(2) Excelを開き、次の順にクリックして、コピーしたフォルダを登録します。", size=10)
    pic(doc, fig_trust(), 17.0)

    # ---- 3ページ目 ----
    doc.add_page_break()
    step_head(doc, "ステップ3　開いて確認する")
    pic(doc, fig_done(), 16.0)

    step_head(doc, "うまくいかなくても大丈夫です")
    for t in [
        "・設定がうまくいかなくても、そのままにしておいてください。Zoomレクチャー当日に画面を見ながら一緒に行います。",
        "・1クリックずつの絵つき手順は、同封の「かんたん導入ガイド（完全図解版）」にあります。",
        "・お急ぎの場合の連絡先は、お送りしたご案内メールに記載しています。",
    ]:
        para(doc, t, size=10.5, space_after=3)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("サポート連絡先（担当者記入欄）：")
    set_jp(r, size=10.5, bold=True)
    r2 = p.add_run("　　　　　　　　　　　　　　　　　　　　　　")
    set_jp(r2, size=10.5)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

    path = os.path.join(_workdir, "start_guide.docx")
    doc.save(path)
    return path


# ==================================================================
# USBラベル印刷シート（手元用）
# ==================================================================
def build_labels():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.3))

    para(doc, "USBラベル印刷シート（手元用・点線で切ってUSBと封筒に貼る）",
         size=13, bold=True, color=NAVY, space_after=8)

    def label_img(no, title, sub, color):
        img = Image.new("RGB", (1200, 360), "white")
        d = ImageDraw.Draw(img)
        d.rectangle((4, 4, 1196, 356), outline=GRAY_S, width=3)
        d.ellipse((40, 60, 280, 300), fill=color, outline=NAVY_S, width=6)
        center_text(d, 160, 110, no, 120, "white", 2)
        d.text((330, 90), title, font=F(72), fill=NAVY_S)
        d.text((330, 210), sub, font=F(44), fill=GRAY_S)
        return _save(img, f"label_{no}.png")

    l1 = label_img("①", "記録側PC用", "積立金マスターがある保管用パソコンへ", BLUE_S)
    l2 = label_img("②", "送信側PC用", "銀行データ・名簿を受け取るパソコンへ", AMBER_S)
    for path in (l1, l1, l2, l2):
        pic(doc, path, 12.5)

    # 封筒用
    img = Image.new("RGB", (1400, 300), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((4, 4, 1396, 296), outline=GRAY_S, width=3)
    center_text(d, 700, 60, "積立金入力アシスタント 導入キット 在中", 64, NAVY_S, 1)
    center_text(d, 700, 180, "USBメモリ2本・印刷物2部", 44, GRAY_S)
    env = _save(img, "label_env.png")
    pic(doc, env, 14.0)

    path = os.path.join(_workdir, "usb_labels.docx")
    doc.save(path)
    return path


def main():
    global _workdir
    _workdir = tempfile.mkdtemp(prefix="startguide_")
    os.makedirs(OUT_DIR, exist_ok=True)
    to_pdf(build_guide(), os.path.join(OUT_DIR, "スタートガイド_図解版.pdf"))
    to_pdf(build_labels(), os.path.join(OUT_DIR, "USBラベル印刷シート.pdf"))


if __name__ == "__main__":
    main()
