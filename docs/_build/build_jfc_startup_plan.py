# -*- coding: utf-8 -*-
"""創業計画書（日本政策金融公庫 国民生活事業の様式対応版）を生成する。

公庫「創業計画書」の記入欄（1〜8＋自由記述欄）と同じ欄立てで、
本事業の内容を記入済みにしたもの。提出時は公庫配布の正式様式へ
転記するか、正式様式に「別紙参照」と書いて本書を添付する使い方を想定。
末尾に任意添付用の「月別収支計画（1期目）」を付す。

生成: docs/06_法人化・創業融資/創業計画書_公庫様式対応版.pdf
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GRAY = (0x66, 0x66, 0x66)


def set_jp(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10, bold=False, color=None, space_after=4, center=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_text(cell, text, size=9.5, bold=False, color=None):
    cell.text = ""
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        set_jp(r, size=size, bold=bold, color=color)


def section_box(doc, num, title, body_lines, min_height_cm=None):
    """公庫様式風の番号付き枠（見出しセル＋本文セル）"""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells[0]
    cell_text(head, f"{num}　{title}", size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    shade(head, "203864")
    body = t.add_row().cells[0]
    cell_text(body, "\n".join(body_lines), size=9.5)
    for row in t.rows:
        row.cells[0].width = Cm(17.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def grid_table(doc, headers, rows, widths, size=9, bold_last=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        cell_text(c, h, size=size, bold=True)
        shade(c, "DCE8FA")
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell_text(t.rows[i + 1].cells[j], str(v), size=size,
                      bold=(bold_last and i == len(rows) - 1))
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin"):
        setattr(sec, m, Cm(1.4))
    for m in ("left_margin", "right_margin"):
        setattr(sec, m, Cm(1.8))

    para(doc, "創　業　計　画　書", size=18, bold=True, color=NAVY, center=True, space_after=2)
    para(doc, "（日本政策金融公庫 国民生活事業「創業計画書」の記入欄に対応した転記用ドラフト）",
         size=9, color=GRAY, center=True, space_after=2)
    para(doc, "お名前：＿＿＿＿＿＿＿＿＿＿＿＿　　作成日：＿＿＿＿年＿＿月＿＿日", size=10, center=True, space_after=8)

    # ============ 1 創業の動機 ============
    section_box(doc, 1, "創業の動機（創業されるのは、どのような目的、動機からですか。）", [
        "高等学校の積立金（学校徴収金）会計は、保護者からの預かり金でありながら、数百人分の手入力と",
        "目視照合に支えられており、担当職員の長時間労働と誤請求リスクが常態化している。",
        "私はこの実態に現場で接し、学校が使い慣れた既存のExcel管理ファイルを一切変更せずに、",
        "入力作業だけを自動化する製品を開発した。実際の学校ファイルでの検証と有償導入を経て",
        "有効性を確認できたため、全国の学校へ提供する事業として創業する。",
        "（※ご自身の原体験を1〜2行追記してください：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿）",
    ])

    # ============ 2 経営者の略歴等 ============
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    head = t.rows[0].cells[0]
    cell_text(head, "2　経営者の略歴等（略歴については、勤務先名だけではなく、担当業務や役職、身につけた技能等についても記載してください。）",
              size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    shade(head, "203864")
    body = t.add_row().cells[0]
    body.width = Cm(17.2)
    head.width = Cm(17.2)
    cell_text(body, "＜略歴＞", size=9.5, bold=True)
    inner = doc.tables[-1]
    # 略歴行は別テーブルで
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    grid_table(doc,
               ["年月", "内容"],
               [
                   ["＿＿年＿月", "【要記入】勤務先・担当業務・役職（本事業に繋がる経験を具体的に）"],
                   ["＿＿年＿月", "【要記入】"],
                   ["＿＿年＿月", "「積立金会計 入力アシスタント」を開発。実物形式ファイルでの検証・導入・レクチャーを実施"],
               ],
               widths=[3.0, 14.2])
    grid_table(doc,
               ["項目", "内容"],
               [
                   ["過去の事業経験", "□事業を経営していたことはない　□経営していたことがある（＿＿＿＿）"],
                   ["取得資格", "【要記入】＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿"],
                   ["知的財産権等", "特になし（製品名称の商標登録を検討中）"],
               ],
               widths=[3.6, 13.6])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============ 3 取扱商品・サービス ============
    section_box(doc, 3, "取扱商品・サービス", [
        "＜取扱商品・サービスの内容＞（売上シェア）",
        "① 私立高校向け 年間利用契約「積立金会計 入力アシスタント」（月額19,800円・税別）…シェア約60％",
        "② 公立高校向け 年間利用契約（年額29,800〜49,800円・税別）…シェア約30％",
        "③ 教育委員会向け 一括ライセンス（1校あたり年額30,000円・10校以上）…シェア約10％",
        "",
        "＜セールスポイント＞",
        "・学校が長年使う既存のExcel管理ファイルを一切変更せず、入力・照合・帳票作成だけを自動化",
        "・インストール不要／完全オフライン／生徒情報は校内から出ない ― 公立校でも導入手続きが軽い",
        "・書き込み前の自動バックアップ等の安全設計。実物形式ファイルで全機能検証済み（検証記録あり）",
        "・導入校の実測で作業時間の削減効果を定量提示（導入前後の作業時間記録）",
        "",
        "＜販売ターゲット・販売戦略＞",
        "・当初は都内の公立高校（先行導入校）で有償実績と実測データを取得",
        "・公立は学校事務職員の研究会・異動ネットワーク経由で横展開（広告に依存しない）",
        "・私立はWebサイトからの直販、将来は教育委員会への一括提案（10校単位）",
        "・導入はUSB郵送＋オンライン指導で完結し、1人体制で月10校の導入が可能",
        "",
        "＜競合・市場など企業を取り巻く状況＞",
        "・全国の高等学校は約4,800校（公立約3,500・私立約1,300）。市場規模の目安は年間約4億円",
        "・既存の校務支援システムは年24万円超／校でシステム移行が前提のため、積立金会計単体には過大",
        "・「既存ファイルを変えない・少額・インストール不要」の領域に直接競合は見当たらない",
    ])

    # ============ 4 取引先・取引関係等 ============
    t4_head = doc.add_table(rows=1, cols=1)
    t4_head.style = "Table Grid"
    c = t4_head.rows[0].cells[0]
    c.width = Cm(17.2)
    cell_text(c, "4　取引先・取引関係等", size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    shade(c, "203864")
    grid_table(doc,
               ["区分", "取引先名（所在地等）", "シェア", "掛取引の割合", "回収・支払の条件"],
               [
                   ["販売先", "高等学校（東京都内中心、順次全国）", "100%", "100%", "請求書払い。公立=年額一括（納品月の翌月末）、私立=月末〆翌月末または年払"],
                   ["仕入先", "なし（自社開発ソフトウェア）", "―", "―", "―"],
                   ["外注先", "税理士・行政書士（顧問・スポット）", "―", "100%", "月末〆翌月末"],
               ],
               widths=[1.8, 5.6, 1.6, 2.2, 6.0])
    grid_table(doc, ["人件費の支払"],
               [["役員報酬のみ（毎月末払い・賞与なし）。従業員なし"]],
               widths=[17.2])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============ 5 従業員 ============
    grid_table(doc,
               ["5　従業員", "常勤役員の人数", "従業員数", "うち家族従業員", "うちパート従業員"],
               [["", "1 人", "0 人", "0 人", "0 人"]],
               widths=[4.6, 3.4, 3.0, 3.2, 3.0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============ 6 お借入の状況 ============
    t6 = doc.add_table(rows=1, cols=1)
    t6.style = "Table Grid"
    c = t6.rows[0].cells[0]
    c.width = Cm(17.2)
    cell_text(c, "6　お借入の状況（法人の場合、代表者の方のお借入）", size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    shade(c, "203864")
    grid_table(doc,
               ["お借入先名", "お使いみち", "お借入残高", "年間返済額"],
               [
                   ["【要記入】＿＿＿＿", "□事業 □住宅 □車 □教育 □カード □その他", "＿＿万円", "＿＿万円"],
                   ["＿＿＿＿", "□事業 □住宅 □車 □教育 □カード □その他", "＿＿万円", "＿＿万円"],
               ],
               widths=[4.6, 7.0, 2.8, 2.8])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============ 7 必要な資金と調達方法 ============
    t7_head = doc.add_table(rows=1, cols=1)
    t7_head.style = "Table Grid"
    c = t7_head.rows[0].cells[0]
    c.width = Cm(17.2)
    cell_text(c, "7　必要な資金と調達方法", size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    shade(c, "203864")
    grid_table(doc,
               ["必要な資金", "見積先等", "金額", "調達の方法", "金額"],
               [
                   ["【設備資金】\n業務用パソコン・周辺機器", "家電量販店（見積書添付）", "20万円",
                    "自己資金", "100万円"],
                   ["【運転資金】\n営業活動費（交通費・USB等資材・研究会参加）", "実費積算", "60万円",
                    "親、兄弟、知人、友人等からの借入", "0万円"],
                   ["広告宣伝費（Webサイト改修・印刷物・広告）", "制作会社見積・実費積算", "60万円",
                    "日本政策金融公庫 国民生活事業からの借入\n（元金据置6ヶ月・返済期間7年 希望）", "300万円"],
                   ["外部委託費（税理士顧問・法務確認・デザイン）", "各事務所見積", "60万円",
                    "他の金融機関等からの借入", "0万円"],
                   ["当面の役員報酬・予備費（12ヶ月分の一部）", "月10万円×12ヶ月＋予備", "200万円", "", ""],
                   ["合　計", "", "400万円", "合　計", "400万円"],
               ],
               widths=[5.8, 3.4, 1.8, 4.4, 1.8], bold_last=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============ 8 事業の見通し ============
    t8_head = doc.add_table(rows=1, cols=1)
    t8_head.style = "Table Grid"
    c = t8_head.rows[0].cells[0]
    c.width = Cm(17.2)
    cell_text(c, "8　事業の見通し（月平均）", size=10.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    shade(c, "203864")
    grid_table(doc,
               ["科目", "創業当初", "軌道に乗った後\n（3年目）", "売上高・経費の根拠"],
               [
                   ["売上高 ①", "3万円", "32万円",
                    "＜創業当初＞導入協力校3校（年29,800円）＋私立1校（月19,800円）の月換算。\n"
                    "＜軌道後＞公立15校（年3〜5万円）＋私立12校（月19,800円）＋教委一括10校（年3万円）の月換算。\n"
                    "単価・校数の詳細は添付「販売計画」参照"],
                   ["売上原価 ②", "0.1万円", "0.4万円", "USB媒体・送料等（1校あたり約1,500円）のみ。仕入れ・在庫なし"],
                   ["経費：人件費", "10万円", "20万円", "役員報酬（創業当初は月10万円に抑制）"],
                   ["経費：家賃", "0万円", "0万円", "自宅を本店とするため発生しない"],
                   ["経費：支払利息", "0.6万円", "0.4万円", "借入300万円・年2.5%想定"],
                   ["経費：その他", "7万円", "11万円", "広告宣伝・交通費・通信費・外注費（税理士等）"],
                   ["経費 計 ③", "17.6万円", "31.4万円", ""],
                   ["利益 ①－②－③", "▲14.7万円", "0.2万円", "創業当初の赤字は自己資金と借入金で賄い、3年目に単月黒字化"],
               ],
               widths=[2.6, 2.2, 2.6, 9.8], bold_last=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ============ 自由記述欄 ============
    section_box(doc, "", "自由記述欄（アピールポイント、事業を行ううえでの悩み、希望するアドバイス等）", [
        "・製品は開発完了済み（v1.0）。実物と同形式のファイルで全機能を検証し、異常データ15項目の",
        "　耐性テストも実施済み（検証記録を添付可能）。開発リスクは残っていない。",
        "・既に有償契約（導入協力校）と、導入前後の作業時間の実測記録があり、需要と効果を定量的に示せる。",
        "・変動費率が約2％と低く、契約校数の増加がほぼそのまま利益改善につながる収益構造。",
        "・契約は年額前受が中心のため資金繰りが安定し、販売先が学校のため貸倒リスクも僅少。",
        "・希望：売上が年度サイクルで立つ事業特性のため、元金据置6ヶ月を希望します。",
    ])

    # ============ 添付: 月別収支計画（1期目） ============
    doc.add_page_break()
    para(doc, "（任意添付）月別収支計画　― 1期目・概算・単位：千円", size=12, bold=True, color=NAVY, space_after=4)
    rows = []
    sales = [10, 10, 10, 20, 20, 20, 30, 30, 30, 49, 49, 49]
    for i, s in enumerate(sales):
        cogs = 1
        exp = 176
        rows.append([f"{i+1}ヶ月目", s, cogs, exp, s - cogs - exp])
    rows.append(["合計", sum(sales), 12, 2112, sum(sales) - 12 - 2112])
    grid_table(doc,
               ["月", "売上高", "売上原価", "経費（役員報酬含む）", "利益"],
               rows, widths=[2.6, 2.8, 2.8, 4.6, 3.0], bold_last=True)
    para(doc, "※売上は契約校数の積み上がりに応じて四半期ごとに増加する想定。経費は月額ほぼ一定"
              "（役員報酬100・広告営業50・外注/通信26）。不足資金は自己資金・借入金で充当。",
         size=9, color=GRAY, space_after=2)

    para(doc, "", space_after=6)
    para(doc, "※本書は公庫「創業計画書」様式の記入欄に対応させた下書きです。提出時は最新の正式様式に"
              "転記のうえ、本書および製品資料・検証記録・契約書類を補足資料として添付してください。",
         size=9, color=GRAY)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, env={**os.environ, "HOME": tempfile.gettempdir()})
    os.replace(os.path.splitext(docx_path)[0] + ".pdf", pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "jfc_startup_plan.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "..", "06_法人化・創業融資", "創業計画書_公庫様式対応版.pdf"))
