# -*- coding: utf-8 -*-
"""学校配布用「オンライン導入手順書」PDFを生成する。

全国展開（郵送USB＋Zoomレクチャー）を想定した、学校側が自力で
Zoom当日までの準備を終えられるようにするための手順書。
生成: docs/04_オンライン導入/オンライン導入手順書.pdf
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
RED = (0xC0, 0x00, 0x00)


def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=13.5, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def step_box(doc, num, title, lines):
    """番号付きの手順ボックス"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"手順{num}　{title}")
    set_jp(r, size=11.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '203864')
    pPr.append(shd)
    for line in lines:
        para(doc, line, size=10.5, space_after=2, indent=0.4)


def checklist(doc, items):
    for t in items:
        para(doc, "□　" + t, size=10.5, space_after=3, indent=0.3)


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("積立金会計 入力アシスタント\nオンライン導入手順書（学校ご担当者さま用）")
    set_jp(r, size=16, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("お手元に USB メモリが2本 届いていることをご確認ください")
    set_jp(r, size=10.5)

    para(doc,
         "この手順書のとおりに準備いただければ、Zoomでのレクチャー（約60分）当日に "
         "すぐ練習を始められます。難しい作業はありません。分からないところは飛ばして、"
         "当日画面を見ながら一緒に行いますのでご安心ください。", space_after=8)

    # ==============================
    h1(doc, "1. 届いた2本のUSBメモリについて")
    para(doc, "USBにはラベルが貼ってあります。差すパソコンが決まっています。", space_after=4)
    para(doc, "■ USB①「記録側PC用」 → 積立金マスターのファイルがある保管用のパソコンへ", bold=True, space_after=2, indent=0.3)
    para(doc, "銀行の振替結果や名簿をもとに、マスターへ入力するためのアシスタント一式が入っています。", size=10, space_after=6, indent=0.6)
    para(doc, "■ USB②「送信側PC用」 → 銀行のデータや名簿を受け取るパソコンへ", bold=True, space_after=2, indent=0.3)
    para(doc, "受け取ったデータを記録側へ渡すための整理フォルダと練習データが入っています。", size=10, space_after=6, indent=0.6)
    para(doc, "※ どちらのUSBにも、生徒の実データは一切入っていません（練習用の架空データのみ）。"
              "万一紛失されても個人情報の事故にはなりません。", size=10, color=RED, space_after=4)

    # ==============================
    h1(doc, "2. Zoom前日までの準備チェックリスト")
    checklist(doc, [
        "保管用パソコンと受取用パソコンの両方が Windows で、Excel（Microsoft 365 / 2016以降）が入っている",
        "学校のルールでUSBメモリが使えることを確認した（使えない場合は事前にご連絡ください。別の受け渡し方法をご案内します）",
        "レクチャーに使うZoomが開ける（カメラは不要、マイクがあれば十分です）",
        "できれば2名で参加（メインのご担当者＋もう1名。引き継ぎのためです）",
        "下の「手順1〜3」を済ませた（うまくいかなくても大丈夫。当日一緒にやります）",
    ])

    # ==============================
    h1(doc, "3. 保管用パソコンでの準備（USB①・所要10分）")
    step_box(doc, 1, "フォルダをパソコンにコピーする", [
        "① USB①を保管用パソコンに差します。",
        "② 中の「積立金入力アシスタント」フォルダを、そのまま デスクトップ にコピーします。",
        "③ 以後の作業はデスクトップにコピーしたほうで行います（USBから直接開かないでください）。",
    ])
    step_box(doc, 2, "マクロを許可する（一番大事な設定です）", [
        "① デスクトップにコピーしたフォルダの中の「積立金入力アシスタント.xlsm」を右クリック →「プロパティ」。",
        "② 一番下に「セキュリティ： このファイルは他のコンピューターから取得したものです…」と表示があれば、"
        "「許可する」にチェックを入れて「OK」。表示が無ければ何もせず閉じてください。",
        "③ Excelを開き、「ファイル」→「オプション」→「トラスト センター」→「トラスト センターの設定」"
        "→「信頼できる場所」→「新しい場所の追加」。",
        "④ 「参照」でデスクトップの「積立金入力アシスタント」フォルダを選んで「OK」。",
        "⑤ 「積立金入力アシスタント.xlsm」を開き、画面上部に黄色や赤の警告バーが出ないことを確認します。",
        "※ ③④の画面に入れない（ボタンが灰色）場合は、学校のPC管理者の設定です。そのまま当日ご相談ください。",
    ])
    step_box(doc, 3, "受取用パソコンでの準備（USB②・所要2分）", [
        "① USB②を受取用パソコンに差します。",
        "② 中の「積立金データ受け渡し」フォルダをデスクトップにコピーします。準備はこれだけです。",
    ])

    # ==============================
    h1(doc, "4. Zoomレクチャー当日の流れ（約60分）")
    rows = [
        ("0〜10分", "ごあいさつ・画面共有の確認。手順2がうまくいっているかを一緒に確認します。"),
        ("10〜20分", "「設定」シートに練習用マスターの場所と年度を入力します（黄色いセル2つだけ）。"),
        ("20〜45分", "練習用データ（架空の80名）で、実際の1か月分の作業をご自身の手で一巡します。"
                    "名簿の取り込み → 支出の一括入力 → 振替結果の照合 → 収入の一括入力。"
                    "未納の2名が自動で見つかるところまで確認します。"),
        ("45〜55分", "年間予定表と、年度替わりの引き継ぎ（○×を付けるだけ）のご説明。"),
        ("55〜60分", "質疑応答。今後1〜2か月の「並走期間」の進め方とご連絡方法を決めます。"),
    ]
    for time, desc in rows:
        para(doc, f"【{time}】 {desc}", size=10.5, space_after=3, indent=0.3)
    para(doc, "レクチャー後も、本物のマスターにはまだ触りません。並走期間中は「コピー」で練習し、"
              "手入力の結果と一致することをご自身で確かめてから切り替えます。", size=10, space_after=4)

    # ==============================
    h1(doc, "5. うまくいかないときは")
    para(doc, "よくあるのは次の3つです。当日Zoomで一緒に解決できますので、そのままにしておいて大丈夫です。", space_after=4)
    para(doc, "・ファイルを開くと「マクロがブロックされました」と赤いバーが出る → 手順2の②（プロパティの許可）が未実施です。", size=10, space_after=2, indent=0.3)
    para(doc, "・「信頼できる場所」が追加できない → 学校の管理者設定です。管理者の方に本手順書の手順2をお見せください。", size=10, space_after=2, indent=0.3)
    para(doc, "・ボタンを押しても動かない → Excelのバージョンが古い可能性があります。事前にご一報ください。", size=10, space_after=6, indent=0.3)

    para(doc, "お困りのときの連絡先：株式会社ココロラボ　info@cocorolab.co.jp", bold=True, space_after=2)
    para(doc, "（エラーの画面をスマートフォンで撮影して添付いただけると、最短で解決できます）", size=9.5)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
        check=True, env={**os.environ, "HOME": tempfile.gettempdir()},
    )
    produced = os.path.splitext(docx_path)[0] + ".pdf"
    os.replace(produced, pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "online_onboarding.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "..", "04_オンライン導入", "オンライン導入手順書.pdf"))
