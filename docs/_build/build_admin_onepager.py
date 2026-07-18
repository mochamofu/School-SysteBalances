# -*- coding: utf-8 -*-
"""学校管理者様向けご説明（技術説明・A4一枚）を生成する。

公立学校で「ソフトウェア導入の許可申請が要るのでは」という懸念に対し、
本製品の実体（Excelブック1つ・インストールなし・通信なし）を
情報担当者・管理職が判断できる形で示す資料。断定的な「許可不要」とは
書かず、判断材料としての事実だけを表で提示する。

生成: docs/01_学校向けマニュアル/学校管理者様向けご説明.pdf
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)


def set_jp(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10, bold=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=12, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths, size=9):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_jp(r, size=size, bold=True, color=(0xFF, 0xFF, 0xFF))
        shade(cell, "203864")
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(v))
            set_jp(r, size=size, bold=(j == 1 and str(v).startswith(("なし", "不要", "行いません"))))
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.5))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("「積立金会計 入力アシスタント」に関する技術的ご説明")
    set_jp(r, size=15, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("― 情報管理ご担当者さま・管理職の皆さまの判断材料としてお使いください ―")
    set_jp(r, size=9.5)

    h1(doc, "1. 本製品の実体")
    para(doc, "本製品は、専用アプリケーションではなく、Microsoft Excel のブック（.xlsm ファイル1つ）と"
              "参照用のExcelファイル・練習用データです。Excelの標準機能であるマクロ（VBA）だけで動作し、"
              "Excel以外のプログラムは一切使用しません。", size=10)

    h1(doc, "2. 導入にあたって「行わないこと」の一覧")
    table(
        doc,
        ["確認項目", "本製品", "補足"],
        [
            ["インストーラーの実行", "なし", "ファイルをフォルダにコピーするだけです"],
            ["管理者権限（Administrator）", "不要", "一般ユーザー権限の範囲で動作します"],
            ["レジストリ・システム設定の変更", "行いません", "OSには手を加えません"],
            ["常駐プログラム・サービスの追加", "なし", "Excelを閉じれば何も動いていません"],
            ["ネットワーク通信", "行いません", "完全オフラインで動作します（外部送信ゼロ）"],
            ["生徒情報の校外持ち出し", "行いません", "データもバックアップも校内PCの中で完結します"],
            ["既存ファイルの改変", "行いません", "現在お使いの積立金ファイルの行・列・様式は変更しません"],
            ["削除（アンインストール）", "ファイル削除のみ", "専用のアンインストール作業は不要です"],
        ],
        widths=[5.8, 3.6, 8.6],
    )

    h1(doc, "3. 必要となる設定（1件のみ）")
    para(doc, "Excelの標準設定画面（ファイル → オプション → トラスト センター）で、本製品を置くフォルダを"
              "「信頼できる場所」に登録します。これはExcelに用意されている通常の設定項目であり、"
              "システムやセキュリティ機構の変更ではありません。設定は導入時の1回のみで、"
              "対象フォルダを限定して登録します。", size=10)

    h1(doc, "4. データの取り扱い")
    for t in [
        "・生徒の氏名・金額・口座情報は、保管用パソコンの中でのみ読み書きされます。",
        "・書き込みの前には必ず実行前のコピー（バックアップ）が同じパソコン内に自動保存されます。",
        "・本製品のファイル自体には生徒の実データは含まれません（同梱の練習用データはすべて架空のものです）。",
    ]:
        para(doc, t, size=10, space_after=2)

    h1(doc, "5. 動作環境")
    para(doc, "Windows 10 / 11 ＋ Microsoft Excel（Microsoft 365 または Excel 2016 以降）。"
              "インターネット接続は不要です。", size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("ご不明な点は、本紙をご覧のうえ導入担当者までお尋ねください。"
                  "校内規程との適合のご判断に必要な追加情報は、随時提供いたします。")
    set_jp(r, size=9.5)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, env={**os.environ, "HOME": tempfile.gettempdir()})
    os.replace(os.path.splitext(docx_path)[0] + ".pdf", pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "admin_onepager.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "..", "01_学校向けマニュアル", "学校管理者様向けご説明.pdf"))
