# -*- coding: utf-8 -*-
"""レクチャー配布資料（職員用）を生成する。

説明を聞きながら手元で見る聴講者用のハンドアウト。
スライド14枚の縮小画像＋メモ欄だけで構成し、
発表者の台本・営業向けの文言は一切入れない。

生成: docs/03_現地デモ・レクチャー/レクチャー配布資料_職員用.pdf
必要: LibreOffice(soffice) と poppler-utils(pdftoppm)
"""
import os
import glob
import shutil
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
PPTX = os.path.join(HERE, "..", "02_営業・商談資料", "積立金入力アシスタント_説明スライド.pptx")
OUT_PDF = os.path.join(HERE, "..", "03_現地デモ・レクチャー", "レクチャー配布資料_職員用.pdf")

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GRAY = (0x88, 0x88, 0x88)


def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def memo_line(cell):
    """メモ用の点線だけの段落（書き込み用の罫線）"""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'dotted')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def slides_to_pngs(workdir):
    ascii_pptx = os.path.join(workdir, "slides.pptx")
    shutil.copy(os.path.abspath(PPTX), ascii_pptx)
    env = {**os.environ, "HOME": workdir}
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", workdir, ascii_pptx], check=True, env=env)
    subprocess.run(["pdftoppm", "-png", "-r", "96",
                    os.path.join(workdir, "slides.pdf"),
                    os.path.join(workdir, "slide")], check=True)
    pngs = sorted(glob.glob(os.path.join(workdir, "slide-*.png")))
    if not pngs:
        raise RuntimeError("スライドのPNG化に失敗しました")
    return pngs


def build():
    workdir = tempfile.mkdtemp(prefix="handout_")
    pngs = slides_to_pngs(workdir)
    print(f"スライド {len(pngs)} 枚をPNG化")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.4))

    # ---- 1ページ目の見出し ----
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("積立金会計 入力アシスタント　ご説明資料")
    set_jp(r, size=17, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("説明をお聞きになりながら、右のメモ欄を自由にお使いください。")
    set_jp(r, size=10)

    # ---- スライド1枚 = 表の1行（左：画像／右：メモ欄） ----
    table = doc.add_table(rows=len(pngs), cols=2)
    table.autofit = False
    # 列幅を固定レイアウトで強制する（LibreOfficeでの右端切れ防止）
    tblPr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    grid = table._tbl.find(qn('w:tblGrid'))
    for col, w in zip(grid.findall(qn('w:gridCol')), (5900, 4300)):  # twips(1/20pt)
        col.set(qn('w:w'), str(w))
    for i, png in enumerate(pngs):
        row = table.rows[i]
        # 行の途中で改ページさせない（番号と画像が泣き別れしない）
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement('w:cantSplit')
        trPr.append(cant)
        left = row.cells[0]
        right = row.cells[1]
        left.width = Cm(10.4)
        right.width = Cm(7.6)

        lp = left.paragraphs[0]
        lp.paragraph_format.space_before = Pt(6)
        lr = lp.add_run(f"{i + 1} / {len(pngs)}")
        set_jp(lr, size=8.5, color=GRAY)
        ip = left.add_paragraph()
        ip.paragraph_format.space_after = Pt(6)
        ip.add_run().add_picture(png, width=Cm(9.6))

        rp = right.paragraphs[0]
        rp.paragraph_format.space_before = Pt(6)
        rr = rp.add_run("メモ")
        set_jp(rr, size=8.5, color=GRAY)
        for _ in range(5):
            memo_line(right)

    # ---- 最終ページ：質問メモと連絡先 ----
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("ご質問・確認したいこと")
    set_jp(r, size=13, bold=True, color=NAVY)
    qtable = doc.add_table(rows=1, cols=1)
    cell = qtable.rows[0].cells[0]
    cell.width = Cm(18)
    for _ in range(10):
        memo_line(cell)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("※本日の説明で使用したデータはすべて架空のものです。実在の生徒情報は含まれていません。")
    set_jp(r, size=9, color=GRAY)

    # ---- 保存 → PDF ----
    docx_path = os.path.join(workdir, "handout.docx")
    doc.save(docx_path)
    env = {**os.environ, "HOME": workdir}
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", workdir, docx_path], check=True, env=env)
    os.replace(os.path.join(workdir, "handout.pdf"), os.path.abspath(OUT_PDF))
    print("pdf saved:", os.path.abspath(OUT_PDF))
    shutil.rmtree(workdir, ignore_errors=True)


if __name__ == "__main__":
    build()
