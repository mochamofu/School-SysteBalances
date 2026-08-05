# -*- coding: utf-8 -*-
"""2台PC実機テスト手順書（自分用・データ反映の検証）を生成する。

手持ちのノートPC2台で「USBで渡したデータがマスターに正しく
反映されるか」を確かめるための、当日そのまま進行できる手順書。
学校へのテスト導入の前段階（自分による最終検証）に使う。

生成: docs/05_開発・検証記録/2台PC実機テスト手順書.pdf
"""
import os
import math
import subprocess
import tempfile
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PDF = os.path.join(HERE, "..", "05_開発・検証記録", "2台PC実機テスト手順書.pdf")

FONT_PATH = "/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf"
FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GREEN = (0x1E, 0x7A, 0x3C)
RED = (0xC0, 0x00, 0x00)
GRAY = (0x66, 0x66, 0x66)
NAVY_S = "#203864"
BLUE_S = "#2F5FA8"
AMBER_S = "#C58F00"
GREEN_S = "#2E7D32"
RED_S = "#C00000"
GRAY_S = "#666666"
LGRAY_S = "#E8EDF5"

_workdir = None


def F(size):
    return ImageFont.truetype(FONT_PATH, size)


def _save(img, name):
    path = os.path.join(_workdir, name)
    img.save(path)
    return path


def txt(d, x, y, s, size, fill=NAVY_S, stroke=0):
    d.text((x, y), s, font=F(size), fill=fill, stroke_width=stroke, stroke_fill=fill)


def ctext(d, cx, y, s, size, fill=NAVY_S, stroke=0):
    w = d.textlength(s, font=F(size))
    d.text((cx - w / 2, y), s, font=F(size), fill=fill, stroke_width=stroke, stroke_fill=fill)


def arrow(d, x1, y1, x2, y2, color=AMBER_S, width=9):
    d.line((x1, y1, x2, y2), fill=color, width=width)
    ang = math.atan2(y2 - y1, x2 - x1)
    L = 28
    for da in (2.6, -2.6):
        d.line((x2, y2, x2 - L * math.cos(ang + da), y2 - L * math.sin(ang + da)),
               fill=color, width=width)


def pc(d, cx, y, label, sub, w=300, h=190, color=NAVY_S):
    d.rectangle((cx - w // 2, y, cx + w // 2, y + h), outline=color, width=6, fill="white")
    ctext(d, cx, y + 40, label, 34, color, 1)
    for i, line in enumerate(sub.split("\n")):
        ctext(d, cx, y + 92 + i * 34, line, 25, GRAY_S)
    d.polygon([(cx - w // 2 - 34, y + h + 44), (cx + w // 2 + 34, y + h + 44),
               (cx + w // 2, y + h), (cx - w // 2, y + h)],
              outline=color, fill=LGRAY_S, width=5)


def usb(d, x, y, label, color):
    d.rounded_rectangle((x, y + 24, x + 140, y + 88), radius=12, fill="white",
                        outline=NAVY_S, width=5)
    d.rectangle((x + 140, y + 40, x + 172, y + 72), outline=NAVY_S, width=4, fill=LGRAY_S)
    d.ellipse((x + 6, y, x + 68, y + 62), fill=color, outline=NAVY_S, width=4)
    ctext(d, x + 37, y + 10, label, 36, "white", 1)


# ------------------------------------------------------------------
def fig_overview():
    img = Image.new("RGB", (1560, 620), "white")
    d = ImageDraw.Draw(img)
    ctext(d, 780, 12, "テストの全体像 ― 学校で起きることを、2台のPCで再現する", 38, NAVY_S, 1)

    pc(d, 330, 110, "PC② 送信側", "銀行データ・名簿が\n届くパソコン", color=AMBER_S)
    pc(d, 1180, 110, "PC① 記録側", "積立金マスターがある\n保管用パソコン", color=BLUE_S)

    usb(d, 690, 190, "→", GREEN_S)
    arrow(d, 500, 225, 680, 225, AMBER_S)
    arrow(d, 875, 225, 1020, 225, AMBER_S)
    ctext(d, 780, 130, "USBで運ぶ", 28, GREEN_S, 1)

    # 下段: 検証の対象
    d.rounded_rectangle((240, 420, 1320, 580), radius=16, fill="#EAF6EA",
                        outline=GREEN_S, width=5)
    ctext(d, 780, 442, "★ このテストで確かめること", 32, GREEN_S, 1)
    ctext(d, 780, 490, "PC②から運んだデータが、PC①のマスターの「正しいセル」に、「正しい数値」で入るか", 27, NAVY_S)
    ctext(d, 780, 530, "（記録ブックの「③ データ反映の突合」でセル単位まで確認します）", 25, GRAY_S)
    return _save(img, "ft_overview.png")


def fig_flow():
    img = Image.new("RGB", (1560, 420), "white")
    d = ImageDraw.Draw(img)
    steps = [
        ("準備1", "2台に\nフォルダを配置", "15分"),
        ("準備2", "PC①にVBAを\n組み込む（初回のみ）", "30分"),
        ("テストA", "名簿を運んで\nクラス替え反映", "20分"),
        ("テストB", "振替結果を運んで\n照合・入金記録", "25分"),
        ("テストC", "支出を一括入力", "15分"),
        ("テストD", "320名で本番規模", "20分"),
        ("テストE", "わざと壊す\n（安全装置）", "10分"),
    ]
    x = 30
    for i, (tag, body, mins) in enumerate(steps):
        w = 195
        fill = "#FFF7E0" if i < 2 else "white"
        outline = AMBER_S if i < 2 else BLUE_S
        d.rounded_rectangle((x, 90, x + w, 300), radius=14, fill=fill, outline=outline, width=5)
        ctext(d, x + w / 2, 108, tag, 28, outline, 1)
        lines = body.split("\n")
        ty = 158 if len(lines) > 1 else 178
        for ln in lines:
            ctext(d, x + w / 2, ty, ln, 23, NAVY_S)
            ty += 32
        ctext(d, x + w / 2, 258, mins, 24, GRAY_S)
        if i < len(steps) - 1:
            arrow(d, x + w + 4, 195, x + w + 22, 195, BLUE_S, 6)
        x += w + 26
    ctext(d, 780, 30, "当日の流れ ― 合計 約2時間（準備45分＋テスト90分）", 34, NAVY_S, 1)
    ctext(d, 780, 340, "※準備2（VBAの組み込み）は初回だけ。2回目以降は .xlsm をコピーするだけで済みます", 24, RED_S)
    return _save(img, "ft_flow.png")


def fig_cellcheck():
    img = Image.new("RGB", (1560, 560), "white")
    d = ImageDraw.Draw(img)
    ctext(d, 780, 10, "データ反映の突合 ― マスターのどこを見るか", 36, NAVY_S, 1)

    # Excelのグリッド風
    left, top = 90, 90
    colw, rowh = 118, 52
    cols = ["A", "…", "E", "…", "G", "H", "…", "J", "…", "BE", "…", "FC"]
    d.rectangle((left, top, left + colw * len(cols), top + rowh), fill=LGRAY_S,
                outline="#999999", width=3)
    for i, c in enumerate(cols):
        ctext(d, left + colw * i + colw / 2, top + 12, c, 24, NAVY_S)
    rows = [("9", "精算番号1の生徒"), ("15", "精算番号7（未納）"), ("331", "合計行")]
    for j, (rn, note) in enumerate(rows):
        y = top + rowh * (j + 1)
        d.rectangle((left, y, left + colw * len(cols), y + rowh), fill="white",
                    outline="#CCCCCC", width=2)
        txt(d, left - 68, y + 12, rn, 24, GRAY_S)
        txt(d, left + colw * len(cols) + 16, y + 12, note, 22, GRAY_S)

    # ハイライト
    def hl(col_idx, row_idx, label, color):
        x = left + colw * col_idx
        y = top + rowh * (row_idx + 1)
        d.rectangle((x + 3, y + 3, x + colw - 3, y + rowh - 3), fill=color, outline=RED_S, width=4)
        ctext(d, x + colw / 2, y + 14, label, 21, NAVY_S)

    hl(7, 0, "76,000", "#DFF0DF")     # J9
    hl(9, 0, "3,500", "#DFF0DF")      # BE9
    hl(11, 0, "3,500", "#EAF0FA")     # FC9
    hl(7, 1, "空欄", "#FFE8E8")        # J15
    hl(9, 2, "280,000", "#DFF0DF")    # BE331

    ty = top + rowh * 4 + 40
    txt(d, 90, ty, "J列 = 収入枠1（⑤で入れた入金額）", 26, GREEN_S)
    txt(d, 90, ty + 40, "BE列 = 支出No.1（④で入れた支出額）", 26, GREEN_S)
    txt(d, 90, ty + 80, "FC列 = 支出合計（数式が自動で反応する列）", 26, BLUE_S)
    txt(d, 820, ty, "★ 未納者の行が「空欄のまま」であることが重要", 26, RED_S)
    txt(d, 820, ty + 40, "　（空欄だからこそ、H列の未納印が自動で立つ）", 24, GRAY_S)
    txt(d, 820, ty + 80, "★ 合計行が正しい ＝ 全員分が漏れなく入った証拠", 26, RED_S)
    return _save(img, "ft_cells.png")


# ==================================================================
def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=4, center=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=13.5, bold=True, color=(0xFF, 0xFF, 0xFF))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), '203864')
    pPr.append(shd)
    return p


def pic(doc, path, width_cm=17.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths, size=9.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        set_jp(r, size=size, bold=True, color=(0xFF, 0xFF, 0xFF))
        shade(c, "203864")
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = ""
            first = True
            for line in str(v).split("\n"):
                p = c.paragraphs[0] if first else c.add_paragraph()
                first = False
                p.paragraph_format.space_after = Pt(1)
                rr = p.add_run(line)
                set_jp(rr, size=size)
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def steps(doc, items):
    for i, t in enumerate(items, start=1):
        para(doc, f"{i}. {t}", size=10.5, space_after=3, indent=0.4)


def expect(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run("✔ 期待値： " + text)
    set_jp(r, size=10.5, bold=True, color=GREEN)


def build():
    global _workdir
    _workdir = tempfile.mkdtemp(prefix="fieldtest_")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.4))

    para(doc, "2台PC 実機テスト手順書", size=20, bold=True, color=NAVY, center=True, space_after=2)
    para(doc, "― 学校へのテスト導入の前に、自分の手で「データが正しく反映されるか」を確かめる ―",
         size=11, center=True, space_after=6)
    para(doc, "所要 約2時間　／　使うデータはすべて架空（本物の生徒情報は一切使いません）",
         size=10, color=GRAY, center=True, space_after=8)

    pic(doc, fig_overview(), 17.0)
    pic(doc, fig_flow(), 17.0)

    # ============ 持ち物 ============
    doc.add_page_break()
    h1(doc, "0. 用意するもの")
    table(doc,
          ["物", "内容", "備考"],
          [
              ["ノートPC 2台", "PC①=記録側（Excel必須）／PC②=送信側", "PC②はExcelなしでも可"],
              ["USBメモリ 1本", "PC②→PC①へデータを運ぶ", "中身は空でよい"],
              ["実機テストキット", "assistant/output/実機テストキット.zip", "展開して2台に配置"],
              ["この手順書（印刷）", "手を動かしながら見るため紙推奨", ""],
              ["テスト結果記録ブック", "キット内 03_記録用紙/ にあるExcel", "PC①で開いて記入"],
          ],
          widths=[3.4, 8.2, 5.6])

    # ============ 準備1 ============
    h1(doc, "準備1　2台にフォルダを配置する（15分）")
    steps(doc, [
        "実機テストキット.zip を展開する。",
        "「01_PC1_記録側」フォルダを丸ごと、PC①のデスクトップにコピーする。",
        "「02_PC2_送信側」フォルダを丸ごと、PC②のデスクトップにコピーする。",
        "「03_記録用紙」フォルダは PC①のデスクトップへ（記入しながら進めるため）。",
    ])
    expect(doc, "両方のPCのデスクトップに、それぞれのフォルダが見える")

    # ============ 準備2 ============
    h1(doc, "準備2　PC①にVBAを組み込む ★初回だけ（30分）")
    para(doc, "この作業は一度だけです。出来上がった .xlsm は、今後そのままコピーして使い回せます"
              "（学校に渡すときもこのファイル）。", size=10.5)
    steps(doc, [
        "PC①の「1_アシスタント」フォルダの 積立金入力アシスタント.xlsx を開く。",
        "Alt キーを押しながら F11 を押す（VBAの画面が開く）。",
        "メニューの「ファイル」→「ファイルのインポート」を選ぶ。",
        "「VBAモジュール」フォルダの中の A00〜A11（12個のファイル）を選んで取り込む。"
        "（Ctrlを押しながら全部選べば一度に入ります）",
        "左側のツリーに12個のモジュールが並んだことを確認する。",
        "Alt キーを押しながら F11 を押してExcelに戻る。",
        "Alt キーを押しながら F8 を押す →「初期設定」を選んで「実行」。",
        "「メニュー」シートにボタン①〜⑮が並べば成功。",
        "「名前を付けて保存」→ ファイルの種類を「Excelマクロ有効ブック (*.xlsm)」にして保存。",
    ])
    expect(doc, "積立金入力アシスタント.xlsm ができ、メニューに15個のボタンが並んでいる")
    para(doc, "▲ うまくいかないとき：コンパイルエラーが出たら、VBA画面のメニュー「デバッグ」→"
              "「VBAProjectのコンパイル」を実行し、赤くなった行を撮影して連絡してください。",
         size=10, color=RED, space_after=4)

    h1(doc, "準備3　設定シートを埋める（5分）")
    steps(doc, [
        "できた .xlsm を開き、「設定」シートを表示する。",
        "C3（マスターファイルの場所）に、PC①の「2_練習用データ」の中の "
        "練習用_令和X年度生積立金.xlsx のフルパスを入れる。"
        "※パスはファイルを Shift+右クリック →「パスのコピー」で取得できます（前後の \" は消す）",
        "C5（年度）に 7 など任意の数字を入れる。",
        "C7（口座マスターの場所）に 練習用_口座マスター.xlsx のフルパスを入れる。",
    ])
    expect(doc, "黄色いセル3か所が埋まっている")

    # ============ テストA ============
    doc.add_page_break()
    h1(doc, "テストA　名簿を運んでクラス替えを反映する（20分）")
    para(doc, "★ここからが本番。「PC②で受け取ったデータをUSBでPC①へ運ぶ」を必ず実際にやってください。",
         size=10.5, bold=True, color=RED)
    steps(doc, [
        "【PC②】「1_届いたデータ」の 練習用_掲示用名簿.xlsx を「2_PC1へ運ぶ箱」にコピーする。",
        "【PC②】USBメモリを挿し、そのファイルをUSBにコピーする。",
        "【PC①】USBを挿し、ファイルをデスクトップにコピーして開く。",
        "【PC①】名簿のシート全体をコピーし、アシスタントの「名簿貼付」シートのA1に貼り付ける。",
        "【PC①】メニューの ①名簿を解析して照合する を押す。",
        "記録ブック「② 動作テスト」のNo.3〜5に、出てきた数値を記入する。",
        "「名簿一覧」シートで結果を確認し、メニューの ②クラス替えをマスターに反映する を押す。",
        "記録ブックのNo.6に人数を記入する。",
    ])
    expect(doc, "4クラス・80名を検出／全員「一致」／反映80名")
    para(doc, "※このとき「バックアップを作成しました」というメッセージが出ます。"
              "練習用データと同じフォルダに「バックアップ」フォルダができていることも確認してください"
              "（これが安全装置の実物です）。", size=10, color=GRAY)

    # ============ テストB ============
    h1(doc, "テストB　振替結果を運んで照合・入金記録（25分）★最重要")
    steps(doc, [
        "【PC②】練習用_振替結果.xlsx を「2_PC1へ運ぶ箱」→USBへコピー。",
        "【PC①】USBから取り込み、ファイルを開く。",
        "【PC①】データ部分（口座記号・口座番号・金額・振替結果の4列×80行）をコピーする。"
        "※1行目の注意書きと2行目の見出しは含めない。3行目からが本体です。",
        "【PC①】アシスタントの「振替結果取込」シートの B12 を選んで貼り付ける。",
        "【PC①】メニューの ⑪振替結果を照合 を押す。",
        "出てきた数値を記録ブックのNo.9〜12に記入する。",
        "「収入入力」シートを開き、未納者表に自動で名前が入っていることを確認（No.13に人数を記入）。",
        "収入入力シートに 収入枠No=空き枠／件名=口座振替テスト／金額=76000 を入れて "
        "⑤収入をマスターへ一括入力 を押す。",
        "出てきた人数を記録ブックのNo.14に記入する。",
    ])
    expect(doc, "読取80／振替済78／未納2（精算番号7と44）／不明0　→　入金あり78名・未納2名")

    # ============ テストC ============
    h1(doc, "テストC　支出を一括入力する（15分）")
    steps(doc, [
        "「支出入力」シートに 支出No=1／件名=校外学習バス代／支払先=○○観光／"
        "日付=今日／一人あたり金額=3500／対象=全員 を入力。",
        "メニューの ④支出をマスターへ一括入力 を押す。",
        "出てきた数値を記録ブックのNo.7〜8に記入する。",
        "「支出承認書」シートが自動で埋まっていることを確認する。",
    ])
    expect(doc, "対象80名／合計280,000円／支出承認書が自動作成される")

    # ============ データ反映の突合 ============
    doc.add_page_break()
    h1(doc, "★ データ反映の突合 ― ここが今日のいちばん大事な確認（20分）")
    para(doc, "テストA〜Cが終わったら、練習用マスター（練習用_令和X年度生積立金.xlsx）を開き、"
              "「データ」シートの以下のセルを実際に見て、記録ブック「③ データ反映の突合」に書き写します。",
         size=10.5)
    pic(doc, fig_cellcheck(), 17.0)
    para(doc, "セルの探し方：Excelの左上の名前ボックス（列名の左）に「BE9」などと打ってEnterを押すと"
              "そのセルに飛べます。", size=10, color=GRAY)
    para(doc, "この突合で「入るべき所に入り、入ってはいけない所に入っていない」ことが確認できれば、"
              "製品としての動作は保証できます。", size=10.5, bold=True, color=GREEN)

    # ============ テストD ============
    h1(doc, "テストD　本番と同じ320名規模で試す（20分）")
    steps(doc, [
        "「設定」シートのC3を 検証用_令和X年度生積立金.xlsx に、C7を 検証用_口座マスター.xlsx に変更。",
        "【PC②→USB→PC①】検証用_新年度名簿.xlsx を運び、「名簿貼付」に貼って ①を実行。",
        "記録ブックのNo.18に人数を記入。",
        "【PC②→USB→PC①】検証用_振替結果.xlsx（321行）を運び、B12に貼って ⑪を実行。",
        "記録ブックのNo.19〜20に記入。",
    ])
    expect(doc, "8クラス320名／読取321・振替済315・未納5・不明1")
    para(doc, "※「不明1件」は異常ではありません。見つからない口座を正しく検出できるか試すため、"
              "わざと仕込んであります。0件だったら逆に問題です。", size=10, color=RED)

    # ============ テストE ============
    h1(doc, "テストE　わざと壊してみる（安全装置の確認・10分）")
    table(doc,
          ["やること", "期待される動き"],
          [
              ["支出入力の例外表に 精算番号 400 を入れて ④を実行", "「範囲外です」とエラーが出て停止する（マスターには何も書かれない）"],
              ["「設定」シートのC3を空にして ④を実行", "設定を促す案内が出て停止する"],
              ["すでに金額が入っている支出No.1を指定して ④を実行", "「すでに○名分の金額が入っています」と確認が出る（いいえで中止できる）"],
              ["C3に、マスターではない普通のExcelを指定して ④を実行", "「データシートの形が想定と違います」と出て書き込みを拒否する"],
          ],
          widths=[8.0, 9.2])
    para(doc, "4つとも「止まる」ことが正解です。止まらずに書き込まれたら重大な不具合なので、"
              "必ず記録ブック④に記入して連絡してください。", size=10.5, bold=True, color=RED)

    # ============ 終わったら ============
    h1(doc, "テストが終わったら")
    steps(doc, [
        "記録ブック「① 実施情報」の総合判定を見る（全項目OKと出れば合格）。",
        "×があった項目は「④ 不具合記録」に4項目（ボタン名・エラー文言・直前の操作・使用データ）を記入。",
        "エラー画面の写真とあわせて連絡する。修正版の .bas ファイルをお返しします。",
        "記録ブックは保存しておく（学校への提示・融資面談の資料としてそのまま使えます）。",
    ])
    para(doc, "全項目○になれば、学校へのテスト導入に進める状態です。おつかれさまでした。",
         size=11.5, bold=True, color=GREEN, space_after=2)

    docx_path = os.path.join(_workdir, "field_test.docx")
    doc.save(docx_path)
    env = {**os.environ, "HOME": _workdir}
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", _workdir, docx_path], check=True, env=env)
    os.replace(os.path.join(_workdir, "field_test.pdf"), os.path.abspath(OUT_PDF))
    print("pdf saved:", os.path.abspath(OUT_PDF))


if __name__ == "__main__":
    build()
