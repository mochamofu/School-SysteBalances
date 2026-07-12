# -*- coding: utf-8 -*-
"""現地レクチャー指南書（発表者用・画像付き）を生成する。

説明スライド(pptx)を1枚ずつPNG化し、各スライドの隣に
「話すこと」「操作すること」「転んだときの逃げ道」を並べた
発表者専用の進行台本PDFを作る。オフライン（ネット無し）前提。

生成: docs/03_現地デモ・レクチャー/現地レクチャー指南書_発表者用.pdf
必要: LibreOffice(soffice) と poppler-utils(pdftoppm)
"""
import os
import glob
import shutil
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
PPTX = os.path.join(HERE, "..", "02_営業・商談資料", "積立金入力アシスタント_説明スライド.pptx")
OUT_PDF = os.path.join(HERE, "..", "03_現地デモ・レクチャー", "現地レクチャー指南書_発表者用.pdf")

SCHOOL = "都立長山高校"   # ←対象校が変わったらここだけ直して再生成

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
RED = (0xC0, 0x00, 0x00)
GREEN = (0x1E, 0x7A, 0x3C)


# ---------- docx ヘルパー ----------
def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=3, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=13.5, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def tag_line(doc, tag, text, color, size=10.5):
    """【話す】【操作】【逃げ道】のような行"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.2)
    r1 = p.add_run(tag + "　")
    set_jp(r1, size=size, bold=True, color=color)
    r2 = p.add_run(text)
    set_jp(r2, size=size)
    return p


def checklist(doc, items):
    for t in items:
        para(doc, "□　" + t, size=10.5, space_after=2, indent=0.3)


# ---------- スライド → PNG ----------
def slides_to_pngs(workdir):
    """pptxをPDF化→1枚ずつPNGにして、ページ順のパスのリストを返す"""
    # LibreOfficeが日本語ファイル名の型判定に失敗する環境があるためASCII名にコピー
    ascii_pptx = os.path.join(workdir, "slides.pptx")
    shutil.copy(os.path.abspath(PPTX), ascii_pptx)
    env = {**os.environ, "HOME": workdir}
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", workdir, ascii_pptx], check=True, env=env)
    pdf = os.path.join(workdir, "slides.pdf")
    subprocess.run(["pdftoppm", "-png", "-r", "96", pdf,
                    os.path.join(workdir, "slide")], check=True)
    pngs = sorted(glob.glob(os.path.join(workdir, "slide-*.png")))
    if not pngs:
        raise RuntimeError("スライドのPNG化に失敗しました")
    return pngs


# ---------- スライドごとの台本 ----------
# (話すこと[複数行], 操作, 逃げ道/補足)  操作・逃げ道はNoneで省略
SCRIPT = [
    (  # 1 タイトル
     ["「本日は資料のご説明というより、実際に動くところをご覧いただくのが中心です。"
      "320人分の入力が1クリックになる瞬間を、この場で実演します」",
      "自己紹介は30秒で切り上げ、すぐ次へ。冒頭で「今日使うのはすべて架空の生徒データです」と宣言する"],
     None, None),
    (  # 2 本日の内容
     ["アジェンダは読み上げない。「一番ご覧いただきたいのは2つ目、毎月の口座振替の山場です」とだけ言う",
      "ハンズオン（15分）があることを予告→「後ほど皆さまにも触っていただきます」で当事者化する"],
     None, None),
    (  # 3 全体像
     ["「新しいシステムを入れる話ではありません。いまの積立金マスターは1ミリも変えず、"
      "隣に『入力の代行係』を1ファイル置くだけです」",
      "図の矢印を左から右へ指でなぞる：届くデータ → アシスタント → マスターと帳票",
      "「やめたら元のファイルがそのまま残ります。囲い込みはありません」を必ず言う"],
     None, None),
    (  # 4 設定
     ["「準備はこの黄色いセルを3ヶ所埋めるだけ。初回の1回だけで、以後は触りません」",
      "「マスターの場所・年度・口座マスター。導入日は私が一緒に入れますのでご安心を」"],
     None, None),
    (  # 5 メニュー
     ["「業務とボタンが1対1です。『クラス替えの日は①②』『振替結果が届いたら⑪⑤』。"
      "Excelの知識ではなく、いつもの業務の言葉で並んでいます」",
      "ここは長居しない（30秒）。次のデモへの助走"],
     None, None),
    (  # 6 振替結果（山場①）
     ["「ここからが本番です。ゆうちょから振替結果が届いた日を再現します」",
      "★数字を先に予告する：「80件貼ります。振替できなかった生徒が2人混ざっています。"
      "いまから3秒でその2人が見つかります」"],
     "【デモA】へ切替（下の手順カード参照）。⑪を実行して結果を見せてからスライドに戻る",
     "デモが動かない場合：このスライド自体が実演画像なので、図を指しながら口頭説明で続行できる"),
    (  # 7 未納者表・⑤
     ["「見つけた未納者は自動で未納者表に入ります。⑤を押せば残り78人に一括で入金記録、"
      "未納の2人には自動で未納の印が立ちます」",
      "「いままで目視で探していた作業が、貼る→ボタン2回、で終わりです」→ここで一拍おく（一番の売り）"],
     "【デモA】の続き：⑤実行→マスターのH列の未納印を見せる",
     None),
    (  # 8 支出
     ["「支出は『全員同じ金額、例外だけ違う』が実務ですよね。なので例外の生徒だけ表に書きます。"
      "転退学は0、給付型は0、途中参加は半額、という書き方です」",
      "「同時に支出承認書も自動で埋まります。二度書きがなくなります」"],
     None, None),
    (  # 9 クラス替え（山場②）
     ["「年に1回の320人クラス替え。いままで何日かかっていましたか？と聞く（答えを待つ）」",
      "「掲示用の名簿をそのまま貼って、ボタン2つです。人が見るのは赤い行"
      "――転入生と同姓同名だけです」"],
     "【デモB】へ切替：①名簿解析→4クラス80名全員一致→②反映",
     "時間が押していたらデモBは省略可。スライドの図で完結できる"),
    (  # 10 年間予定表
     ["「年間の徴収計画は年度初めに1回書けば、毎回は『行番号を呼ぶだけ』です」",
      "「年度末には支出を業者ごとに集計して、来年も使う項目に○を付けるだけで"
      "翌年度の予定表に引き継がれます。毎年のファイル作りが○×で済みます」"],
     None, None),
    (  # 11 締めの点検と帳票
     ["「月末の締めも毎月同じ流れです。整合性チェックで収支を点検し、精算書はPDFでも印刷でも。"
      "マスターと一緒に保管用PCへ戻して終わりです」",
      "「決算のときは⑦を押せば、全項目の人数・一人あたり・執行総額が一覧で出ます。"
      "電卓の集計作業がなくなります」"],
     None, None),
    (  # 12 安全装置
     ["「事務室で一番大事なのは『壊さないこと』だと思います」と切り出す",
      "4枚のカードを1枚ずつ：自動バックアップ／構造チェック／上書き確認／触らない設計",
      "「書き込みのたびに実行前のコピーが自動で残ります。間違えても、いつでも戻せます」"],
     "ハンズオン（15分）はこのスライドの後に挟む：【デモC】参加者に④支出一括入力を操作してもらう",
     "ハンズオンで参加者がエラーを出したら喜ぶ：「いま止まりましたよね。壊れる前に止まるんです」"),
    (  # 13 料金
     ["「月額はファイルの代金ではありません。ファイルは無料トライアルでお渡しします。"
      "お支払いいただくのは、エラー対応・毎年の年度更新の立ち会い・担当交代時の再レクチャー・"
      "様式が変わったときの改修、この4つの引き受け料です」",
      "比較の詳細は口頭で言わず、価格A4資料を配って「稟議にはこちらをお使いください」"],
     None,
     "「高い」と言われたら：「いま手入力に月5〜10時間、人件費換算で月1〜2万円かかっています。"
     "その置き換えです」とだけ返す"),
    (  # 14 クロージング
     ["「ご判断は今日でなくて結構です。実物のコピーで30日試して、手入力と1円でも違ったら"
      "この話は忘れてください」",
      "★金額の合意ではなく日付の合意で締める：「トライアルを始める日だけ、決めさせてください」"],
     None, None),
]

DEMO_CARDS = [
    ("デモA　振替結果の照合（スライド6〜7で実施・所要5分）",
     [
         "1. Excelに切替え、練習用マスターが「設定」C3に入っていることを事前確認（前日までに）",
         "2. 練習用_振替結果.xlsx を開き、データ部分（4列×80行）をコピー",
         "3. アシスタントの「振替結果取込」シートB12を選んで貼り付け",
         "4. 「実行しますよ」と一声かけて ⑪振替結果を照合 を押す",
     ],
     "読取 80件／振替済 78件／未納 2名（精算番号7と44）／不明口座 0件",
     [
         "5. G〜I列に精算番号・氏名・判定が並んだのを見せる",
         "6. 「収入入力」シートへ→未納者表に2名が自動で入っている",
         "7. 金額欄に 76000 を入れて ⑤収入をマスターへ一括入力",
         "8. 練習用マスターを開き、未納2名のH列に未納印が立ったのを見せる",
     ]),
    ("デモB　クラス替えの名簿照合（スライド9で実施・所要3分）",
     [
         "1. 練習用_掲示用名簿.xlsx を開き、シート全体をコピー",
         "2. アシスタントの「名簿貼付」シートA1に貼り付け",
         "3. ①名簿を解析して照合する を押す",
     ],
     "4クラス80名を検出・全員一致（赤い行なし）",
     [
         "4. 「名簿一覧」シートで結果を見せる：「実物では転入生だけ赤くなり、そこだけ人が見ます」",
         "5. 時間があれば ②クラス替えをマスターに反映する まで実行",
     ]),
    ("デモC　ハンズオン：支出の一括入力（スライド12の後・所要15分）",
     [
         "1. 参加者に席を替わってもらう（必ず本人に操作させる。見学では覚えない）",
         "2. 「支出入力」シートに入力してもらう：支出No=空き番号／件名=校外学習バス代／"
         "日付=当日／金額=3500／対象=全員",
         "3. ④支出をマスターへ一括入力 を押してもらう",
     ],
     "対象80名／合計 280,000円。「支出承認書」シートが自動で埋まる",
     [
         "4. わざと失敗も体験してもらう：例外表に精算番号 400 を入れて④→エラーで止まる",
         "5. 「間違えても書き込む前に止まります。壊れません」で締める",
     ]),
]

QA = [
    ("パソコンが苦手な職員でも使えますか",
     "覚えるのはボタンだけです。今日のハンズオンで実際に押していただいた操作がすべてです。"
     "導入時に練習用データで一緒に練習し、担当が替わったら再レクチャーに伺います。"),
    ("Excelのマクロは学校のセキュリティ的に大丈夫?",
     "ネット通信は一切せず、生徒データはこのPCから出ません。初回に1度だけ"
     "「信頼できる場所」の設定をします（管理者権限が必要ならその手順書もお渡しします）。"),
    ("いまのマスターファイルが壊れたりしませんか",
     "書き込みのたびに実行前のコピーが自動保存されます。またマスターの形が想定と違う場合は"
     "書き込み自体を拒否します。行や列の追加・削除は一切しません。"),
    ("作った人がいなくなったら終わりでは（属人化の心配）",
     "まさにそれを解決する月額です。改修・年度更新・再レクチャーを契約で引き受けます。"
     "逆に解約されても、いまのマスターは何も変わらず手元に残ります。"),
    ("無料のマクロやテンプレートと何が違うの",
     "いまのマスターに合わせて作ってあること、壊さない安全装置、そして毎年の伴走です。"
     "買い切りのテンプレートは様式が合わず、作者が去ると直せなくなります。"),
    ("導入を決める前に試せますか",
     "30日間、実物のコピーでお試しください。手入力の結果と1円でも違えば契約不要です。"
     "今日このあと、トライアル開始日だけ決めさせてください。"),
]


def build():
    workdir = tempfile.mkdtemp(prefix="playbook_")
    pngs = slides_to_pngs(workdir)
    print(f"スライド {len(pngs)} 枚をPNG化")

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.5))

    # ---- 表紙 ----
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n現地レクチャー指南書（発表者用）")
    set_jp(r, size=22, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"対象校：{SCHOOL}　｜　所要60分　｜　オフライン実施（ネット接続なし前提）")
    set_jp(r, size=12, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("スライド1枚ごとに「話すこと」と「操作」を並べた進行台本です。当日はこれだけ見れば進行できます。")
    set_jp(r, size=10.5)

    # ---- 持ち物・前日準備 ----
    h1(doc, "0. 前日までの準備と持ち物（オフライン前提）")
    checklist(doc, [
        "ノートPC2台（フル充電＋ACアダプタ）。1台目=スライド投影＋デモ、2台目=送信側PC役兼バックアップ",
        "デモキットUSB 2本（本体＋同内容の予備）。前日に自分のPCで⑪①④を各1回リハーサル済みであること",
        "スライドはpptxに加えてPDF版も両方のPCとUSBに入れる（PowerPointが無い環境対策）",
        "HDMIとVGA両対応の映像アダプタ、電源タップ（教室のコンセントは遠い）",
        "印刷物：本指南書／業務改善提案書／価格とサービスのご説明（人数分＋2部）／スライドの紙焼き1部（プロジェクタ全滅時の紙芝居用）",
        "アシスタントの「設定」C3・C7が練習用ファイルのパスになっているか出発前に確認",
        "実在の生徒データはPC・USBのどこにも入れていないことを最終確認",
    ])

    # ---- タイムテーブル ----
    h1(doc, "1. 当日の時間割（60分）")
    for t in [
        "00-05分　挨拶・接続。「今日は実演中心・全部架空データ」を宣言（スライド1〜2）",
        "05-15分　これは何か：全体像・設定・メニュー（スライド3〜5）",
        "15-30分　山場：振替結果の照合を実演【デモA】（スライド6〜7）",
        "30-38分　支出とクラス替え【デモB】（スライド8〜9）",
        "38-42分　年間予定表・締めの点検（スライド10〜11）",
        "42-45分　安全装置（スライド12）",
        "45-55分　ハンズオン【デモC】＝参加者が④を操作",
        "55-60分　料金（スライド13）→ クロージング（スライド14）→ トライアル開始日を決める",
    ]:
        para(doc, "・" + t, size=10.5, space_after=2, indent=0.2)
    para(doc, "押したときに削る順番：デモB → スライド10〜11 → ハンズオンを10分に短縮。デモAとクロージングは絶対に削らない。",
         size=10, bold=True, color=RED, space_after=6)

    # ---- スライド別台本 ----
    h1(doc, "2. スライド別の進行台本")
    para(doc, "左の画像＝いま映っているスライド。【話す】をそのまま読んでも進行できます。", size=10)

    for i, png in enumerate(pngs):
        script = SCRIPT[i] if i < len(SCRIPT) else (["（台本未設定）"], None, None)
        talk, ops, escape = script

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(f"◆ スライド {i+1} / {len(pngs)}")
        set_jp(r, size=12, bold=True, color=NAVY)

        doc.add_picture(png, width=Cm(11.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

        for line in talk:
            tag_line(doc, "【話す】", line, GREEN)
        if ops:
            tag_line(doc, "【操作】", ops, NAVY)
        if escape:
            tag_line(doc, "【逃げ道】", escape, RED)

    # ---- デモ手順カード ----
    doc.add_page_break()
    h1(doc, "3. デモ手順カード（Excelに切り替えたらこのページ）")
    for title, steps, expect, steps2 in DEMO_CARDS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run("■ " + title)
        set_jp(r, size=11.5, bold=True, color=NAVY)
        for s in steps:
            para(doc, s, size=10.5, space_after=2, indent=0.4)
        tag_line(doc, "【期待値】", expect + "　←この数字を実行前に口頭で予告する", GREEN, size=10.5)
        for s in steps2:
            para(doc, s, size=10.5, space_after=2, indent=0.4)

    # ---- 想定問答 ----
    h1(doc, "4. 想定問答（聞かれたらこのまま答える）")
    for q, a in QA:
        tag_line(doc, "Q.", q, NAVY, size=10.5)
        tag_line(doc, "A.", a, GREEN, size=10.5)

    # ---- トラブル対応 ----
    h1(doc, "5. 当日のトラブルと逃げ道")
    for t in [
        "プロジェクタに映らない → 2台目のPCで映す → それも駄目なら紙焼きスライドで紙芝居＋Excelデモは画面を囲んで見てもらう",
        "Excelのマクロがブロックされる → 自分のPCなら前日リハで潰してあるはず。学校PCでのデモは求められても行わない（トライアル時に正式手順で）",
        "デモでエラーが出た → 慌てず画面をそのまま見せる：「書き込む前に止まる設計です」。エラー文言を控えて持ち帰る",
        "答えられない質問 → 「持ち帰って書面で回答します」と言い、その場で指南書の余白にメモ。当日中にメールで一次回答",
        "時間が半分しかない → デモA＋安全装置＋クロージングの3点のみ（20分版）。スライド3→6→7→12→14の順",
    ]:
        para(doc, "・" + t, size=10.5, space_after=3, indent=0.2)

    para(doc, "\n最後に必ず：トライアル開始日を決めてから帰る。「ご検討ください」で終わらせない。",
         size=11.5, bold=True, color=RED)

    # ---- 保存 → PDF ----
    docx_path = os.path.join(workdir, "playbook.docx")
    doc.save(docx_path)
    env = {**os.environ, "HOME": workdir}
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", workdir, docx_path], check=True, env=env)
    os.replace(os.path.join(workdir, "playbook.pdf"), os.path.abspath(OUT_PDF))
    print("pdf saved:", os.path.abspath(OUT_PDF))
    shutil.rmtree(workdir, ignore_errors=True)


if __name__ == "__main__":
    build()
