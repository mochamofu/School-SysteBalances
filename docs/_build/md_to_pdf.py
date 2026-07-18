# -*- coding: utf-8 -*-
"""内部用Markdown文書をPDF化する簡易コンバータ。

対応: 見出し(#/##/###)・箇条書き(-)・番号リスト(1.)・表(|)・
コードブロック(```)・太字(**)・区切り線(---)。
用途: docs配下の内部メモをPDFで確認・印刷したいとき。

使い方: python md_to_pdf.py <input.md> <output.pdf>
"""
import os
import re
import subprocess
import sys
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GRAY = (0x66, 0x66, 0x66)


def set_jp(run, size=10.5, bold=False, color=None, mono=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_runs(p, text, size=10.5, color=None, base_bold=False):
    """**bold** と `code` を解釈してrunを追加"""
    text = text.replace("`", "")
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = base_bold
        if part.startswith("**") and part.endswith("**"):
            bold = True
            part = part[2:-2]
        r = p.add_run(part)
        set_jp(r, size=size, bold=bold, color=color)


def heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    size = {1: 16, 2: 13.5, 3: 11.5}.get(level, 11)
    add_runs(p, text, size=size, color=NAVY, base_bold=True)
    if level == 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '2')
        bottom.set(qn('w:color'), '203864')
        pBdr.append(bottom)
        pPr.append(pBdr)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def flush_table(doc, rows):
    if not rows:
        return
    headers = rows[0]
    body = [r for r in rows[1:] if not all(re.match(r"^[-: ]*$", c) for c in r)]
    t = doc.add_table(rows=1 + len(body), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, h, size=9.5, color=(0xFF, 0xFF, 0xFF), base_bold=True)
        shade(cell, "203864")
    for i, row in enumerate(body):
        for j in range(len(headers)):
            v = row[j] if j < len(row) else ""
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            add_runs(cell.paragraphs[0], v, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def convert(md_path, pdf_path):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.7))

    in_code = False
    table_rows = []
    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(line)
            set_jp(r, size=9.5, color=GRAY)
            continue

        if stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            continue
        else:
            flush_table(doc, table_rows)
            table_rows = []

        if not stripped:
            continue
        if stripped == "---":
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            heading(doc, m.group(2), len(m.group(1)))
            continue
        m = re.match(r"^([-・])\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            add_runs(p, "・" + m.group(2), size=10)
            continue
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.4)
            add_runs(p, f"{m.group(1)}. " + m.group(2), size=10)
            continue
        # 通常段落（インデント継続行も含む）
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if line.startswith(("   ", "　")):
            p.paragraph_format.left_indent = Cm(0.8)
            add_runs(p, stripped, size=10)
        else:
            add_runs(p, stripped, size=10.5)
    flush_table(doc, table_rows)

    workdir = tempfile.mkdtemp(prefix="md2pdf_")
    docx_path = os.path.join(workdir, "doc.docx")
    doc.save(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", workdir, docx_path],
                   check=True, env={**os.environ, "HOME": workdir})
    os.replace(os.path.join(workdir, "doc.pdf"), os.path.abspath(pdf_path))
    print("pdf saved:", os.path.abspath(pdf_path))


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
