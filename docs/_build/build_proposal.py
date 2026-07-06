# -*- coding: utf-8 -*-
"""業務改善提案書（PDF元のdocx）を生成するスクリプト

docs/build_manual.py と同じヘルパーを流用する。
生成: docs/assistant_manual.docx → LibreOfficeでPDF化して
      docs/入力アシスタント_手順書.pdf として保存する。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"


def set_japanese_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text, size=20, color=(0x20, 0x38, 0x64)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_japanese_font(r, size=size, bold=True, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_japanese_font(r, size=15, bold=True, color=(0x20, 0x38, 0x64))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("■ " + text)
    set_japanese_font(r, size=12.5, bold=True, color=(0x44, 0x72, 0xC4))
    return p


def add_body(doc, text, size=10.5, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_japanese_font(r, size=size, bold=bold, color=color)
    return p


def add_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"Step {num}　")
    set_japanese_font(r, size=10.5, bold=True, color=(0x20, 0x38, 0x64))
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p


def add_note(doc, text, label="⚠ 注意"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FFF2CC')
    pPr.append(shd)
    r = p.add_run(f"{label}: ")
    set_japanese_font(r, size=10, bold=True, color=(0xC0, 0x50, 0x00))
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r2 = p.add_run(line)
        set_japanese_font(r2, size=10)
    return p


def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("☐ ")
    set_japanese_font(r, size=10.5, bold=True)
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_japanese_font(r, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '203864')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_japanese_font(r, size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table



doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ============ 表紙 ============
for _ in range(6):
    doc.add_paragraph()
add_title(doc, "積立金会計業務 改善提案書", size=24)
add_title(doc, "～「令和○年度生積立金」への転記業務の自動化～", size=14, color=(0x44, 0x72, 0xC4))
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("現場ヒアリング（録音）および実務ファイル5点の構造解析に基づく提案")
set_japanese_font(r, size=11, color=(0x60, 0x60, 0x60))
doc.add_page_break()

# ============ 1. 現状と課題 ============
add_h1(doc, "1．現状の業務構造と課題")
add_body(doc, "貴校の積立金会計は、すべての情報が「令和○年度生積立金」（学校徴収金個人別管理票）へ"
              "人の手で転記される構造になっています。")
add_body(doc, "掲示用名簿（クラス発表）─┐\n"
              "銀行の振替結果 ──────┼─▶ 令和○年度生積立金（データシート） ─▶ 精算書・決算書・督促状\n"
              "業者の請求書 ────────┘　　　　│\n"
              "　　　　　　　　　　　　支出承認書・収入承認書（同じ情報をもう一度記入）", size=9.5)
add_h2(doc, "ヒアリングで挙がった課題")
add_table(doc,
          ["課題（録音より）", "実ファイルで確認した実態"],
          [["クラス替えのたびに全員分の組・番号を手修正", "生徒321行×毎年。掲示用名簿からの目視転記"],
           ["支出のたびに全員分の金額を手入力", "支出100枠。1件あたり最大315名分の同一金額入力＋転退学・給付型・返金の例外を個別修正"],
           ["口座振替の結果を1人ずつ照合・入力", "収入43枠。毎月の振替で未納者を目視抽出"],
           ["承認書に同じ情報を二重記入", "支出承認書・収入承認書は独立ファイルで手書き運用（#REF!切れの旧リンクも放置）"],
           ["決算書は電卓集計で数字を直打ち", "決算書の総額・人数欄がハードコードされていた"],
           ["マクロが古く・保護されていて触れない", "内蔵マクロは1999年作の一括印刷1本のみ。VBAプロジェクトはパスワード保護"]],
          col_widths=[7.2, 9.2])

# ============ 2. 提案する仕組み ============
add_h1(doc, "2．渡すべき仕組み：3層構造の「転記レス化」")
add_body(doc, "マスターファイルは今のまま一切変更せず、入口（取込）→中継（一括入力）→出口（帳票）の"
              "3層でマスターを包み込みます。")

add_h2(doc, "第1層：入口ツール（新規開発を提案）")
add_table(doc,
          ["ツール", "置き換える手作業", "頻度"],
          [["口座マスター（生徒⇔口座の紐付け台帳）", "生徒と口座の対応を記憶と紙で管理", "年1回更新"],
           ["振替結果取込シート", "銀行の結果票を目視→未納者を手で抽出→入力。CSVを貼れば未納者表まで自動生成", "毎月"],
           ["年間徴収・支出予定表", "支出のたびに請求書から件名・単価を転記。年度初めの計画入力1回で、実行はボタン1つ。決算書の予算列も自動化", "年1回＋随時"]],
          col_widths=[5.5, 8.2, 2.7])

add_h2(doc, "第2層：中継ツール（納品済み・積立金入力アシスタント）")
add_body(doc, "名簿取込（自動照合）／支出・収入の一括入力／承認書の自動作成／決算集計／整合性チェック／"
              "精算書のPDF・印刷（計10機能）。ヒアリングの主訴であった「全員分の手入力」「番号紐付け」"
              "「承認書の二重記入」はこの層で解決済みです。全操作に自動バックアップ・構造チェック付き。")

add_h2(doc, "第3層：出口ツール（増築を提案）")
add_table(doc,
          ["ツール", "効果"],
          [["転退学処理ウィザード", "精算番号1つで、残金計算→返還金記録→精算書PDF→在籍終了まで一気通貫。返還金の計算ミス（保護者トラブルの火種）を排除"],
           ["督促状一括生成＋分納台帳", "未納者一覧から督促状PDFを一括作成。分納の約束も台帳で管理"],
           ["年度繰越ウィザード", "前年度残金→新年度マスターの繰越金へ自動転記。3月の年度替わり作業をボタン化"],
           ["月次残高突合シート", "通帳残高を1つ入れると帳簿残高との差額を即表示。監査・引き継ぎ対策"]],
          col_widths=[5.0, 11.4])

# ============ 3. 業務改善策 ============
add_h1(doc, "3．あわせて実施すべき業務改善策（運用ルール）")
add_step(doc, 1, "「マスターに直接触らない」原則 ― 入力はすべてアシスタント経由。数式破壊事故と入力ズレを根絶")
add_step(doc, 2, "承認書は「書くもの」から「出るもの」へ ― 決裁作業を内容確認と押印だけに")
add_step(doc, 3, "月次締めルーチンの定例化 ― 整合性チェック→通帳残高突合→保管用PCへの格納を毎月1セットで実施")
add_step(doc, 4, "保管の標準化 ― 精算書PDF（組-番号_氏名）＋マスター＋チェック結果を月次フォルダでワンセット保管")
add_step(doc, 5, "引き継ぎの型化 ― 練習用ダミーファイルで習熟→本番、という人事異動時の標準手順を確立")

# ============ 4. ロードマップ ============
add_h1(doc, "4．導入ロードマップ")
add_table(doc,
          ["フェーズ", "内容", "削減見込み"],
          [["Phase 1（済）", "入力アシスタント導入：一括入力・名簿照合・承認書自動化・点検・帳票出力", "月5〜10時間"],
           ["Phase 2", "振替結果取込＋口座マスター＋年間予定表：毎月の照合作業と転記の残りを自動化", "月2〜4時間＋ミス削減"],
           ["Phase 3", "転退学ウィザード・督促状一括・年度繰越・残高突合：例外業務と年度替わりの自動化", "繁忙期の残業削減"]],
          col_widths=[3.2, 9.6, 3.6])
add_note(doc, "Phase 2の「振替結果取込」を完全自動にするため、銀行（ゆうちょ）の振替結果ファイルの"
              "サンプル（口座番号・氏名はダミーで可）のご提供をお願いします。", label="ご協力のお願い")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("以上")
set_japanese_font(r, size=11, bold=True)

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'proposal.docx')
doc.save(out_path)
print(f"saved {out_path}")
