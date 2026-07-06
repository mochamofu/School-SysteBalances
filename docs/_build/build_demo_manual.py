# -*- coding: utf-8 -*-
"""デモ実施手順書（PDF元のdocx）を生成するスクリプト

docs/build_manual.py と同じヘルパーを流用する。
生成: docs/assistant_manual.docx → LibreOfficeでPDF化して
      docs/入力アシスタント_手順書.pdf として保存する。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"


def set_japanese_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text, size=20, color=(0x20, 0x38, 0x64)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_japanese_font(r, size=size, bold=True, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_japanese_font(r, size=15, bold=True, color=(0x20, 0x38, 0x64))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("■ " + text)
    set_japanese_font(r, size=12.5, bold=True, color=(0x44, 0x72, 0xC4))
    return p


def add_body(doc, text, size=10.5, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_japanese_font(r, size=size, bold=bold, color=color)
    return p


def add_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"Step {num}　")
    set_japanese_font(r, size=10.5, bold=True, color=(0x20, 0x38, 0x64))
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p


def add_note(doc, text, label="⚠ 注意"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FFF2CC')
    pPr.append(shd)
    r = p.add_run(f"{label}: ")
    set_japanese_font(r, size=10, bold=True, color=(0xC0, 0x50, 0x00))
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r2 = p.add_run(line)
        set_japanese_font(r2, size=10)
    return p


def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("☐ ")
    set_japanese_font(r, size=10.5, bold=True)
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_japanese_font(r, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '203864')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_japanese_font(r, size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table



doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ============ 表紙 ============
for _ in range(6):
    doc.add_paragraph()
add_title(doc, "積立金会計 入力アシスタント", size=24)
add_title(doc, "デモ実施手順書（営業同行用）", size=18, color=(0x44, 0x72, 0xC4))
doc.add_paragraph()
add_title(doc, "～ USB1本とPC2台で、学校の業務フローを目の前に再現する ～",
          size=11, color=(0x60, 0x60, 0x60))
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("所要時間: 約25分\n必要なもの: ノートPC2台（Microsoft 365/Excel 2019以降）・USBメモリ・デモキット")
set_japanese_font(r, size=11, color=(0x20, 0x38, 0x64))
doc.add_page_break()

# ============ 0. 考え方 ============
add_h1(doc, "0．デモの考え方")
add_body(doc, "お客様（学校）の実際の業務フローは「担当者のPCで入力 → 保管用PCに格納して完了」。"
              "デモもこの形をそのまま2台のPCでなぞります。機能を並べて見せるのではなく、"
              "『いつもの仕事が、そのまま速くなる』ことを目の前で再現するのが狙いです。")
add_table(doc,
          ["役割", "使い方"],
          [["PC1（事務担当のPC役）", "アシスタント＋練習用マスターでライブデモを行う主役機"],
           ["PC2（保管用PC役）", "前半はLP・アニメーションを映すプレゼン機。最後にUSB経由で成果物を受け取る"],
           ["USBメモリ", "①デモ環境の運搬 ②デモ中に「PC1→PC2へ成果物を渡す」小道具"]],
          col_widths=[5.5, 10.9])
add_note(doc, "このキットに本物の生徒データは絶対に入れないこと。練習用（架空の氏名）のみを使う。"
              "『実データの投入は契約後、貴校のPCの中だけで行います』と言い切れること自体が営業トークになる。", label="鉄則")

# ============ 1. 事前準備 ============
add_h1(doc, "1．事前準備（自宅・前日まで）")
add_h2(doc, "1-1．デモキットを作る")
add_step(doc, 1, "リポジトリで python assistant/src/build_assistant.py と make_practice_files.py を実行")
add_step(doc, 2, "生成された「積立金入力アシスタント.xlsx」を自分のPCで開き、VBA（A00〜A07）をインポート → .xlsm保存 → 「初期設定」まで済ませる（お客様の前でセットアップは見せない）")
add_step(doc, 3, "python assistant/src/make_usb_kit.py を実行 → assistant/output/デモキット/ が完成")
add_step(doc, 4, "デモキット内の .xlsx を手順2で作った .xlsm に差し替え、フォルダごとUSBへコピー")
add_h2(doc, "1-2．前日リハーサル（PC1で必ず1回）")
for item in [
    "USB→デスクトップへコピー → xlsmのプロパティで「許可する（ブロックの解除）」",
    "設定シートC3のマスターパスをこのPCの実パスに修正",
    "④支出一括入力 → マスターに金額が入る → バックアップフォルダが自動生成される、まで通す",
    "⑨精算書PDFが出力されることを確認",
]:
    add_check(doc, item)
add_note(doc, "最近のExcelはUSB/ネット由来のマクロを既定でブロックする。「許可する（ブロックの解除）」を忘れると当日デモが止まる。これが最大の落とし穴。")

# ============ 2. 当日進行 ============
add_h1(doc, "2．当日の進行（約25分）")
add_h2(doc, "①課題の共有（5分・PC2）")
add_body(doc, "製品LP（02_プレゼン資料の「製品LP.html」）のBefore/Afterを映す。"
              "『クラス替えで320人分を手修正』『支出のたび315人分入力』を読み上げ、"
              "相手に「うちのことだ」と言わせるパート。")
add_h2(doc, "②ライブデモ（10分・PC1）")
add_step(doc, 1, "「支出入力」シート: 件名・単価520円を入力、例外表に「45番=0（転退学）」「12番=-1040（返金）」だけ書く")
add_step(doc, 2, "④のボタン → 「対象80名・合計〇〇円を書き込みました」のダイアログを見せる")
add_step(doc, 3, "マスターを開き、全員に金額・支出合計・残金の自動計算・精算書への項目出現を見せる")
add_step(doc, 4, "続けて ⑤収入一括入力（未納者2名だけ指定→未納の印が自動で立つ）")
add_step(doc, 5, "⑧整合性チェック → 「問題なし」と言い切る画面を見せる")
add_step(doc, 6, "⑨精算書PDF一括生成（組-番号_氏名.pdfがずらっと並ぶ）")
add_h2(doc, "③保管用PCへの格納を再現（3分・USB＋PC2）")
add_body(doc, "PC1で更新されたマスターと精算書PDFフォルダをUSBの「03_保管用」へコピー → "
              "PC2に貼り付けて開く。『いつもの保管業務がそのまま、中身だけが一瞬で出来上がる』と締める。")
add_h2(doc, "④安全装置の実演（3分・PC1）")
add_body(doc, "バックアップフォルダに日時付きコピーが増えているのを見せ、その場で1つ前に戻す。"
              "学校相手では「速い」より「壊れない・戻せる」が響く。")
add_h2(doc, "⑤クロージング（4分・PC2）")
add_body(doc, "LPの料金セクション → 30日間無料トライアルの案内 → 導入相談フォームのQRコードを紙で渡す。"
              "その場でスマホから申し込んでもらえれば商談化率が上がる。")

# ============ 3. トーク集 ============
add_h1(doc, "3．手離れをよくする一言集")
add_table(doc,
          ["場面", "トーク"],
          [["構造チェックの説明", "いまお使いの積立金ファイルは1ミリも変えません。隣に置くだけです"],
           ["バックアップの説明", "操作を間違えても、実行前のコピーが毎回自動で残ります"],
           ["稟議の懸念潰し", "インターネット不要・データは校内から出ません"],
           ["価格の説明", "事務1時間分の時給で、1ヶ月分の入力が終わります"],
           ["トライアル誘導", "まずは架空データの練習用ファイルで、30日間ノーリスクで試せます"]],
          col_widths=[4.5, 11.9])

# ============ 4. 当日チェックリスト ============
add_h1(doc, "4．当日チェックリスト")
for item in [
    "USB内: 01_デモ環境（xlsm差し替え済み）／02_プレゼン資料／03_保管用（空）",
    "PC1・PC2ともに電源アダプタ持参、Excelのバージョン確認済み",
    "PC1でブロック解除 → 設定C3修正 → ④⑧⑨の3ボタンが動くことを開始前に確認",
    "導入相談フォームのQRコード印刷物",
    "名刺・会社案内・料金表（LPの料金セクション印刷でも可）",
]:
    add_check(doc, item)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("以上")
set_japanese_font(r, size=11, bold=True)

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'demo_manual.docx')
doc.save(out_path)
print(f"saved {out_path}")
