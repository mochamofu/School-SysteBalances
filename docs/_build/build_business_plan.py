# -*- coding: utf-8 -*-
"""事業計画書（金融機関提出用）を生成する。

日本政策金融公庫・信用金庫等に提出することを想定した正式な事業計画書。
販売戦略メモ・創業計画書下書き・創業ロードマップの数値と整合させてある。
代表者名・会社名・日付などは記入欄（＿＿）としてあり、提出前に記入する。

生成: docs/06_法人化・創業融資/事業計画書.pdf
"""
import os
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"
NAVY = (0x20, 0x38, 0x64)
GRAY = (0x66, 0x66, 0x66)
RED = (0xC0, 0x00, 0x00)


def set_jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=10.5, bold=False, color=None, space_after=5, center=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_jp(r, size=size, bold=bold, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=14, bold=True, color=NAVY)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_jp(r, size=11.5, bold=True, color=NAVY)
    return p


def bullets(doc, items, size=10.5):
    for t in items:
        para(doc, "・" + t, size=size, space_after=3, indent=0.4)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def table(doc, headers, rows, widths, size=9.5, bold_last_row=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_jp(r, size=size, bold=True, color=(0xFF, 0xFF, 0xFF))
        shade(cell, "203864")
    for i, row in enumerate(rows):
        is_last = bold_last_row and i == len(rows) - 1
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(v))
            set_jp(r, size=size, bold=is_last)
            if is_last:
                shade(cell, "EAF0FA")
    for j, w in enumerate(widths):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def build(docx_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(2.0))

    # ================= 表紙 =================
    for _ in range(5):
        doc.add_paragraph()
    para(doc, "事　業　計　画　書", size=28, bold=True, color=NAVY, center=True, space_after=16)
    para(doc, "学校徴収金（積立金）会計 入力自動化サービス", size=15, bold=True, center=True, space_after=4)
    para(doc, "「積立金会計 入力アシスタント」", size=15, bold=True, center=True, space_after=40)
    para(doc, "作成日：＿＿＿＿年＿＿月＿＿日", size=11, center=True, space_after=6)
    para(doc, "商号（予定含む）：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", size=11, center=True, space_after=6)
    para(doc, "代表者：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", size=11, center=True, space_after=6)
    para(doc, "所在地：＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", size=11, center=True, space_after=40)
    para(doc, "本計画書の数値は作成時点の見込みであり、市場環境等により変動する可能性があります。",
         size=9, color=GRAY, center=True)

    doc.add_page_break()

    # ================= 1. 事業の要旨 =================
    h1(doc, "1. 事業の要旨（エグゼクティブサマリー）")
    para(doc, "本事業は、高等学校の学校徴収金（積立金）会計における入力・照合・帳票作成業務を"
              "自動化するソフトウェアサービスである。学校が長年使用してきた既存のExcel管理ファイルを"
              "一切変更せず、隣に置いた1つのファイルから安全に入力を代行する点が最大の特長であり、"
              "システム移行を伴わないため、予算とIT環境に制約の大きい公立高校でも導入できる。", space_after=6)
    bullets(doc, [
        "製品は開発完了済み（バージョン1.0）。実物と同形式のファイルで全機能の検証を実施済み",
        "先行導入校（有償）との契約・導入前後の作業時間の実測により、効果を定量的に示せる体制",
        "販売単価は公立年額29,800〜49,800円・私立月額19,800円・教育委員会一括は1校年額30,000円",
        "仕入れ・在庫なし。1校あたりの変動費は約1,500円（USB・送料）であり、売上のほぼ全額が粗利",
        "3年目に年商約384万円、以後は教育委員会一括契約により拡大を計画",
        "必要資金は400万円（自己資金100万円・借入希望300万円）。全額運転資金中心",
    ])

    # ================= 2. 創業の動機と代表者 =================
    h1(doc, "2. 創業の動機・代表者の背景")
    para(doc, "高等学校の積立金会計は、保護者からの預かり金でありながら、現在も数百人分の手入力と"
              "目視照合に支えられている。担当職員は毎月、銀行の振替結果と生徒名簿を突き合わせ、"
              "同じ金額を数百回入力し、承認書類に同じ内容を再度記入している。作業時間の負担に加え、"
              "誤請求が1件でも発生すれば保護者への説明対応に発展する、心理的負荷の大きい業務である。", space_after=4)
    para(doc, "代表者はこの実態に現場で接し、既存の管理ファイルを変えずに入力だけを自動化する本製品を"
              "開発した。実際の学校ファイルによる検証と先行校への導入を経て、事業として全国の学校へ"
              "提供するため法人を設立する。", space_after=6)
    para(doc, "【代表者略歴】（提出時に記入）", size=10, color=GRAY, space_after=2)
    for _ in range(3):
        para(doc, "＿＿年＿月〜　＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿＿", space_after=4, indent=0.4)

    # ================= 3. 製品・サービス =================
    h1(doc, "3. 製品・サービスの内容")
    h2(doc, "3-1. 製品概要")
    para(doc, "「積立金会計 入力アシスタント」は、Microsoft Excelのマクロ機能（VBA）のみで動作する"
              "入力自動化パッケージである。専用ソフトのインストール・サーバー・ネットワーク接続は"
              "一切不要で、生徒情報は校内のパソコンから外に出ない。", space_after=4)
    table(doc,
          ["業務", "従来", "本製品導入後"],
          [
              ["クラス替え（年1回）", "320人分を1行ずつ手修正（数日）", "名簿を貼り付けてボタン2回"],
              ["口座振替の照合（毎月）", "振替不能者を目視で探し全員分を手入力", "結果を貼るだけで未納者を自動判定・一括記録"],
              ["支出の記録（随時）", "同じ金額を全員分入力し例外を個別修正", "例外の生徒だけ指定して一括入力"],
              ["承認書類の作成", "同じ内容を書類へ二度書き", "一括入力と同時に自動作成"],
              ["決算集計（年度末）", "電卓で項目別に集計", "ボタン1回で全項目を集計"],
              ["年度更新", "翌年度ファイルを一から作成", "残す項目に○を付けるだけで引き継ぎ"],
          ],
          widths=[3.6, 6.2, 6.2])
    h2(doc, "3-2. 安全設計（学校会計に求められる信頼性）")
    bullets(doc, [
        "書き込み前に必ず自動バックアップを作成（いつでも実行前に戻せる）",
        "既存ファイルの行・列・様式・数式は一切変更しない",
        "誤ったファイルへの書き込みを拒否する構造チェック、入力ミスはその場でエラー停止",
        "異常データ15項目の耐性テストを実施し、判明した問題はすべて修正済み（検証記録あり）",
    ])
    h2(doc, "3-3. 提供・サポート体制")
    bullets(doc, [
        "USB郵送＋オンライン指導（Zoom）により、訪問なしで全国に導入可能（導入キャパ月10校）",
        "図解マニュアル・完全図解の導入ガイド・動作確認チェックシートを整備済み",
        "年間保守として、障害対応・年度更新支援・担当者交代時の再指導を提供",
    ])

    # ================= 4. 市場環境 =================
    h1(doc, "4. 市場環境と事業機会")
    para(doc, "全国の高等学校は約4,800校（公立約3,500校・私立約1,300校、文部科学省「学校基本調査」に"
              "基づく概数）。積立金・学校徴収金の管理事務はほぼ全校に存在する。", space_after=4)
    bullets(doc, [
        "既存の校務支援システムは月額2〜3万円／校または生徒1人あたり月約300円が相場であり、"
        "積立金会計単体のために導入するには過大。既存Excelからの移行負担も大きい",
        "一方、現場はExcel管理を継続しており、入力業務の負担は解消されていない",
        "本製品は「既存ファイルを変えない・少額・インストール不要」により、"
        "既存システムが取り込めていない層（特に公立）に到達できる",
        "市場規模の目安: 私立1,300校×年額19.8万円＋公立3,500校×年額3〜5万円 ≒ 年間約4億円。"
        "3年目計画（年商384万円）はこの1%未満であり、達成可能性の高い水準に設定",
    ])

    # ================= 5. 競合と優位性 =================
    h1(doc, "5. 競合と本製品の優位性")
    table(doc,
          ["比較項目", "校務支援システム", "市販テンプレート", "本製品"],
          [
              ["既存ファイルの継続利用", "不可（移行が前提）", "様式が合わない", "可（一切変更しない）"],
              ["価格（年間）", "24〜36万円／校〜", "数千円（買切）", "3〜20万円／校"],
              ["導入作業", "移行・研修が必要", "自力で作り直し", "コピーと設定のみ（約15分）"],
              ["保守・年度更新", "ベンダー保守", "なし", "年間保守に含む"],
              ["公立校の予算適合", "困難", "－", "少額執行の範囲で導入可"],
          ],
          widths=[4.0, 4.0, 3.6, 4.4])
    para(doc, "参入障壁について: 本製品の価値は技術そのものではなく、実際の学校ファイルの様式・"
              "業務フローへの適合と、導入実績・実測データの蓄積にある。先行して実績を積むことが"
              "そのまま防壁となる。", size=10, space_after=4)

    # ================= 6. 販売戦略 =================
    h1(doc, "6. 販売戦略")
    h2(doc, "6-1. 価格体系（税別）")
    table(doc,
          ["区分", "価格", "内容"],
          [
              ["導入協力校（先行3校限定）", "年額 29,800円", "事例・実測データ提供に協力いただく特別価格"],
              ["公立高校（標準）", "年額 49,800円", "少額執行で稟議しやすい価格設定"],
              ["私立高校", "月額 19,800円（年払 198,000円）", "導入支援・優先サポート込み"],
              ["教育委員会 一括", "1校あたり年額 30,000円（10校以上）", "集合研修・窓口一元化で提供原価を圧縮"],
          ],
          widths=[5.2, 5.0, 5.8])
    h2(doc, "6-2. 販売チャネルと展開手順")
    bullets(doc, [
        "第1段階: 先行導入校（都内公立）で有償実績と削減効果の実測データを取得",
        "第2段階: 学校事務職員の研究会・異動ネットワークを通じた公立の横展開（広告に依存しない）、"
        "および私立高校への直販（Webサイト・私学事務長ネットワーク）",
        "第3段階: 同一区市で5校以上の実績を基に、教育委員会への一括ライセンス提案",
        "導入はUSB郵送＋オンライン指導で完結するため、1人体制でも月10校の導入が可能",
    ])

    # ================= 7. 収支計画 =================
    h1(doc, "7. 収支計画（3ヶ年）")
    h2(doc, "7-1. 売上計画")
    table(doc,
          ["区分（単価）", "1期目", "2期目", "3期目"],
          [
              ["導入協力校（年29,800円）", "3校　89千円", "3校　89千円", "3校　89千円"],
              ["公立標準（年49,800円）", "―", "7校　349千円", "12校　598千円"],
              ["私立（月19,800円）", "2校※　238千円", "6校　1,426千円", "12校　2,851千円"],
              ["教育委員会（校3万円）", "―", "―", "10校　300千円"],
              ["売上高合計", "327千円", "1,864千円", "3,838千円"],
          ],
          widths=[5.4, 3.4, 3.6, 3.6], bold_last_row=True)
    para(doc, "※1期目の私立は年度途中加入（平均6ヶ月）として算定。", size=9, color=GRAY, space_after=4)
    h2(doc, "7-2. 損益計画（概算・千円）")
    table(doc,
          ["科目", "1期目", "2期目", "3期目"],
          [
              ["売上高", "327", "1,864", "3,838"],
              ["売上原価（USB・送料等）", "15", "30", "50"],
              ["役員報酬", "1,200", "1,800", "2,400"],
              ["広告宣伝・営業費", "300", "400", "500"],
              ["外注費（税理士・法務等）", "300", "360", "400"],
              ["通信・消耗品・その他", "250", "300", "350"],
              ["支払利息（年2.5%想定）", "70", "60", "50"],
              ["経常利益", "▲1,808", "▲1,086", "88"],
          ],
          widths=[5.4, 3.4, 3.6, 3.6], bold_last_row=True)
    bullets(doc, [
        "1〜2期目の赤字は役員報酬を含む先行投資期間であり、借入金と自己資金で賄う計画",
        "変動費率が約2%と極めて低いため、契約校数の増加がほぼそのまま利益改善に直結する",
        "単月黒字化は3期目上期を計画（累計契約27校時点）。教育委員会一括契約が取れた場合は前倒し",
    ], size=10)

    # ================= 8. 資金計画 =================
    h1(doc, "8. 資金計画")
    h2(doc, "8-1. 必要な資金と調達方法")
    table(doc,
          ["必要な資金", "金額", "調達方法", "金額"],
          [
              ["設備資金（業務用PC等）", "20万円", "自己資金", "100万円"],
              ["運転資金（営業・広告・外注）", "180万円", "日本政策金融公庫 借入", "300万円"],
              ["運転資金（役員報酬・予備費）", "200万円", "", ""],
              ["合計", "400万円", "合計", "400万円"],
          ],
          widths=[5.6, 2.6, 5.6, 2.6], bold_last_row=True)
    h2(doc, "8-2. 借入の返済計画")
    bullets(doc, [
        "借入希望額300万円・返済期間7年・元金据置6ヶ月を希望",
        "毎月返済額（据置後）: 元金約38千円＋利息（年2.5%想定で当初約6千円）＝ 約44千円",
        "3期目の計画月商 約320千円に対し返済は約44千円（返済負担率 約14%）であり、無理のない水準",
        "契約は年額前受が中心のため、期初に現金が積み上がる資金繰り構造（貸倒れリスクも学校向けのため僅少）",
    ], size=10)

    # ================= 9. リスクと対応 =================
    h1(doc, "9. 想定されるリスクと対応策")
    table(doc,
          ["リスク", "対応策"],
          [
              ["学校の意思決定が遅い（年度予算サイクル）", "少額価格で稟議を軽くし、年度替わり（1〜3月）に営業を集中。30日無料トライアルで判断を後押し"],
              ["1人体制への依存", "導入をUSB郵送＋オンラインに標準化し月10校のキャパを確保。全手順を文書化済みで外部委託にも移行可能"],
              ["Excel環境の変化", "Office標準機能（VBA）のみで構成。バージョン追随は保守契約内で対応"],
              ["個人情報の懸念による導入見送り", "生徒情報に一切接触しない設計（オフライン・校内完結）を技術資料で提示"],
              ["類似品の出現", "実物様式への適合ノウハウ・導入実績・実測データの先行蓄積で優位を維持"],
              ["販売不振", "変動費率2%のため損失は限定的。役員報酬の圧縮と営業期間の延長で対応し、撤退時も負債は借入のみ"],
          ],
          widths=[6.0, 10.0], size=9.5)

    # ================= 10. 添付資料 =================
    h1(doc, "10. 添付資料一覧")
    bullets(doc, [
        "製品資料一式（サービス紹介・機能一覧・図解マニュアル）",
        "検証記録（実物同形式ファイルでの全機能検証・異常データ15項目の耐性テスト）",
        "導入協力校との覚書・請求書控え（有償実績）",
        "導入前後の作業時間の実測記録",
        "販売価格表・3年販売計画の詳細",
    ])
    para(doc, "", space_after=10)
    para(doc, "以上", size=11, center=True)

    doc.save(docx_path)
    print("docx saved:", docx_path)


def to_pdf(docx_path, pdf_out):
    outdir = os.path.dirname(docx_path)
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path],
                   check=True, env={**os.environ, "HOME": tempfile.gettempdir()})
    os.replace(os.path.splitext(docx_path)[0] + ".pdf", pdf_out)
    print("pdf saved:", pdf_out)


if __name__ == "__main__":
    scratch = os.environ.get("BUILD_TMP", tempfile.gettempdir())
    docx_path = os.path.join(scratch, "business_plan.docx")
    build(docx_path)
    here = os.path.dirname(os.path.abspath(__file__))
    to_pdf(docx_path, os.path.join(here, "..", "06_法人化・創業融資", "事業計画書.pdf"))
