# -*- coding: utf-8 -*-
"""学次会計 業務改善ツール ─ 動作確認・2台PCシミュレーション手順書 を生成"""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "IPAGothic"

def set_japanese_font(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_title(doc, text, size=20, color=(0x20, 0x38, 0x64)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_japanese_font(r, size=size, bold=True, color=color)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_japanese_font(r, size=15, bold=True, color=(0x20, 0x38, 0x64))
    # 下線的に細いボーダーを引く
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '203864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("■ " + text)
    set_japanese_font(r, size=12.5, bold=True, color=(0x44, 0x72, 0xC4))
    return p

def add_body(doc, text, size=10.5, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        set_japanese_font(r, size=size, bold=bold, color=color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_japanese_font(r, size=10.5)
    return p

def add_step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"Step {num}　")
    set_japanese_font(r, size=10.5, bold=True, color=(0x20, 0x38, 0x64))
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p

def add_note(doc, text, label="⚠ 注意"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'FFF2CC')
    pPr.append(shd)
    r = p.add_run(f"{label}: ")
    set_japanese_font(r, size=10, bold=True, color=(0xC0, 0x50, 0x00))
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        r2 = p.add_run(line)
        set_japanese_font(r2, size=10)
    return p

def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("☐ ")
    set_japanese_font(r, size=10.5, bold=True)
    r2 = p.add_run(text)
    set_japanese_font(r2, size=10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        set_japanese_font(r, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '203864')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_japanese_font(r, size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_pagebreak(doc):
    doc.add_page_break()


doc = Document()

# ページ設定
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(10.5)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), FONT_NAME)

# ============================================================
# 表紙
# ============================================================
for _ in range(4):
    doc.add_paragraph()
add_title(doc, "学次会計 業務改善ツール", size=26)
add_title(doc, "動作確認・2台PCシミュレーション手順書", size=16, color=(0x44, 0x72, 0xC4))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("～ 学校のパソコンに導入する前に、ご自身で動作を確認するための手引き ～")
set_japanese_font(r, size=11, color=(0x59, 0x59, 0x59))
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("対象ファイル：")
set_japanese_font(r, size=10.5, bold=True)
add_body(doc, "① 番号紐付けテンプレート.xlsx\n② CoPilotプロンプトガイド_学次会計用.xlsx\n"
              "③ master_マスターブック.xlsx\n④ 口座データバリデーション.xlsx", size=10.5)
for p_ in doc.paragraphs[-1:]:
    p_.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_pagebreak(doc)

# ============================================================
# 目次的な概要
# ============================================================
add_h1(doc, "0. この手順書について")
add_body(doc,
    "このツール群は、学校の事務PC1台で使うことを前提に作られていますが、"
    "実際の学次会計業務は「入力する人」と「確認・転記する人」が分かれている場面が多くあります。\n"
    "この手順書では、PCを2台用意して「担当者A（入力する人）」「担当者B（確認・提出する人）」"
    "の役割に分けてファイルをやり取りするシミュレーションを行い、"
    "実際の運用に近い形で正しく動くかどうかを確認します。"
)
add_note(doc,
    "このシミュレーションは「2台のPCでリアルタイムに同時編集する」ものではありません。\n"
    "Excelファイルはメール添付・共有フォルダ・USBメモリなどで「受け渡し」しながら使うため、"
    "1人が順番にPC1→PC2の役割を演じても、本当に2人・2台で行っても、どちらでも確認できます。",
    label="ポイント"
)

add_h1(doc, "1. 事前準備")
add_h2(doc, "1-1. 必要なもの")
add_bullet(doc, "Windows PC（できれば2台。1台でも代用可）　※Microsoft 365 の Excel が入っていること")
add_bullet(doc, "今回お渡しする output フォルダ内の4つの .xlsx ファイル")
add_bullet(doc, "vba フォルダ内の .bas マクロファイル（③ master_マスターブック.xlsx に追加する場合）")
add_note(doc,
    "本ツールは XLOOKUP 関数を使用しています。XLOOKUP は Microsoft 365 / Excel 2021 以降で利用できます。\n"
    "古いバージョンの Excel（2019以前）や、無料のExcel代替ソフト（LibreOffice Calc 24.8未満など）では\n"
    "正しく動作しない場合があります。学校のPCのExcelバージョンを事前に確認してください。"
)

add_h2(doc, "1-2. フォルダ構成の準備（2台のPCを模擬する）")
add_body(doc, "デスクトップなどに、以下の2つのフォルダを作成してください。")
add_step(doc, 1, "「PC1_担当者A」フォルダを作成し、output フォルダの4ファイルをコピーする")
add_step(doc, 2, "「PC2_担当者B」フォルダを作成する（最初は空でよい）")
add_body(doc,
    "実際の運用では、担当者Aが入力したファイルを担当者Bに渡して確認してもらう、という流れになります。\n"
    "シミュレーションでは「PC1_担当者A」で作業 → 完成したファイルを「PC2_担当者B」にコピー → "
    "そこで確認作業を行う、という形で再現します。"
)

add_pagebreak(doc)

# ============================================================
# ツール①
# ============================================================
add_h1(doc, "2. ツール①：番号紐付けテンプレート.xlsx の確認")
add_body(doc, "【想定シーン】クラス替えが決まり、新クラス名簿を生徒マスターに反映する作業")

add_h2(doc, "2-1. PC1（担当者A）での作業：新クラス名簿の入力")
add_step(doc, 1, "「PC1_担当者A」フォルダの 番号紐付けテンプレート.xlsx を開く")
add_step(doc, 2, "オレンジ色タブの「新年度名簿入力」シートを開く")
add_step(doc, 3, "A〜C列（氏名・新クラス・新出席番号）に、サンプルで入っている5名のデータがあることを確認する"
                 "（実際の運用では先生からもらった名簿をここに貼り付ける）")
add_step(doc, 4, "上書き保存する（Ctrl+S）")
add_step(doc, 5, "完成したファイルを「PC2_担当者B」フォルダにコピーする")

add_h2(doc, "2-2. PC2（担当者B）での作業：照合結果の確認")
add_step(doc, 1, "「PC2_担当者B」フォルダの 番号紐付けテンプレート.xlsx を開く")
add_step(doc, 2, "赤色タブの「照合結果」シートを開く")
add_step(doc, 3, "E列「照合状態」の表示を確認する")

add_body(doc, "【期待される結果（チェックリスト）】", bold=True)
add_check(doc, "山田 太郎・鈴木 花子・田中 一郎・佐藤 美咲・高橋 健一の5行とも「一致」（緑色）と表示される")
add_check(doc, "D列「生徒番号（自動）」に 04001〜04005 が自動表示されている")
add_check(doc, "F列「旧クラス（参考）」・G列「ふりがな（参考）」も自動で表示されている")

add_body(doc, "【応用：エラーケースも試してみる】", bold=True)
add_bullet(doc, "「新年度名簿入力」シートの氏名を「山田太郎」（スペースを削除）に変更すると → 「氏名不一致」（黄色）になることを確認")
add_bullet(doc, "氏名を「存在しない 生徒」に変更すると → 「未マッチ」（赤色）になることを確認")
add_note(doc,
    "「山田 太郎」と「山田太郎」（スペース有無）は別人として判定されます。これは仕様です。\n"
    "名簿の表記を生徒マスターと揃えることが、このツールを正しく使う一番のポイントです。"
)

add_pagebreak(doc)

# ============================================================
# ツール④
# ============================================================
add_h1(doc, "3. ツール④：口座データバリデーション.xlsx の確認")
add_body(doc, "【想定シーン】自動払込利用申込書（紙）の内容をExcelに転記し、銀行提出前にミスをチェックする作業")

add_h2(doc, "3-1. PC1（担当者A）での作業：申込書の転記")
add_step(doc, 1, "「PC1_担当者A」フォルダの 口座データバリデーション.xlsx を開く")
add_step(doc, 2, "緑色タブの「口座データ入力」シートを確認する。サンプルとして10名分のデータが入っている")
add_step(doc, 3, "上書き保存し、「PC2_担当者B」フォルダにコピーする")

add_h2(doc, "3-2. PC2（担当者B）での作業：バリデーション確認と銀行提出データの作成")
add_step(doc, 1, "「PC2_担当者B」フォルダの 口座データバリデーション.xlsx を開く")
add_step(doc, 2, "赤色タブの「バリデーション結果」シートを開く")
add_step(doc, 3, "K列「総合判定」を確認する")

add_body(doc, "【期待される結果（チェックリスト）】", bold=True)
add_check(doc, "佐藤 美咲（04004）の行：G列「番号チェック」が「NG: 7桁」になる（番号が8桁でないため）")
add_check(doc, "小林 あい（05003）の行：F列「記号チェック」が「NG: 4桁」になる（記号が5桁でないため）")
add_check(doc, "上記2名以外の8名は K列「総合判定」が「OK」（緑色）になる")
add_check(doc, "B列「氏名」に、口座データ入力シートで入力した氏名がそのまま表示される")
add_check(doc, "紺色タブ「銀行提出用」シートで、OKの8名のみ16桁コードが生成され、エラーの2名は「─」と表示される")

add_body(doc, "【応用：重複チェックを試す】", bold=True)
add_bullet(doc, "「口座データ入力」シートの空いている行に、既存の記号・番号と同じ値をコピーして追加する")
add_bullet(doc, "→「バリデーション結果」シートで該当する2行とも I列「重複チェック」が「重複あり」になることを確認")

add_note(doc,
    "H列「カナチェック」（全角/半角の判定）は、日本語版Excelの「DBCS（全角/半角を区別する）」機能を前提にしています。\n"
    "これはExcel独自の標準的な仕組みで、日本語Windows + 日本語Excelであれば正しく動作しますが、"
    "英語版OSや一部の互換ソフトでは正しく判定できない場合があります。\n"
    "学校のPCで実際に確認する際は、必ずこのH列が正しく「全角混在」を検出できているか目視確認してください。"
)

add_pagebreak(doc)

# ============================================================
# ツール③
# ============================================================
add_h1(doc, "4. ツール③：master_マスターブック.xlsx の確認")
add_body(doc, "【想定シーン】日々の月次徴収・支出記録の入力から、年度末の精算書・決算書作成までの一連の流れ")

add_h2(doc, "4-1. PC1（担当者A）での作業：日々の記録入力")
add_step(doc, 1, "「PC1_担当者A」フォルダの master_マスターブック.xlsx を開く")
add_step(doc, 2, "赤色タブの「支出記録」シートを開く。サンプルとして10件の支出記録が入っている")
add_step(doc, 3, "一番下の空行に、新しい支出を1行追加する（日付・支出番号・生徒番号・金額・備考）"
                 "※生徒番号を空欄にすると「全員分」として扱われる")
add_step(doc, 4, "上書き保存し、「PC2_担当者B」フォルダにコピーする")

add_h2(doc, "4-2. PC2（担当者B）での作業：精算書・決算書の確認")
add_step(doc, 1, "「PC2_担当者B」フォルダの master_マスターブック.xlsx を開く")
add_step(doc, 2, "紫色タブの「精算書(出力)」シートを開く")
add_step(doc, 3, "B3の黄色セルに生徒番号「04001」を入力する")

add_body(doc, "【期待される結果（チェックリスト）】", bold=True)
add_check(doc, "【支出の部】に「各教科フラットファイル 500円」「修学旅行積立① 10,000円」など、サンプル支出記録に応じた金額が自動表示される")
add_check(doc, "B3を「04002」に変更すると、「芸術鑑賞教室（1年）」が 0円 になる"
               "（サンプルでは欠席による返還[-1,696円]が記録されており、全員分[1,696円]と相殺されて0円になる仕組み）")
add_check(doc, "紺色タブ「決算書(出力)」シートで、各支出項目の「支出合計（自動）」が "
               "単価×（対象人数－免除等）に近い数字で集計されている")

add_body(doc, "【給付型奨学金の自動判定を試す】", bold=True)
add_bullet(doc, "「生徒マスター」シートで、04003（田中 一郎）のK列「給付型」セルに「○」を入力する")
add_bullet(doc, "「精算書(出力)」シートでB3に「04003」を入力し、保存・再度開く")
add_bullet(doc, "→ 支出項目マスターでG列「給付型対象」フラグが付いている項目だけ「給」と表示され、金額が0円扱いになることを確認")

add_note(doc,
    "「精算書(出力)」「決算書(出力)」「個人別管理表(出力)」の3シートは数式で自動生成されるシートです。"
    "ここを直接編集すると数式が壊れるため、入力は「生徒マスター」「月次徴収」「支出記録」の3シートのみで行ってください。"
)

add_pagebreak(doc)

# ============================================================
# ツール②
# ============================================================
add_h1(doc, "5. ツール②：CoPilotプロンプトガイド_学次会計用.xlsx の確認")
add_body(doc,
    "このファイルは数式を使わない「参照用の資料」です。動作確認というよりも、"
    "実際にCoPilotを使う担当者のPCで開けること、印刷して机に置けることを確認してください。"
)
add_step(doc, 1, "ファイルを開き、「プロンプトテンプレート」「Outlookルール設定」「運用カレンダー」の3シートが表示されることを確認")
add_step(doc, 2, "黄色いセルのプロンプト文をコピーし、実際にMicrosoft 365 CoPilotに貼り付けて動作するか試す"
                 "（CoPilot for Microsoft 365 のライセンスが必要です）")

add_pagebreak(doc)

# ============================================================
# VBAマクロ確認
# ============================================================
add_h1(doc, "6. VBAマクロの導入・動作確認（master_マスターブック.xlsx）")
add_h2(doc, "6-1. マクロの追加方法")
add_step(doc, 1, "master_マスターブック.xlsx を開く")
add_step(doc, 2, "Alt + F11 キーを押して VBA エディタを開く")
add_step(doc, 3, "メニューの「ファイル」→「ファイルのインポート」→ vba/MasterBook_Macros.bas を選択")
add_step(doc, 4, "「ファイル」→「閉じてMicrosoft Excelへ戻る」")
add_step(doc, 5, "「ファイル」→「名前を付けて保存」→ファイルの種類を「Excel マクロ有効ブック (*.xlsm)」に変更して保存")

add_h2(doc, "6-2. 主なマクロの動作確認")
add_table(doc,
    headers=["マクロ名", "確認方法", "期待される結果"],
    rows=[
        ["精算書_行非表示", "Alt+F8 → 実行", "精算書(出力)で金額0円の行が非表示になる"],
        ["精算書_行表示", "Alt+F8 → 実行", "非表示にした行が再表示される"],
        ["精算書_一括印刷", "Alt+F8 → 実行（確認ダイアログが出る）", "全生徒分の精算書が順番に印刷される"],
        ["精算書_一括PDF出力", "Alt+F8 → 実行", "ブックと同じ場所に「精算書PDF」フォルダが作成され、生徒ごとのPDFが保存される"],
        ["個人別管理表_クラス更新", "Alt+F8 → 実行", "クラス替え後の新クラス情報が反映される"],
    ],
    col_widths=[4.5, 6.5, 6.0]
)
add_note(doc,
    "マクロ実行時に「セキュリティの警告：マクロが無効化されました」と表示された場合は、\n"
    "「コンテンツの有効化」ボタンをクリックしてください。\n"
    "VBAコードには平易な日本語コメントを付けているため、Alt+F11でコードを開けば事務職員でも内容を確認できます。"
)

add_pagebreak(doc)

# ============================================================
# 検証結果まとめ
# ============================================================
add_h1(doc, "7. 開発時に実施した検証について（参考情報）")
add_body(doc,
    "このツールをお渡しする前に、実際にExcel互換ソフトでファイルを開き、数式を再計算させて"
    "結果が正しいかを検証しました。その過程で見つかった問題は、すべて修正済みです。"
)
add_table(doc,
    headers=["検証内容", "見つかった問題", "対応"],
    rows=[
        ["③ 支出記録の返還処理",
         "欠席等による返還（マイナス金額）が正しく相殺されず、満額のまま表示されていた",
         "全員分金額と個人別の調整額を加算する方式に修正済み"],
        ["③ 決算書の集計",
         "学校全体の支出合計が、生徒数を反映せずに記録の生値をそのまま合計していた",
         "個人別管理表の確定金額を合算する方式に修正済み"],
        ["④ バリデーション結果の氏名表示",
         "存在しない「生徒マスター」シートを参照し、常に「未マッチ」と表示されていた",
         "同一ブック内の「口座データ入力」シートを直接参照するよう修正済み"],
    ],
    col_widths=[4.5, 6.5, 6.0]
)
add_note(doc,
    "本検証はLinux環境のLibreOffice（XLOOKUPに未対応の古いバージョン）で行ったため、"
    "XLOOKUPを使う一部の数式（氏名・クラスの自動表示など）はこの検証環境では確認できませんでした。\n"
    "XLOOKUPはMicrosoft 365で標準搭載されている関数のため、実際の学校PCでは問題なく動作する見込みですが、"
    "本手順書の各チェックリストで実機にて最終確認をお願いします。",
    label="検証の限界"
)

add_h1(doc, "8. うまく動かないときは")
add_table(doc,
    headers=["症状", "考えられる原因と対処"],
    rows=[
        ["#NAME? エラーが出る", "Excelのバージョンが古くXLOOKUPに対応していない可能性。Microsoft 365 / Excel 2021以降が必要"],
        ["#REF! エラーが出る", "シート名を変更・削除した可能性。シート名は変更しないでください"],
        ["カナチェックが全部NGになる", "日本語版Excelであることを確認。それでも直らない場合はLENB関数の言語設定を確認"],
        ["マクロが実行できない", "ファイルが .xlsm 形式で保存されているか確認。.xlsx のままだとマクロは保存されません"],
        ["数式が ##### と表示される", "列幅が足りないだけです。列の境界をダブルクリックすると自動調整されます"],
    ],
    col_widths=[5.5, 11.5]
)

add_pagebreak(doc)
add_h1(doc, "9. 最終チェックリスト（学校PC導入前）")
for item in [
    "PC1→PC2のファイル受け渡しシミュレーションで、①③④すべて期待結果どおりになった",
    "VBAマクロをインポートし、.xlsm 形式で保存できた",
    "主要マクロ（行非表示・一括印刷・PDF出力）が動作した",
    "実際の生徒データ・口座データに置き換えても数式が壊れないことを確認した（サンプル行を削除しすぎない）",
    "学校PCのExcelがMicrosoft 365（またはExcel 2021以降）であることを確認した",
    "CoPilot for Microsoft 365 のライセンスがあり、②のプロンプトが実際に使えることを確認した",
]:
    add_check(doc, item)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("以上")
set_japanese_font(r, size=11, bold=True)

import os
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'manual.docx')
doc.save(out_path)
print(f"saved {out_path}")
