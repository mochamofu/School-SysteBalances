# -*- coding: utf-8 -*-
"""図解つき使い方ガイド（PDF元のdocx）を生成するスクリプト

各手順に「その場面で見る画面」の図（docs/animations/_engine.py の
描画部品で作成）を添えた操作マニュアルを組み立てる。
生成: docs/guide.docx（→LibreOfficeでPDF化して docs/使い方ガイド_図解版.pdf）
図は一時フォルダに生成し、docxへ埋め込んだあとは不要（コミットしない）。
"""
import os
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(HERE, "animations"))
import _engine as E  # noqa: E402  (docs/animations の描画部品を流用)

from docx import Document  # noqa: E402
from docx.shared import Pt, Cm, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

FONT = "IPAGothic"
IMG_DIR = tempfile.mkdtemp(prefix="guide_figs_")

# ==================================================================
# 図の描画（アニメーションエンジンの部品を静止画として利用）
# ==================================================================

def _save(img, name):
    path = os.path.join(IMG_DIR, name)
    img.save(path)
    return path


def fig_overview():
    """図1: 全体像（データの流れ）"""
    img, d = E.new_frame()
    d.rectangle([0, 0, E.W, E.H], fill=(0xF7, 0xF9, 0xFC))

    def box(x, y, w, h, title, lines, fill, tcolor=E.NAVY):
        d.rounded_rectangle([x, y, x + w, y + h], radius=10, fill=E.WHITE,
                            outline=E.BORDER, width=2)
        d.rectangle([x, y, x + w, y + 34], fill=fill)
        d.text((x + 12, y + 8), title, font=E.F_H, fill=E.WHITE)
        for i, t in enumerate(lines):
            d.text((x + 14, y + 46 + i * 24), t, font=E.F_SMALL, fill=tcolor)

    box(30, 60, 250, 170, "外から届くデータ", ["・銀行の振替結果", "・掲示用名簿(クラス発表)", "・業者の請求書"], (0x8a, 0x6d, 0x3b))
    box(360, 40, 280, 230, "入力アシスタント", ["⑪ 振替結果取込", "①② 名簿の照合・反映", "④⑤ 支出・収入の一括入力",
                                            "⑫ 年間予定表から転送", "⑧ 整合性チェック"], E.NAVY)
    box(720, 60, 250, 170, "令和○年度生積立金", ["(マスターは今のまま)", "金額・未納印・合計が", "自動で埋まっていく"], (0x1e, 0x6e, 0x3e))
    box(360, 330, 280, 150, "自動で出てくる帳票", ["・支出/収入承認書", "・精算書(PDF/印刷)", "・決算集計"], (0x2f, 0x5f, 0xa8))
    E.draw_arrow(d, (285, 145), (355, 145), color=E.BLUE, width=5)
    E.draw_arrow(d, (645, 145), (715, 145), color=E.BLUE, width=5)
    E.draw_arrow(d, (500, 275), (500, 325), color=E.BLUE, width=5)
    d.text((300, 108), "貼るだけ", font=E.F_SMALL, fill=E.BLUE)
    d.text((652, 108), "ボタン1つ", font=E.F_SMALL, fill=E.BLUE)
    d.text((30, 528), "ポイント: マスターには一切手を加えず、アシスタントが決まった場所に安全に書き込みます。", font=E.F_BODY, fill=E.GRAY)
    d.text((30, 556), "書き込みの前には毎回バックアップが自動で作られます。", font=E.F_BODY, fill=E.GRAY)
    return _save(img, "fig1.png")


def fig_settings():
    """図2: 設定シート"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "積立金入力アシスタント.xlsm ― 設定シート")
    rows = [
        ("マスターファイルの場所", "C:\\Users\\jimu\\Desktop\\練習用_令和X年度生積立金.xlsx", True),
        ("バックアップの保存先", "(空欄でOK)", False),
        ("年度（数字）", "7", True),
        ("精算書PDFの保存先", "(空欄でOK)", False),
        ("口座マスターの場所", "C:\\Users\\jimu\\Desktop\\練習用_口座マスター.xlsx", True),
    ]
    y = 90
    for label, val, hot in rows:
        d.rectangle([60, y, 300, y + 40], fill=E.LBLUE, outline=E.BORDER)
        d.text((72, y + 10), label, font=E.F_BODY, fill=E.NAVY)
        d.rectangle([300, y, 860, y + 40], fill=E.YELLOW, outline=(0xE0, 0xC3, 0x41), width=2)
        d.text((312, y + 10), val, font=E.F_SMALL, fill=(0x33, 0x33, 0x33))
        if hot:
            E.draw_arrow(d, (920, y + 20), (872, y + 20), color=E.RED)
        y += 56
    d.text((60, y + 12), "※黄色いセルが入力欄。矢印の3ヶ所（C3・C5・C7）を自分のPCのパスに合わせるだけで準備完了。",
           font=E.F_BODY, fill=E.GRAY)
    E.callout(d, "最初に1回だけ: 設定シートの黄色いセルを埋める")
    return _save(img, "fig2.png")


def fig_menu():
    """図3: メニュー画面（ボタン12個）"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "積立金入力アシスタント.xlsm ― メニュー")
    groups = [
        ("◆ 名簿・クラス替え", 50, ["① 名簿を解析して照合する", "② クラス替えをマスターに反映する", "③ 新入生としてマスターに登録する"]),
        ("◆ 日々の入力", 360, ["④ 支出をマスターへ一括入力", "⑤ 収入をマスターへ一括入力", "⑥ 収入枠の一覧を表示",
                             "⑪ 振替結果を照合", "⑫ 予定を入力フォームへ転送"]),
        ("◆ 年度末・点検", 670, ["⑦ 決算用の集計を実行", "⑧ マスターの整合性をチェック", "⑨ 精算書を一括PDF保存", "⑩ 精算書を一括印刷"]),
    ]
    for title, x, items in groups:
        d.text((x, 74), title, font=E.F_BODY, fill=E.GRAY)
        for i, t in enumerate(items):
            y = 104 + i * 56
            hot = t.startswith("⑪")
            d.rounded_rectangle([x, y, x + 280, y + 42], radius=8,
                                fill=(0xF5, 0xB7, 0x31) if hot else E.WHITE,
                                outline=(0xD5, 0xDE, 0xED) if not hot else (0xC9, 0x8F, 0x00), width=2)
            d.text((x + 14, y + 11), t, font=E.F_SMALL,
                   fill=(0x12, 0x20, 0x3D) if hot else E.NAVY)
    E.callout(d, "日々の操作はこの12個のボタンだけ（Excelの知識は不要）")
    return _save(img, "fig3.png")


def fig_bank_paste():
    """図4: 振替結果取込シートに貼り付け→⑪実行"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "振替結果取込シート ― 銀行の結果を貼るだけ")
    # 左: 貼り付け領域
    heads = ["口座記号", "口座番号", "金額", "振替結果"]
    for i, h in enumerate(heads):
        E.fill_cell(d, (60 + i * 110, 120, 170 + i * 110, 148), E.NAVY, h, E.WHITE, E.F_SMALL)
    rows = [("10000", "80000000", "76,000", "0"), ("10001", "80000111", "76,000", "0"),
            ("10006", "80000666", "0", "1"), ("10007", "80000777", "76,000", "0")]
    for r, row in enumerate(rows):
        ng = row[3] == "1"
        for i, v in enumerate(row):
            E.fill_cell(d, (60 + i * 110, 148 + r * 28, 170 + i * 110, 176 + r * 28),
                        E.LRED if ng else E.YELLOW, v, E.RED if ng else (0x33, 0x33, 0x33), E.F_SMALL)
    d.text((60, 268), "…(80行つづく)…", font=E.F_SMALL, fill=E.GRAY)
    d.text((60, 92), "①銀行の結果票から4列をコピーして12行目に貼り付け", font=E.F_BODY, fill=E.NAVY)

    # 右: 照合結果
    heads2 = ["精算番号", "氏名", "判定"]
    for i, h in enumerate(heads2):
        E.fill_cell(d, (560 + i * 120, 120, 680 + i * 120, 148), E.NAVY, h, E.WHITE, E.F_SMALL)
    res = [("1", "あおき あおい", "振替済"), ("2", "いしだ いつき", "振替済"),
           ("7", "きたむら くるみ", "未納"), ("8", "くどう けんと", "振替済")]
    for r, row in enumerate(res):
        ng = row[2] == "未納"
        for i, v in enumerate(row):
            E.fill_cell(d, (560 + i * 120, 148 + r * 28, 680 + i * 120, 176 + r * 28),
                        E.LRED if ng else E.LGREEN, v, E.RED if ng else E.GREEN, E.F_SMALL)
    d.text((560, 92), "②「⑪振替結果を照合」で自動判定", font=E.F_BODY, fill=E.NAVY)
    E.draw_arrow(d, (505, 200), (552, 200), color=E.BLUE, width=5)

    # サマリー
    d.rounded_rectangle([560, 300, 940, 400], radius=8, fill=E.LGREEN, outline=E.GREEN, width=2)
    d.text((580, 314), "読取80件 ／ 振替済78名 ／ 未納2名", font=E.F_BODY, fill=E.GREEN)
    d.text((580, 348), "未納2名は「収入入力」シートの未納者表へ", font=E.F_SMALL, fill=E.GREEN)
    d.text((580, 370), "自動で転記されました", font=E.F_SMALL, fill=E.GREEN)
    E.callout(d, "探す作業がゼロに: 未納者は自動で見つかる")
    return _save(img, "fig4.png")


def fig_income():
    """図5: 収入入力（未納者表が自動で埋まった状態）→⑤"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "収入入力シート")
    rows = [("収入枠No（1〜43）", "2"), ("件名", "マスター口座振替 10月分"), ("日付", "2026/10/13"), ("一人あたり金額（円）", "76000")]
    y = 80
    for label, val in rows:
        d.rectangle([60, y, 300, y + 36], fill=E.LBLUE, outline=E.BORDER)
        d.text((72, y + 8), label, font=E.F_SMALL, fill=E.NAVY)
        d.rectangle([300, y, 560, y + 36], fill=E.YELLOW, outline=(0xE0, 0xC3, 0x41), width=2)
        d.text((312, y + 8), val, font=E.F_SMALL, fill=(0x33, 0x33, 0x33))
        y += 44
    d.text((60, y + 8), "―未納者表―（⑪が自動で書いた分。手入力は不要）", font=E.F_BODY, fill=E.NAVY)
    heads = ["精算番号", "氏名（メモ）"]
    for i, h in enumerate(heads):
        E.fill_cell(d, (60 + i * 160, y + 40, 220 + i * 160, y + 68), E.NAVY, h, E.WHITE, E.F_SMALL)
    for r, row in enumerate([("7", "きたむら くるみ"), ("44", "はやし ひなた")]):
        for i, v in enumerate(row):
            E.fill_cell(d, (60 + i * 160, y + 68 + r * 28, 220 + i * 160, y + 96 + r * 28),
                        E.LGREEN, v, E.GREEN, E.F_SMALL)
    btn = E.button(d, 640, 300, 280, 52, "⑤ 収入をマスターへ一括入力", hot=True)
    E.draw_cursor(d, (780, 330), clicked=True)
    d.rounded_rectangle([600, 390, 940, 470], radius=8, fill=E.WHITE, outline=E.NAVY, width=2)
    d.text((620, 404), "入金あり78名 ／ 未納2名", font=E.F_BODY, fill=E.NAVY)
    d.text((620, 436), "未納2名にはマスターのH列に「未納」の印", font=E.F_SMALL, fill=E.GRAY)
    E.callout(d, "78名分の入金がワンクリックでマスターへ")
    return _save(img, "fig5.png")


def fig_expense():
    """図6: 支出入力＋例外表→④"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "支出入力シート")
    rows = [("支出No（1〜100）", "5"), ("件名", "遠足バス代"), ("支払先", "○○観光バス"), ("一人あたり金額（円）", "3500"), ("対象", "全員")]
    y = 80
    for label, val in rows:
        d.rectangle([60, y, 300, y + 34], fill=E.LBLUE, outline=E.BORDER)
        d.text((72, y + 7), label, font=E.F_SMALL, fill=E.NAVY)
        d.rectangle([300, y, 560, y + 34], fill=E.YELLOW, outline=(0xE0, 0xC3, 0x41), width=2)
        d.text((312, y + 7), val, font=E.F_SMALL, fill=(0x33, 0x33, 0x33))
        y += 40
    d.text((60, y + 8), "―例外表―（例外の生徒だけ書く）", font=E.F_BODY, fill=E.NAVY)
    heads = ["精算番号", "氏名（メモ）", "金額（0=対象外）"]
    for i, h in enumerate(heads):
        E.fill_cell(d, (60 + i * 150, y + 40, 210 + i * 150, y + 68), E.NAVY, h, E.WHITE, E.F_SMALL)
    for r, row in enumerate([("45", "転退学", "0"), ("12", "欠席の返金", "-1,040")]):
        for i, v in enumerate(row):
            E.fill_cell(d, (60 + i * 150, y + 68 + r * 28, 210 + i * 150, y + 96 + r * 28),
                        E.YELLOW, v, (0x33, 0x33, 0x33), E.F_SMALL)
    E.button(d, 640, 300, 280, 52, "④ 支出をマスターへ一括入力", hot=True)
    E.draw_cursor(d, (780, 330), clicked=True)
    d.rounded_rectangle([600, 390, 940, 470], radius=8, fill=E.WHITE, outline=E.NAVY, width=2)
    d.text((620, 404), "対象78名（うち例外2名）", font=E.F_BODY, fill=E.NAVY)
    d.text((620, 436), "支出承認書シートも同時に自動作成", font=E.F_SMALL, fill=E.GRAY)
    E.callout(d, "全員分の入力と承認書づくりが同時に終わる")
    return _save(img, "fig6.png")


def fig_roster():
    """図7: 名簿貼付→①→名簿一覧"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "名簿一覧シート ― ①解析の結果")
    heads = ["組", "番号", "氏名", "照合結果", "精算番号"]
    widths = [70, 70, 200, 200, 110]
    x = 60
    xs = []
    for w_, h in zip(widths, heads):
        E.fill_cell(d, (x, 100, x + w_, 130), E.NAVY, h, E.WHITE, E.F_SMALL)
        xs.append((x, w_))
        x += w_
    data = [("2", "1", "あおき あおい", "一致", "17"),
            ("2", "2", "いしだ いつき", "一致", "203"),
            ("2", "3", "うえの うた", "一致", "88"),
            ("2", "4", "えんどう えま", "見つからず", "?"),
            ("2", "5", "おかべ おうすけ", "一致", "156")]
    for r, row in enumerate(data):
        ng = row[3] != "一致"
        for i, v in enumerate(row):
            x0, w_ = xs[i]
            E.fill_cell(d, (x0, 130 + r * 30, x0 + w_, 160 + r * 30),
                        E.LRED if ng else (E.WHITE if r % 2 else (0xFA, 0xFC, 0xFF)),
                        v, E.RED if ng else (0x33, 0x33, 0x33), E.F_SMALL)
    E.draw_arrow(d, (800, 265), (700, 258), color=E.RED)
    d.text((810, 250), "赤い行だけE列に精算番号を手入力", font=E.F_SMALL, fill=E.RED)
    d.rounded_rectangle([60, 330, 700, 400], radius=8, fill=E.LBLUE, outline=E.BLUE, width=2)
    d.text((80, 344), "手順: 掲示用名簿を「名簿貼付」に貼る → ①解析 → 赤い行を確認 → ②反映", font=E.F_SMALL, fill=E.NAVY)
    d.text((80, 370), "80名の照合が数秒。同姓同名や転入生だけ人が判断する", font=E.F_SMALL, fill=E.NAVY)
    E.callout(d, "クラス替え: 320人分の手修正が「貼って2クリック」に")
    return _save(img, "fig7.png")


def fig_plan():
    """図8: 年間予定表→⑫"""
    img, d = E.new_frame()
    body = E.app_window(d, 20, 20, 960, 560, "年間予定表シート")
    heads = ["行No", "予定月", "区分", "No", "件名", "一人あたり"]
    widths = [60, 70, 80, 60, 240, 110]
    x = 60
    xs = []
    for w_, h in zip(widths, heads):
        E.fill_cell(d, (x, 100, x + w_, 130), E.NAVY, h, E.WHITE, E.F_SMALL)
        xs.append((x, w_))
        x += w_
    data = [("1", "4", "収入", "2", "学年積立金 前期", "76,000"),
            ("2", "5", "支出", "5", "遠足バス代", "3,500"),
            ("3", "6", "支出", "6", "芸術鑑賞教室", "1,600"),
            ("4", "9", "収入", "3", "学年積立金 後期", "76,000")]
    for r, row in enumerate(data):
        hot = r == 1
        for i, v in enumerate(row):
            x0, w_ = xs[i]
            E.fill_cell(d, (x0, 130 + r * 30, x0 + w_, 160 + r * 30),
                        (0xFF, 0xE9, 0xB8) if hot else E.YELLOW, v, (0x33, 0x33, 0x33), E.F_SMALL)
    E.button(d, 640, 300, 300, 52, "⑫ 予定を入力フォームへ転送", hot=True)
    E.draw_cursor(d, (790, 330), clicked=True)
    d.rounded_rectangle([60, 320, 560, 420], radius=8, fill=E.LGREEN, outline=E.GREEN, width=2)
    d.text((80, 334), "行No「2」を指定 → 支出入力シートに", font=E.F_SMALL, fill=E.GREEN)
    d.text((80, 360), "No・件名・支払先・金額が自動で入る", font=E.F_SMALL, fill=E.GREEN)
    d.text((80, 386), "→ あとは④を押すだけ（実質2クリック）", font=E.F_SMALL, fill=E.GREEN)
    E.callout(d, "年度初めに計画を1回書けば、あとは呼び出すだけ")
    return _save(img, "fig8.png")


def fig_reports():
    """図9: 点検と帳票"""
    img, d = E.new_frame()
    d.rectangle([0, 0, E.W, E.H], fill=(0xF7, 0xF9, 0xFC))
    # 左: チェック結果
    E.app_window(d, 20, 30, 460, 300, "⑧ チェック結果", icon_color=E.NAVY)
    d.text((44, 90), "【点検1】収入−支出＝残金", font=E.F_SMALL, fill=E.NAVY)
    d.rounded_rectangle([44, 116, 140, 140], radius=12, fill=E.LGREEN)
    d.text((58, 119), "✓ 一致", font=E.F_SMALL, fill=E.GREEN)
    d.text((44, 156), "【点検4】未納の生徒", font=E.F_SMALL, fill=E.NAVY)
    d.text((60, 182), "精算番号 7   2組14番", font=E.F_SMALL, fill=(0x33, 0x33, 0x33))
    d.text((60, 206), "精算番号 44  5組3番", font=E.F_SMALL, fill=(0x33, 0x33, 0x33))
    d.text((44, 250), "→ 督促状づくりの材料にそのまま", font=E.F_SMALL, fill=E.GRAY)
    # 右: 精算書PDF
    E.app_window(d, 520, 30, 460, 300, "⑨ 精算書PDF フォルダ", icon_color=(0x1E, 0x6E, 0x3E))
    for i, name in enumerate(["1-1_あおきあおい.pdf", "1-2_いしだいつき.pdf", "1-3_うえのうた.pdf", "…(78ファイル)"]):
        d.text((560, 90 + i * 32), ("📄 " if i < 3 else "　 ") + name, font=E.F_SMALL, fill=(0x33, 0x33, 0x33))
    d.text((560, 230), "組-番号_氏名.pdf で自動命名", font=E.F_SMALL, fill=E.GRAY)
    # 下: 保管
    d.rounded_rectangle([220, 400, 780, 500], radius=10, fill=E.WHITE, outline=E.NAVY, width=2)
    d.text((250, 420), "マスター ＋ 精算書PDF ＋ チェック結果 をワンセットで", font=E.F_BODY, fill=E.NAVY)
    d.text((250, 452), "保管用PCへ（毎月の締めルーチン）", font=E.F_BODY, fill=E.NAVY)
    E.draw_arrow(d, (250, 350), (330, 395), color=E.BLUE, width=4)
    E.draw_arrow(d, (750, 350), (670, 395), color=E.BLUE, width=4)
    return _save(img, "fig9.png")


# ==================================================================
# docx の組み立て
# ==================================================================

def jp(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = RGBColor(*color)


def title(doc, text, size=20, color=(0x20, 0x38, 0x64)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    jp(p.add_run(text), size=size, bold=True, color=color)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    jp(p.add_run(text), size=15, bold=True, color=(0x20, 0x38, 0x64))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "203864")
    pBdr.append(bottom)
    pPr.append(pBdr)


def body(doc, text, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        jp(p.add_run(line), size=size, color=color)


def step(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    jp(p.add_run(f"手順{num}　"), size=10.5, bold=True, color=(0x20, 0x38, 0x64))
    jp(p.add_run(text), size=10.5)


def figure(doc, path, caption):
    doc.add_picture(path, width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    jp(p.add_run(caption), size=9.5, bold=True, color=(0x44, 0x72, 0xC4))


def build():
    figs = {
        1: fig_overview(), 2: fig_settings(), 3: fig_menu(),
        4: fig_bank_paste(), 5: fig_income(), 6: fig_expense(),
        7: fig_roster(), 8: fig_plan(), 9: fig_reports(),
    }

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = sec.bottom_margin = Cm(1.8)

    # ---- 表紙 ----
    for _ in range(5):
        doc.add_paragraph()
    title(doc, "積立金会計 入力アシスタント", size=24)
    title(doc, "使い方ガイド（図解版）", size=18, color=(0x44, 0x72, 0xC4))
    doc.add_paragraph()
    title(doc, "～ どの画面で・何を入れて・どのボタンを押すか、図で分かる ～", size=11, color=(0x60, 0x60, 0x60))
    for _ in range(3):
        doc.add_paragraph()
    figure(doc, figs[1], "図1　全体像: 届いたデータが、ボタン操作だけでマスターと帳票になる")
    doc.add_page_break()

    # ---- 1. 準備 ----
    h1(doc, "1．最初の準備（1回だけ・約10分）")
    step(doc, 1, "「積立金入力アシスタント.xlsm」を開き、「コンテンツの有効化」を押す")
    step(doc, 2, "「設定」シートの黄色いセルに、マスター（C3）・年度（C5）・口座マスター（C7）を入力")
    figure(doc, figs[2], "図2　設定シート: 黄色いセル3ヶ所を埋めるだけ")
    body(doc, "設定が済むと、以後の操作はすべて「メニュー」シートのボタンだけで行えます。")
    figure(doc, figs[3], "図3　メニュー画面: 業務は12個のボタンに対応")
    doc.add_page_break()

    # ---- 2. 毎月: 振替結果 ----
    h1(doc, "2．毎月の口座振替 ― 貼って、押して、終わり")
    step(doc, 1, "銀行から届いた振替結果の4列（口座記号・口座番号・金額・振替結果）をコピー")
    step(doc, 2, "「振替結果取込」シートの12行目に貼り付け、メニューの「⑪振替結果を照合」を押す")
    figure(doc, figs[4], "図4　振替結果取込: 貼り付けると生徒を自動特定し、未納だけ赤く出る")
    step(doc, 3, "「収入入力」シートを開くと未納者表がもう埋まっている。枠No・件名・日付・金額を入れて「⑤」を押す")
    figure(doc, figs[5], "図5　収入入力: 未納者表は自動。⑤で78名分が一括記録")
    body(doc, "未納の生徒はマスターのH列に「未納」の印が自動で立ち、⑧のチェックや督促状づくりにそのまま使えます。")
    doc.add_page_break()

    # ---- 3. 支出 ----
    h1(doc, "3．支出があったとき ― 例外だけ書けばよい")
    step(doc, 1, "「支出入力」シートに 支出No・件名・支払先・一人あたり金額・対象（全員）を入力")
    step(doc, 2, "転退学・給付型・返金など例外の生徒だけ、下の例外表に精算番号と金額を書く（0=対象外、マイナス=返金）")
    step(doc, 3, "メニューの「④支出をマスターへ一括入力」を押す")
    figure(doc, figs[6], "図6　支出入力: 全員分の記録と支出承認書の作成が同時に終わる")
    doc.add_page_break()

    # ---- 4. クラス替え ----
    h1(doc, "4．毎年4月のクラス替え ― 名簿を貼って2クリック")
    step(doc, 1, "掲示用名簿のシート全体をコピーし、「名簿貼付」シートのA1に貼り付ける")
    step(doc, 2, "メニューの「①名簿を解析して照合する」を押す → 名簿一覧に照合結果が出る")
    step(doc, 3, "赤い行（見つからず・同姓同名）だけE列に精算番号を手入力し、「②クラス替えをマスターに反映する」")
    figure(doc, figs[7], "図7　名簿一覧: ほぼ全員が自動で「一致」。人が見るのは赤い行だけ")
    doc.add_page_break()

    # ---- 5. 年間予定表 ----
    h1(doc, "5．年間予定表 ― 計画は1回、実行は2クリック")
    step(doc, 1, "年度初めに「年間予定表」シートへ、月・区分・No・件名・一人あたり金額を書いておく")
    step(doc, 2, "実行するとき、メニューの「⑫予定を入力フォームへ転送」で行Noを指定")
    step(doc, 3, "支出入力（または収入入力）が自動で埋まるので、④または⑤を押す")
    figure(doc, figs[8], "図8　年間予定表: 請求書からの打ち直しがなくなる")
    doc.add_page_break()

    # ---- 6. 点検と帳票 ----
    h1(doc, "6．締めの点検と帳票 ― 月1回のルーチン")
    step(doc, 1, "「⑧マスターの整合性をチェック」→ 金額のズレ・処理漏れ・未納者一覧を自動点検")
    step(doc, 2, "「⑨精算書を一括PDF保存」→ 生徒別PDF（組-番号_氏名.pdf）が一括生成")
    step(doc, 3, "マスター・精算書PDF・チェック結果をワンセットで保管用PCへ")
    figure(doc, figs[9], "図9　点検と帳票: 保管まで含めて毎月同じ流れで終わる")

    # ---- 7. 困ったとき ----
    h1(doc, "7．困ったとき")
    body(doc, "・ボタンがない → Alt+F8 →「初期設定」を実行\n"
              "・「設定が足りません」→ 設定シートのC3（マスター）やC7（口座マスター）のパスを確認\n"
              "・「不明口座」が出る → 口座マスターにその口座を追記して⑪を再実行\n"
              "・間違えて実行した → 「バックアップ」フォルダの直近のコピーを元の名前に戻す\n"
              "・マクロがブロックされる → ファイルを右クリック→プロパティ→「許可する」にチェック")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    jp(p.add_run("以上"), size=11, bold=True)

    out = os.path.join(HERE, "guide.docx")
    doc.save(out)
    print("saved", out)


if __name__ == "__main__":
    build()
